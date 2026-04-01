"""
Revision after peer review — action-map-driven edit application.

Pipeline ordering:
  A. Triage each reviewer comment (via RevisionActionMap)
  B. Decide manuscript action for each comment
  C. Apply edits to manuscript (deterministic string matching)
  D. Generate point-by-point response letter (AFTER edits are applied)
  E. Run verification (consistency audit)

Each concern produces a structured response with:
  - response_category: Accepted | Partially Accepted | Respectfully Declined | Already Addressed
  - exact manuscript changes via replace/insert_after/delete operations
  - verification that changes match the action map
"""

from __future__ import annotations

import json
import logging
import re
from typing import TYPE_CHECKING

from models import AuditRecommendedEdit, PeerReviewReport, RepairTask, ResponseQCResult, RevisionActionMap, RevisionRepairTelemetry, RevisionResult
from services.article_builder import build_revision_writing_requirements
from services.manuscript_utils import (
    apply_manuscript_changes,
    audit_revision,
    build_compact_evidence,
    build_full_manuscript_context,
    build_manuscript_chunk_coverage_context,
    build_relevant_passage_context,
    build_section_index,
    locate_text_span,
    validate_change_operations,
)

if TYPE_CHECKING:
    from models import PaperSummary
    from services.ai_provider import AIProvider

logger = logging.getLogger(__name__)


# ── Post-revision structural integrity check ─────────────────────────────────

def _validate_revised_manuscript(original: str, revised: str) -> list[str]:
    """Check that revision didn't corrupt the manuscript structure.

    Returns list of critical warnings. If non-empty, the revision should be
    reverted to the original.
    """
    warnings: list[str] = []
    if not revised or not revised.strip():
        warnings.append("Revised manuscript is empty")
        return warnings

    # 1. Major section headings preserved
    orig_headings = set(re.findall(r'^#{1,3}\s+(.+)$', original, re.MULTILINE))
    revised_headings = set(re.findall(r'^#{1,3}\s+(.+)$', revised, re.MULTILINE))
    missing = orig_headings - revised_headings
    # Allow minor heading changes (e.g., "## Discussion" still present even if slightly rephrased)
    critical_missing = {h for h in missing if any(
        kw in h.lower() for kw in ('abstract', 'introduction', 'method', 'result', 'discussion', 'conclusion', 'reference')
    )}
    if critical_missing:
        warnings.append(f"Critical section headings lost: {', '.join(sorted(critical_missing))}")

    # 2. References section preserved
    has_orig_refs = bool(re.search(r'(?im)^\s*#{1,6}\s*(references|bibliography)', original))
    has_revised_refs = bool(re.search(r'(?im)^\s*#{1,6}\s*(references|bibliography)', revised))
    if has_orig_refs and not has_revised_refs:
        warnings.append("References section was deleted by revision")

    # 3. Word count within ±20%
    orig_words = len(original.split())
    revised_words = len(revised.split())
    if orig_words > 100:
        ratio = revised_words / orig_words
        if ratio < 0.8:
            warnings.append(f"Word count dropped by {(1-ratio)*100:.0f}% ({orig_words} → {revised_words})")
        elif ratio > 1.2:
            warnings.append(f"Word count increased by {(ratio-1)*100:.0f}% ({orig_words} → {revised_words})")

    # 4. No embedded AI instructions
    instruction_patterns = [
        r'\[INSTRUCTION[^\]]*\]',
        r'NOTE:\s*Do not include line numbers',
        r'IMPORTANT:\s*Copy only',
        r'Return ONLY the JSON',
    ]
    for pat in instruction_patterns:
        if re.search(pat, revised, re.IGNORECASE):
            warnings.append(f"Embedded AI instruction leaked into manuscript: {pat}")

    # 5. Workflow residue and obvious prose-collision artifacts
    residue_patterns = [
        (r'\[FUT\]', "Workflow marker [FUT] remains in manuscript"),
        (
            r'Complete the citation(?: formatting later)?(?: and restore the remainder of the paragraph)?',
            "Editorial placeholder text remains in manuscript",
        ),
        (r'\.ilable evidence(?:\.ilable evidence)+', "Duplicated suffix artifact remains in manuscript"),
        (r'\.r narratives(?:\.r narratives)?', "Collided repeated fragment '.r narratives' remains in manuscript"),
        (r'summaries\s+\[CITE:', "Workflow residue 'summaries [CITE:' remains in manuscript"),
    ]
    for pat, message in residue_patterns:
        if re.search(pat, revised, re.IGNORECASE):
            warnings.append(message)

    return warnings


def _repair_issue_type(*parts: str) -> str:
    text = " ".join(p for p in parts if p).lower()
    if "placeholder" in text:
        return "placeholder"
    if any(term in text for term in ("truncat", "incomplete", "structural", "coherence")):
        return "structural"
    if any(term in text for term in ("corrupt", "duplicat", "artifact")):
        return "corruption"
    if "citation" in text or "reference" in text:
        return "citation"
    return "other"


def _repair_rewrite_scope(issue_type: str) -> str:
    if issue_type in {"placeholder", "structural", "corruption"}:
        return "paragraph"
    if issue_type == "citation":
        return "sentence"
    return "paragraph"


def _review_concern_by_id(review: PeerReviewReport, concern_id: str):
    prefix, _, raw_index = concern_id.partition("_")
    if not raw_index.isdigit():
        return None
    idx = int(raw_index) - 1
    if prefix == "major" and 0 <= idx < len(review.major_concerns):
        return review.major_concerns[idx]
    if prefix == "minor" and 0 <= idx < len(review.minor_concerns):
        return review.minor_concerns[idx]
    return None


def _serialize_repair_tasks(tasks: list[RepairTask]) -> str:
    lines: list[str] = []
    for i, task in enumerate(tasks, 1):
        location = task.section or "Unknown section"
        line_span = ""
        if len(task.line_span) == 2 and all(task.line_span):
            line_span = f" | lines {task.line_span[0]}-{task.line_span[1]}"
        passage = f'\n  QUOTED PASSAGE: "{task.quoted_passage}"' if task.quoted_passage else ""
        safe_ops = ""
        if task.safe_edit_ops:
            safe_ops = "\n  SAFE EDIT OPS: " + json.dumps(
                [op.model_dump(mode="json") for op in task.safe_edit_ops],
                ensure_ascii=False,
            )
        lines.append(
            f"[task_{i}] source={task.source} severity={task.severity} scope={task.rewrite_scope}\n"
            f"  LOCATION: {location}{line_span}\n"
            f"  ISSUE TYPE: {task.issue_type}\n"
            f"  EXPECTED OUTCOME: {task.expected_outcome}"
            f"{passage}{safe_ops}"
        )
    return "\n\n".join(lines) or "(none)"


def _repair_task_group_key(task: RepairTask) -> tuple[str, ...]:
    quoted_passage = re.sub(r"\s+", " ", task.quoted_passage).strip().lower()
    if quoted_passage:
        return ("passage", quoted_passage)
    if len(task.line_span) == 2 and all(task.line_span):
        return (
            "lines",
            task.section.strip().lower(),
            task.issue_type.strip().lower(),
            str(task.line_span[0]),
            str(task.line_span[1]),
        )
    return (
        "fallback",
        task.section.strip().lower(),
        task.issue_type.strip().lower(),
        re.sub(r"\s+", " ", task.expected_outcome).strip().lower(),
    )


def _merge_repair_task_list(tasks: list[RepairTask]) -> list[RepairTask]:
    merged: dict[tuple[str, ...], RepairTask] = {}
    scope_rank = {"sentence": 1, "paragraph": 2, "multi_paragraph": 3, "section": 4}

    for task in tasks:
        key = _repair_task_group_key(task)
        current = merged.get(key)
        if not current:
            merged[key] = task.model_copy(deep=True)
            continue

        source_values = [value.strip() for value in f"{current.source},{task.source}".split(",") if value.strip()]
        current.source = ", ".join(dict.fromkeys(source_values))
        if task.severity == "blocking":
            current.severity = "blocking"
        if not current.section and task.section:
            current.section = task.section
        if not current.quoted_passage and task.quoted_passage:
            current.quoted_passage = task.quoted_passage
        if not current.line_span and task.line_span:
            current.line_span = list(task.line_span)
        elif len(current.line_span) == 2 and len(task.line_span) == 2:
            current.line_span = [
                min(current.line_span[0], task.line_span[0]),
                max(current.line_span[1], task.line_span[1]),
            ]
        if scope_rank.get(task.rewrite_scope, 0) > scope_rank.get(current.rewrite_scope, 0):
            current.rewrite_scope = task.rewrite_scope

        outcomes = [item.strip() for item in current.expected_outcome.split(" | ") if item.strip()]
        if task.expected_outcome and task.expected_outcome not in outcomes:
            outcomes.append(task.expected_outcome)
        current.expected_outcome = " | ".join(outcomes) if outcomes else current.expected_outcome

        seen_ops = {
            (op.edit_type, op.find, op.replace_with)
            for op in current.safe_edit_ops
        }
        for op in task.safe_edit_ops:
            op_key = (op.edit_type, op.find, op.replace_with)
            if op_key in seen_ops:
                continue
            current.safe_edit_ops.append(op)
            seen_ops.add(op_key)

    return list(merged.values())


def _validate_repair_tasks(
    manuscript_text: str,
    tasks: list[RepairTask],
) -> tuple[list[RepairTask], list[RepairTask]]:
    valid: list[RepairTask] = []
    invalid: list[RepairTask] = []
    for task in tasks:
        if task.quoted_passage:
            span = locate_text_span(manuscript_text, task.quoted_passage)
            if not span:
                invalid.append(task)
                continue
            task.quoted_passage = span["text"]
            task.section = task.section or span["section"]
            if not task.line_span:
                task.line_span = [span["start_line"], span["end_line"]]
        if task.safe_edit_ops:
            safe_ops, _unsafe_ops = validate_change_operations(
                manuscript_text,
                [
                    {
                        "type": op.edit_type,
                        "find": op.find,
                        "anchor": op.find if op.edit_type == "insert_after" else "",
                        "replace_with": op.replace_with,
                        "text": op.replace_with,
                    }
                    for op in task.safe_edit_ops
                ],
            )
            task.safe_edit_ops = [
                AuditRecommendedEdit(
                    edit_type=str(op.get("type", "replace")),
                    find=str(op.get("find", "") or op.get("anchor", "")),
                    replace_with=str(op.get("replace_with", "") or op.get("text", "")),
                )
                for op in safe_ops
            ] if safe_ops else []
        if task.severity == "blocking" and not task.quoted_passage and not task.safe_edit_ops:
            invalid.append(task)
            continue
        valid.append(task)
    return valid, invalid


def _collect_repair_tasks(
    review: PeerReviewReport,
    manuscript_text: str,
    consistency_audit: dict | None,
    re_review: dict | None,
    editorial_review: dict | None,
) -> tuple[list[RepairTask], RevisionRepairTelemetry]:
    tasks: list[RepairTask] = []
    if consistency_audit:
        for item in (consistency_audit.get("repair_tasks") or []):
            try:
                tasks.append(RepairTask.model_validate(item))
            except Exception:
                continue
    if re_review:
        for item in (re_review.get("repair_tasks") or []):
            try:
                tasks.append(RepairTask.model_validate(item))
            except Exception:
                continue
        for res in (re_review.get("concern_resolutions") or []):
            if not isinstance(res, dict):
                continue
            status = str(res.get("status", "")).strip().lower()
            if status not in {"unresolved", "partially_resolved"}:
                continue
            concern = _review_concern_by_id(review, str(res.get("concern_id", "")))
            quoted_passage = getattr(concern, "quoted_passage", "") if concern else ""
            tasks.append(RepairTask(
                source="rereview",
                severity="blocking" if status == "unresolved" else "advisory",
                section=getattr(concern, "location", "") if concern else "",
                line_span=[],
                quoted_passage=quoted_passage,
                issue_type=_repair_issue_type(str(res.get("original_concern", "")), str(res.get("explanation", ""))),
                expected_outcome=str(res.get("explanation", "") or res.get("original_concern", "")),
                safe_edit_ops=[],
                rewrite_scope=_repair_rewrite_scope(_repair_issue_type(str(res.get("original_concern", "")), str(res.get("explanation", "")))),
            ))
    if editorial_review:
        for item in (editorial_review.get("repair_tasks") or []):
            try:
                tasks.append(RepairTask.model_validate(item))
            except Exception:
                continue
    merged = _merge_repair_task_list(tasks)
    valid, invalid = _validate_repair_tasks(manuscript_text, merged)
    discarded_blockers = sum(1 for task in invalid if task.severity == "blocking")
    logger.info(
        "Repair task normalization: raw=%d grouped=%d valid=%d invalid=%d discarded_blockers=%d",
        len(tasks),
        len(merged),
        len(valid),
        len(invalid),
        discarded_blockers,
    )
    if invalid:
        logger.warning(
            "Discarded %d invalid QA finding(s) before follow-up revision",
            len(invalid),
        )
    return valid, RevisionRepairTelemetry(
        invalid_qa_findings=len(invalid),
        discarded_blockers=discarded_blockers,
        merged_repair_groups=len(merged),
        structural_repair_invocations=0,
    )


# ── Call 1: Generate edits driven by the action map ──────────────────────────

_EDITS_SYSTEM = """\
You are a senior academic revision specialist with deep experience preparing
manuscripts for high-impact journals (Nature, Science, The Lancet, PLOS Medicine,
BMJ, JAMA). You have received three inputs:

  (A) The PEER-REVIEW REPORT — reviewer comments, numbered by severity.
  (B) The REVISION ACTION MAP — a triage and edit plan for each concern,
      including action_type, target location, and verification_criterion.
  (C) The ORIGINAL MANUSCRIPT — shown with line numbers for reference only.

Your task: produce precise, minimal, publication-ready edits that faithfully
implement the action map while preserving every aspect of the manuscript that
is not under revision.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
§1  ABSOLUTE PRESERVATION RULES
    Violating any of these constitutes a critical, disqualifying error.
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

 1. TITLE — Never alter the manuscript title unless the reviewer explicitly
    requests a title change.

 2. SECTION HEADINGS — Preserve all section and subsection headings verbatim.
    Do not rename, merge, split, reorder, or insert headings unless the
    reviewer specifically instructs it.

 3. FIGURES & ILLUSTRATIONS — Do not add, remove, relocate, or modify
    figure references (e.g., "Figure 1", "Fig. 2", "Supplementary Fig. S3"),
    figure blocks, captions, alt text, image syntax, or placeholders
    unless the reviewer explicitly targets a specific figure.

 4. TABLES — Do not add, remove, or modify table references (e.g., "Table 1"),
    table content, column headers, row data, footnotes, legends, or any
    markdown/HTML table markup unless the reviewer explicitly targets a
    specific table.

 5. CITATIONS & REFERENCES — Do not add, remove, rewrite, reorder, or
    renumber any citation or reference entry unless:
      (a) the reviewer explicitly asks for a new citation, or
      (b) the reviewer flags an incorrect/missing reference, or
      (c) the action map contains an "add_citation" action.
    Preserve the entire References section character-for-character unless
    a specific entry must be updated.

 6. WORD COUNT — Keep total word count within ±5 % of the original. Never
    pad prose with filler. Never delete content the reviewer did not flag.

 7. SCOPE DISCIPLINE — Change ONLY the text that directly addresses an
    accepted reviewer concern. Every other character stays identical.

 8. NO UNSOLICITED EDITS — Do not fix grammar, spelling, punctuation, or
    phrasing outside the targeted edit region. Do not beautify, reorganize,
    modernize, or "improve" anything the reviewer did not raise.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
§2  PROSE QUALITY — WRITING AT JOURNAL STANDARD
    Every word you add or rewrite must be indistinguishable from the
    surrounding manuscript in voice, register, and argumentative rigour.
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

VOICE MATCHING
- Before writing, read two full paragraphs around the edit site.
  Match sentence length distribution, hedging conventions, passive/active
  balance, and technical vocabulary density.
- If the manuscript uses British conventions (behaviour, analysed, centre),
  continue in British English — and vice versa for American English.
- If the manuscript favours the authorial "we" ("We hypothesised…"),
  maintain that perspective. If it uses impersonal constructions
  ("It was hypothesised…"), preserve those.
- Adopt the same connective tissue the authors use. If they write
  "Notably, …" or "Of particular relevance, …", work within that
  register rather than importing your own preferred transitions.

ARGUMENTATION
- State claims with the same epistemic calibration the authors use.
  In a Lancet paper, "Our findings suggest…" carries different weight
  from "Our findings demonstrate…". Match the existing gradient.
- When the reviewer asks you to strengthen or temper a claim, adjust
  the SPECIFIC verb or hedge — do not restructure the sentence.
    WRONG rewrite: "This unequivocally proves that X causes Y."
    RIGHT rewrite: "These findings provide strong evidence that X is
    associated with Y" (if the manuscript already uses associational
    language).
- Every causal or directional claim must align with the study design.
  Cross-sectional data supports "association", not "effect". An RCT
  supports "effect" but only within the studied population.

CONCISION
- High-impact journals prize economy. Do not inflate a 12-word fix
  into a 40-word paragraph.
- Prefer a single, well-constructed sentence over two that say the
  same thing with different words.
- Delete throat-clearing ("It is important to note that…",
  "It should be mentioned that…", "It is worth highlighting that…").
  Start with the substance.

LIMITATION ACKNOWLEDGEMENTS (special guidance)
- When a reviewer requests acknowledgement of a limitation, write
  ONE surgically placed sentence at the most logical location —
  typically within the Limitations subsection or the paragraph
  discussing the specific methodological choice.
- Weave the caveat INTO the surrounding argument; do not append
  it as a standalone disclaimer.
- NEVER scatter echoes of the same caveat across multiple sections.
  One acknowledgement, one location, fully integrated.
    BAD:  "A limitation is that the sample size was small."
          (dropped as a free-standing sentence)
    GOOD: "The modest sample (n = 87), while adequately powered for
          the primary endpoint, may have limited our ability to detect
          smaller subgroup effects."
          (integrated, specific, properly qualified)

ANTI-PATTERNS — never produce any of the following:
- Boilerplate hedges copy-pasted across paragraphs:
  "future research is needed", "this is a limitation",
  "caution should be exercised in interpreting…" reappearing verbatim.
- Reviewer language or action-map instructions leaking into the manuscript.
- Defensive or apologetic tone ("We apologise for this oversight…").
  Revisions should read as confident improvements, not confessions.
- Inflated claims that overreach beyond the study design.
- Weasel phrases that evacuate all meaning ("This might potentially
  perhaps be somewhat related…").
- New jargon or acronyms not already established in the manuscript.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
§3  STRUCTURAL COHERENCE AFTER EDITS
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

- After each edit, mentally re-read the paragraph. If the insert breaks
  logical flow or creates a non sequitur, adjust the seam — the one
  sentence before and after the insertion — to restore continuity.
  Do NOT rewrite the whole paragraph.
- If two reviewer-driven edits land in the same paragraph, ensure they
  do not contradict each other or repeat the same point.
- Preserve existing paragraph boundaries. Do not merge or split
  paragraphs unless the edit makes it unavoidable (e.g., the reviewer
  asked to restructure an argument).

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
§4  LINE NUMBER RULE
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

The manuscript is displayed with line numbers: "  42  text here".
These numbers are for reference only.

When you copy text into "find" or "anchor", copy ONLY the text content —
strip the leading spaces and line number.

  WRONG:  "  42  Psychology is a discipline"
  RIGHT:  "Psychology is a discipline"

Any "find" or "anchor" value that begins with digits + spaces WILL FAIL
to match and the edit will be silently lost.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
§5  CITATIONS AND REFERENCES — DETAILED PROTOCOL
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

EXISTING CITATIONS
- [CITE:key] markers are valid manuscript-internal placeholders.
  Preserve them exactly unless the reviewer explicitly targets that
  citation. Final bibliography normalisation happens at export.

ADDING NEW CITATIONS (when action_type is "add_citation" or new
evidence must be cited in revised text):
- Use the EVIDENCE CORPUS (provided separately) to retrieve the correct
  authors, year, title, journal, volume, pages, and DOI.
- Match the citation style already used in the manuscript:
    • Numbered style [1], [2]: insert the new reference at the correct
      position and renumber all subsequent entries.
    • Author-year style (Smith et al., 2023): insert the in-text tag
      at the appropriate clause boundary and add the full entry to the
      References in alphabetical order.
    • [CITE:key] placeholder style: insert [CITE:new_key] in-text and
      note the full reference details in your changes_summary so the
      export pipeline can resolve it.
- Every in-text citation MUST have a matching entry in the References
  section (or in the placeholder registry if the manuscript uses
  [CITE:key] workflow).
- NEVER fabricate a citation. If the evidence corpus does not contain the
  paper, flag it in changes_summary rather than inventing bibliographic
  details.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
§6  RESPONSE CATEGORIES — choose the most accurate label
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

"Accepted"
  — The concern was valid and fully addressed. Manuscript text changed.

"Partially Accepted"
  — Some aspects of the concern were valid and addressed; others were
    not applicable or beyond the scope of this revision. Explain what
    was done and what was not, and why.

"Respectfully Declined"
  — The concern was carefully considered but is not supported by the
    evidence or conflicts with the study's stated scope. Provide a
    concrete, evidence-backed rebuttal — not a vague dismissal.

"Already Addressed"
  — The manuscript already handles this concern at a specific,
    citable location. Quote the relevant passage in the summary.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
§7  EDIT OPERATIONS — JSON SCHEMA
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

For each action in the map, produce a JSON object with exactly these keys:

  "reviewer_comment_id"  — matches the action map ID (e.g., "major_1")
  "response_category"    — one of the four categories above
  "manuscript_changes"   — a JSON array of edit operations (see below)
  "changes_summary"      — 1-2 sentences: what was changed, where, and
                           why (include line-range references)

Edit operations (inside "manuscript_changes"):

  REPLACE — swap existing text with new text
  {{
    "type": "replace",
    "find": "<EXACT text from manuscript, ≥ 50 chars>",
    "replace_with": "<new text>"
  }}

  INSERT AFTER — add new text immediately after an anchor passage
  {{
    "type": "insert_after",
    "anchor": "<EXACT text from manuscript, ≥ 50 chars>",
    "text": "<new text to insert>"
  }}

  DELETE — remove a passage entirely
  {{
    "type": "delete",
    "find": "<EXACT text to remove, ≥ 50 chars>"
  }}

  Use an empty array [] if no text change is needed (e.g., for
  "no_change_rebut" actions or "Already Addressed" responses).

STRING-MATCHING RULES:
- "find" and "anchor" values MUST be copied CHARACTER-FOR-CHARACTER from
  the manuscript text. Do NOT paraphrase, abridge, or approximate.
- Minimum length: 50 characters, to guarantee unique matching.
- Include enough surrounding context to be unambiguous if the target
  phrase appears more than once.
- Prefer ending your "find" string at a sentence boundary to avoid
  splitting a sentence across edit operations.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
§8  DECISION HEURISTICS — WHEN IN DOUBT
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

- Smaller edit beats larger edit. A five-word substitution that satisfies
  the verification_criterion is superior to a paragraph rewrite.
- If two phrasings are equally valid, choose the one that preserves more
  of the original wording.
- If a reviewer comment is ambiguous, interpret it conservatively —
  address the most concrete, actionable reading.
- If the action map says "no_change_rebut", produce an empty
  manuscript_changes array and a substantive, respectful rebuttal
  in changes_summary.
- Never fabricate data, effect sizes, p-values, confidence intervals,
  or sample sizes.
- Never introduce a claim that is not supported by the study's own data
  or by a citation from the evidence corpus.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

OUTPUT FORMAT:
Return a JSON array — one object per action map entry, in the same order
as the action map:

[
  {{
    "reviewer_comment_id": "major_1",
    "response_category": "Accepted",
    "manuscript_changes": [
      {{
        "type": "replace",
        "find": "exact original text (≥50 chars)…",
        "replace_with": "revised text…"
      }}
    ],
    "changes_summary": "Replaced the causal claim in the Discussion (Lines 204-206) with associational language consistent with the cross-sectional design."
  }},
  {{
    "reviewer_comment_id": "minor_2",
    "response_category": "Already Addressed",
    "manuscript_changes": [],
    "changes_summary": "The requested sample-size justification already appears in the Methods (Lines 98-102): 'A priori power analysis (G*Power 3.1) indicated…'."
  }}
]"""


_EDITS_USER_TMPL = """\
Research topic: {query}
Target journal: {journal}

{manuscript_context}

REVISION ACTION MAP:
{action_map_text}

EVIDENCE CORPUS ({n} papers):
{evidence_json}

Generate the manuscript_changes for every action in the map. Return ONLY the JSON array."""


# ── Call 2: Generate structured point-by-point response data ─────────────────

_RESPONSE_SYSTEM = """\
You are drafting a journal-quality point-by-point response to reviewers.

Your job is to produce structured per-comment response data that will be formatted
into a professional response letter with a table layout.

Rules:
- Address every reviewer comment individually.
- For each comment, restate the reviewer's concern fairly and concisely.
- Provide a detailed, scholarly reply addressing the concern directly.
  Reference specific evidence, describe exactly what was changed.
- Describe changes made in the manuscript with section and line references.
  If wording was changed, provide the new wording.
- If the suggestion was not followed, explain why respectfully and substantively.
- Never claim a change was made unless it appears in the revised manuscript.
- Maintain a respectful, non-defensive tone.
- Do not over-thank. Be concise and precise.
- Prefer "We revised…" over vague phrasing like "This has been improved."
- Distinguish between textual clarification, substantive reframing, added evidence,
  and softened claims.

Also produce a "novelty_summary" (2-3 sentences about the novelty of the research
and main changes in the revised version) and a "major_changes_list" (3-5 bullet
points summarizing the most significant changes and additions).

Output a JSON object:
{{
  "novelty_summary": "2-3 sentences about the novelty and main improvements",
  "major_changes_list": [
    "Added framework paragraph defining success criteria for adversarial collaboration (Discussion, Lines 145-160)",
    "Softened overclaiming in Introduction regarding theory elimination (Lines 23-28)",
    "..."
  ],
  "responses": [
    {{
      "reviewer_comment_id": "major_1",
      "reviewer_number": 1,
      "comment_number": 1,
      "reviewer_comment": "Restated reviewer concern — concise, fair restatement",
      "author_reply": "Detailed scholarly reply (2-4 sentences). References evidence, states agreement or rationale for disagreement, describes changes.",
      "changes_done": "Section X, Lines Y-Z: [description of change]. New wording: '...'"
    }}
  ]
}}

CRITICAL:
- "reviewer_comment" must restate the concern in the reviewer's voice, not yours.
- "author_reply" is the substantive response — explain reasoning, cite evidence.
- "changes_done" describes the actual manuscript change with location. Use "" (empty)
  if no change was made, and explain why in author_reply instead.
- Group comments by reviewer_number (1, 2, etc.). For single-reviewer reports, use 1.
"""


_RESPONSE_USER_TMPL = """\
Journal: {journal}
Manuscript title: {manuscript_title}

ORIGINAL REVIEW CONCERNS:
{review_concerns}

APPLIED EDITS SUMMARY:
{edits_summary}

REVISED MANUSCRIPT:
{revised_manuscript}

Generate the structured JSON response. Address every concern."""


_FOLLOWUP_EDITS_SYSTEM = """\
You are an expert academic revision assistant performing a FOLLOW-UP revision pass.

This is not a rewrite. A first revision already exists. Your task is to make the
smallest justified manuscript edits needed to address internal QA findings from:
- consistency audit
- re-review
- editorial review

MANDATORY RULES:
- Every edit must be justified by at least one explicit finding ID provided below.
- Do NOT revisit resolved areas unless a finding specifically identifies them.
- Preserve title, headings, figures, illustrations, tables, citations, references,
  and overall structure unless a finding explicitly requires a change.
- Preserve [CITE:key] grounding markers as valid temporary citations when present.
- Keep Abstract revisions citation-free.
- Do NOT rewrite broadly. Apply minimal, surgical corrections only.
- Do NOT generate any point-by-point reply text in this step.

Return ONLY a JSON array:
[
  {
    "justification_ids": ["audit_1", "rereview_2"],
    "changes_summary": "What was fixed and where",
    "manuscript_changes": [
      {"type": "replace", "find": "<EXACT text>", "replace_with": "<new text>"},
      {"type": "insert_after", "anchor": "<EXACT text>", "text": "<new text>"},
      {"type": "delete", "find": "<EXACT text>"}
    ]
  }
]

Rules for manuscript_changes:
- "find" and "anchor" must be copied exactly from the manuscript.
- Do not include line numbers in copied text.
- Use [] when no manuscript change is warranted for a finding.
- Each operation must be directly tied to the cited justification_ids.
"""


_FOLLOWUP_EDITS_AGGRESSIVE_SYSTEM = """\
You are an expert academic revision assistant performing an AGGRESSIVE follow-up revision.

Previous surgical edit attempts failed to resolve the blocking issues below. You must now
make SUBSTANTIVE changes — rewrite full paragraphs, restructure arguments, soften or
remove unsupported claims, and add missing content where needed.

CONTEXT: The blocking issues have persisted across multiple rounds of minimal edits.
The AI reviewer keeps flagging the same problems. You MUST resolve them decisively.

MANDATORY RULES:
- Each edit MUST directly address one or more of the blocking findings below.
- You MAY rewrite entire paragraphs or sections to fully resolve an issue.
- If a claim is flagged as unsupported or overstated, SOFTEN or REMOVE it — do not
  merely add a hedge word. Restructure the argument.
- If the manuscript is missing positive scope conditions, ADD a dedicated paragraph.
- If the manuscript is truncated, COMPLETE the missing sections.
- Preserve [CITE:key] grounding markers.
- Keep Abstract citation-free.
- Do NOT generate any point-by-point reply text in this step.

Return ONLY a JSON array:
[
  {
    "justification_ids": ["audit_1", "rereview_2"],
    "changes_summary": "What was fixed and where",
    "manuscript_changes": [
      {"type": "replace", "find": "<EXACT text to replace — can be a full paragraph>", "replace_with": "<new text>"},
      {"type": "insert_after", "anchor": "<EXACT text>", "text": "<new text>"},
      {"type": "delete", "find": "<EXACT text>"}
    ]
  }
]

Rules for manuscript_changes:
- "find" and "anchor" must be copied exactly from the manuscript.
- For paragraph-level rewrites, copy the FULL paragraph as "find".
- Use [] when no manuscript change is warranted for a finding.
- Each operation must be directly tied to the cited justification_ids.
"""


_FOLLOWUP_EDITS_USER_TMPL = """\
Journal: {journal}

CURRENT MANUSCRIPT:
{manuscript_context}

JUSTIFIED FOLLOW-UP FINDINGS:
{findings_text}

Apply only the minimal justified manuscript fixes. Return ONLY the JSON array."""


# ── Table-format markdown builder ────────────────────────────────────────────

def _build_revision_system_prompt(
    base_system: str,
    article_type: str,
    journal_style: object | None,
) -> str:
    shared_rules = build_revision_writing_requirements(
        article_type=article_type or "review",
        journal_style=journal_style if hasattr(journal_style, "get_sections") else None,
    )
    if not shared_rules:
        return base_system
    return f"{base_system}\n\n{shared_rules}"


def _build_table_format_md(
    response_data: dict,
    journal: str = "",
    manuscript_title: str = "",
    raw_fallback: str = "",
) -> str:
    """Build a table-format markdown point-by-point response matching the journal template."""
    responses = response_data.get("responses", [])
    if not responses and not isinstance(response_data, dict):
        return raw_fallback  # Fallback to raw text if parsing failed

    novelty = response_data.get("novelty_summary", "")
    major_changes = response_data.get("major_changes_list", [])

    lines: list[str] = [
        "# Point-by-Point Response",
        "",
        f"**Manuscript Title:** {manuscript_title}" if manuscript_title else "",
        f"**Journal:** {journal}" if journal else "",
        "",
        "Dear Editor,",
        "",
    ]

    if novelty:
        lines += [novelty, ""]
    else:
        lines += [
            "We sincerely thank the reviewer(s) for their constructive feedback. "
            "We have carefully addressed each point below.",
            "",
        ]

    if major_changes:
        lines += [
            "**Major changes and additions to the revised manuscript:**",
            "",
        ]
        for i, change in enumerate(major_changes, 1):
            lines.append(f"{i}. {change}")
        lines.append("")

    lines += ["---", "", "**Specific Responses:**", ""]

    # Group responses by reviewer number
    by_reviewer: dict[int, list[dict]] = {}
    for r in responses:
        rn = r.get("reviewer_number", 1)
        by_reviewer.setdefault(rn, []).append(r)

    for reviewer_num in sorted(by_reviewer.keys()):
        reviewer_responses = by_reviewer[reviewer_num]
        lines += [
            f"## Response to Reviewer #{reviewer_num}",
            "",
            "We thank the Reviewer for their valuable time and constructive comments.",
            "",
            "| S.No | Reviewer's Comments | Reply from the Author | Changes Done |",
            "|------|--------------------|-----------------------|-------------|",
        ]
        for r in reviewer_responses:
            sno = r.get("comment_number", "")
            comment = r.get("reviewer_comment", "").replace("|", "\\|").replace("\n", " ")
            reply = r.get("author_reply", "").replace("|", "\\|").replace("\n", " ")
            changes = r.get("changes_done", "").replace("|", "\\|").replace("\n", " ")
            if not changes:
                changes = "No manuscript change required."
            lines.append(f"| {sno} | {comment} | {reply} | {changes} |")
        lines += ["", ""]

    lines += [
        "Sincerely,",
        "The Authors",
    ]

    return "\n".join(lines)


def build_response_letter_docx(
    response_data: dict,
    journal: str = "",
    manuscript_title: str = "",
) -> bytes:
    """
    Build a professional .docx point-by-point response letter matching the
    reference template format with colored table columns.

    Returns the docx file as bytes.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()

    # Page setup: Letter, 1 inch margins
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    BLUE = RGBColor(0x1D, 0x48, 0x95)
    DARK_BLUE = RGBColor(0x00, 0x20, 0x60)
    BROWN = RGBColor(0x83, 0x3C, 0x0B)
    BLACK = RGBColor(0x00, 0x00, 0x00)
    FONT_SIZE = Pt(11)

    def add_heading_para(text: str, bold: bool = True, color: RGBColor = BLUE, size: Pt = Pt(14)):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.color.rgb = color
        run.font.size = size
        run.font.name = "Calibri"
        return p

    def add_body_para(text: str, bold: bool = False, color: RGBColor = BLACK, italic: bool = False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.color.rgb = color
        run.font.size = FONT_SIZE
        run.font.name = "Calibri"
        run.italic = italic
        return p

    def set_cell_text(cell, text: str, bold: bool = False, color: RGBColor = BLACK):
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = bold
        run.font.color.rgb = color
        run.font.size = Pt(10)
        run.font.name = "Calibri"

    # ── Title ─────────────────────────────────────────────────────────────
    add_heading_para("Point-by-Point Response", size=Pt(16))

    # ── Metadata ──────────────────────────────────────────────────────────
    if manuscript_title:
        add_body_para(f"Title of Article: {manuscript_title}", bold=True, color=BLUE)
    if journal:
        add_body_para(f"Journal: {journal}", bold=True, color=BLUE)
    add_body_para("")

    # ── Dear Editor ───────────────────────────────────────────────────────
    add_body_para("Dear Editor,")
    add_body_para("")

    novelty = response_data.get("novelty_summary", "")
    if novelty:
        add_body_para(novelty)
    else:
        add_body_para(
            "We sincerely thank the reviewer(s) for their constructive feedback. "
            "We have carefully addressed each point below."
        )

    # ── Major changes ─────────────────────────────────────────────────────
    major_changes = response_data.get("major_changes_list", [])
    if major_changes:
        add_body_para("")
        add_body_para("Major changes and additions to the revised manuscript:", bold=True, color=BLUE)
        for i, change in enumerate(major_changes, 1):
            add_body_para(f"{i}. {change}")

    # ── Specific Responses ────────────────────────────────────────────────
    add_body_para("")
    add_heading_para("Specific Responses:", size=Pt(12))

    responses = response_data.get("responses", [])
    by_reviewer: dict[int, list[dict]] = {}
    for r in responses:
        rn = r.get("reviewer_number", 1)
        by_reviewer.setdefault(rn, []).append(r)

    for reviewer_num in sorted(by_reviewer.keys()):
        reviewer_responses = by_reviewer[reviewer_num]

        add_heading_para(f"Response to Reviewer #{reviewer_num}:", size=Pt(12))
        add_body_para(
            "We thank the Reviewer for their valuable time and constructive comments.",
            italic=True,
        )
        add_body_para("")

        # Build the table
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"

        # Column widths (approximate)
        for i, width in enumerate([Cm(1.2), Cm(4.5), Cm(6.5), Cm(4.5)]):
            for cell in table.columns[i].cells:
                cell.width = width

        # Header row
        hdr = table.rows[0]
        set_cell_text(hdr.cells[0], "S.No", bold=True, color=BLACK)
        set_cell_text(hdr.cells[1], "Reviewer's Comments", bold=True, color=BLACK)
        set_cell_text(hdr.cells[2], "Reply from the Author", bold=True, color=DARK_BLUE)
        set_cell_text(hdr.cells[3], "Changes Done", bold=True, color=BROWN)

        # Header row shading
        for cell in hdr.cells:
            shading = cell._element.get_or_add_tcPr()
            shading_elem = shading.makeelement(qn('w:shd'), {
                qn('w:val'): 'clear',
                qn('w:color'): 'auto',
                qn('w:fill'): 'D6E4F0',
            })
            shading.append(shading_elem)

        # Data rows
        for r in reviewer_responses:
            row = table.add_row()
            sno = str(r.get("comment_number", ""))
            comment = r.get("reviewer_comment", "")
            reply = r.get("author_reply", "")
            changes = r.get("changes_done", "")
            if not changes:
                changes = "No manuscript change required."

            set_cell_text(row.cells[0], sno, bold=True)
            set_cell_text(row.cells[1], comment)
            set_cell_text(row.cells[2], reply, color=DARK_BLUE)
            set_cell_text(row.cells[3], changes, color=BROWN)

        add_body_para("")

    # ── Closing ───────────────────────────────────────────────────────────
    add_body_para("")
    add_body_para("Sincerely,")
    add_body_para("The Authors")

    # Save to bytes
    from io import BytesIO
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _normalize_ws(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def _verify_final_response(
    review: PeerReviewReport,
    revised_article: str,
    response_data: dict,
) -> ResponseQCResult:
    """Lightweight response QA against the finalized manuscript state."""
    responses = response_data.get("responses", []) if isinstance(response_data, dict) else []
    if not isinstance(responses, list):
        responses = []

    blocking_issues: list[str] = []
    advisory_issues: list[str] = []
    normalized_article = _normalize_ws(revised_article)
    expected_count = len(review.major_concerns) + len(review.minor_concerns)

    if expected_count and len(responses) < expected_count:
        blocking_issues.append(
            f"Final response covers {len(responses)} reviewer comments, but {expected_count} comments were expected."
        )

    quote_patterns = [
        r'New wording:\s*["“\'](.{15,400}?)["”\']',
        r'Revised text:\s*["“\'](.{15,400}?)["”\']',
    ]
    change_verbs = ("revised", "added", "inserted", "updated", "changed", "clarified", "expanded")

    for idx, item in enumerate(responses, 1):
        if not isinstance(item, dict):
            continue
        changes_done = str(item.get("changes_done", "")).strip()
        author_reply = str(item.get("author_reply", "")).strip()
        combined = f"{changes_done}\n{author_reply}"

        for pattern in quote_patterns:
            for match in re.finditer(pattern, combined, flags=re.IGNORECASE | re.DOTALL):
                snippet = _normalize_ws(match.group(1))
                if snippet and snippet not in normalized_article:
                    blocking_issues.append(
                        f"Response item {idx} cites new wording that does not appear in the finalized manuscript."
                    )

        if any(verb in author_reply.lower() for verb in change_verbs) and not changes_done:
            advisory_issues.append(
                f"Response item {idx} describes manuscript changes but does not list a concrete change location."
            )

    if blocking_issues:
        summary = "Final response QA found blocking mismatches between the response letter and the finalized manuscript."
    elif advisory_issues:
        summary = "Final response QA passed with advisory cautions about response specificity."
    else:
        summary = "Final response QA passed. No phantom claims were detected against the finalized manuscript."

    return ResponseQCResult(
        checked=True,
        blocking_issues=blocking_issues,
        advisory_issues=advisory_issues,
        summary=summary,
    )


# ── Conflict resolution ──────────────────────────────────────────────────────

_CONFLICT_SYSTEM = """\
You are applying specific text edits to an academic manuscript. Some edits failed
because the exact text was not found (likely due to minor whitespace or formatting
differences).

For each failed operation, find the closest matching text in the manuscript and
produce the corrected operation.

CRITICAL: Copy all unchanged paragraphs VERBATIM. Only alter text explicitly
targeted by the change operations below. Do NOT improve, rephrase, or reorganize
any text outside the targeted regions.

Output: a JSON array of corrected operations (same format as input)."""

_CONFLICT_USER_TMPL = """\
MANUSCRIPT (excerpt around failed operations):
{manuscript_excerpt}

FAILED OPERATIONS:
{failed_ops_json}

Return corrected operations as a JSON array, or [] if they should be skipped."""


# ── Main function ─────────────────────────────────────────────────────────────

async def generate_revision_package(
    provider: "AIProvider",
    summaries: list["PaperSummary"],
    query: str,
    article: str,
    review: PeerReviewReport,
    journal: str = "",
    *,
    article_type: str = "review",
    journal_style: object | None = None,
    manuscript_packs: dict | None = None,
    word_limit: int = 4000,
    manuscript_title: str = "",
    action_map: RevisionActionMap | None = None,
    generate_response_letter: bool = True,
) -> RevisionResult:
    """
    Generate a revision package using the 4-stage pipeline:
    1. Action map drives what edits to make (already provided or generated)
    2. AI generates manuscript_changes operations per action
    3. Operations applied deterministically (exact string matching)
    4. Point-by-point response letter generated from actual edits
    5. Post-revision audit checks for quality regressions
    """
    # ── Build manuscript context with line numbers ────────────────────────
    section_index = build_section_index(article) if article else []
    manuscript_context = build_full_manuscript_context(
        manuscript_text=article,
        section_index=section_index,
    )

    # ── Build action map text for the edit prompt ─────────────────────────
    action_map_lines: list[str] = []
    if action_map and action_map.actions:
        for a in action_map.actions:
            passage_line = f'  Target text: "{a.quoted_passage}"\n' if a.quoted_passage else ""
            action_map_lines.append(
                f"[{a.reviewer_comment_id}] {a.concern_title}\n"
                f"  Action: {a.action_type}\n"
                f"  Location: {a.manuscript_location}\n"
                f"{passage_line}"
                f"  Target section: {a.target_section}\n"
                f"  Instruction: {a.revision_instruction}\n"
                f"  Edit size: {a.estimated_edit_size}\n"
                f"  Verification: {a.verification_criterion}"
            )
    else:
        # Fallback: build from review concerns directly
        for i, c in enumerate(review.major_concerns, 1):
            passage_line = f'  Target text: "{c.quoted_passage}"\n' if getattr(c, 'quoted_passage', '') else ""
            action_map_lines.append(
                f"[major_{i}] {c.concern}\n"
                f"  Location: {c.location}\n"
                f"{passage_line}"
                f"  Revision request: {c.revision_request}\n"
                f"  Satisfaction criterion: {c.satisfaction_criterion}"
            )
        for i, c in enumerate(review.minor_concerns, 1):
            passage_line = f'  Target text: "{c.quoted_passage}"\n' if getattr(c, 'quoted_passage', '') else ""
            action_map_lines.append(
                f"[minor_{i}] {c.concern}\n"
                f"  Location: {c.location}\n"
                f"{passage_line}"
                f"  Revision request: {c.revision_request}"
            )
    action_map_text = "\n\n".join(action_map_lines) or "(No actions)"

    # ── Compact evidence ─────────────────────────────────────────────────
    evidence = build_compact_evidence(summaries, max_papers=30, max_results=6)

    # ── Call 1: Generate edits driven by the action map ───────────────────
    edits_prompt = _EDITS_USER_TMPL.format(
        query=query or "general academic research",
        journal=journal or "Not specified",
        manuscript_context=manuscript_context,
        action_map_text=action_map_text,
        n=len(evidence),
        evidence_json=json.dumps(evidence, indent=2),
    )
    edits_system = _build_revision_system_prompt(_EDITS_SYSTEM, article_type, journal_style)

    raw_edits = await provider.complete(
        system=edits_system,
        user=edits_prompt,
        json_mode=True,
        temperature=0.15,
        max_tokens=16384,
    )

    # Parse the edits array
    edit_responses: list[dict] = []
    try:
        raw_cleaned = re.sub(r'^```(?:json)?\s*', '', raw_edits.strip())
        raw_cleaned = re.sub(r'\s*```$', '', raw_cleaned.strip())
        parsed = json.loads(raw_cleaned)
        if isinstance(parsed, list):
            edit_responses = parsed
        elif isinstance(parsed, dict) and "responses" in parsed:
            edit_responses = parsed["responses"]
    except json.JSONDecodeError:
        logger.warning("Failed to parse revision edits JSON")

    # ── Collect all manuscript_changes operations ────────────────────────
    all_changes: list[dict] = []
    for resp in edit_responses:
        mc = resp.get("manuscript_changes", [])
        if isinstance(mc, str):
            try:
                mc = json.loads(mc)
            except (json.JSONDecodeError, TypeError):
                mc = []
        if isinstance(mc, list):
            all_changes.extend(mc)

    # ── Apply changes deterministically ──────────────────────────────────
    revised_article = article
    applied_ops: list[dict] = []
    failed_ops: list[dict] = []

    if all_changes:
        revised_article, applied_ops, failed_ops = apply_manuscript_changes(
            article, all_changes,
        )
        logger.info(
            "Revision applied: %d operations succeeded, %d failed",
            len(applied_ops), len(failed_ops),
        )

    # Log failed ops — they'll be retried in next audit round if needed
    if failed_ops:
        logger.warning(
            "Revision: %d operations failed to apply (will be retried in next round if blocking)",
            len(failed_ops),
        )

    # ── Structural integrity check ───────────────────────────────────────
    integrity_warnings = _validate_revised_manuscript(article, revised_article)
    if integrity_warnings:
        for w in integrity_warnings:
            logger.error("REVISION INTEGRITY FAILURE: %s", w)
        # Revert to original — a corrupted revision is worse than no revision
        logger.error("Reverting to original manuscript due to integrity failures")
        revised_article = article
        applied_ops = []
        failed_ops = list(all_changes)

    response_data: dict = {}
    point_by_point = ""
    response_qc: ResponseQCResult | None = None
    if generate_response_letter:
        point_by_point, response_data, response_qc = await generate_final_response_letter(
            provider=provider,
            review=review,
            revised_article=revised_article,
            journal=journal,
            manuscript_title=manuscript_title,
            action_map=action_map,
            change_justifications=[
                str(resp.get("changes_summary", "")).strip()
                for resp in edit_responses
                if str(resp.get("changes_summary", "")).strip()
            ],
        )

    # ── Strip leaked drafting tags before audit sees them ──────────────
    from services.manuscript_citation_formatter import strip_evidence_purpose_tags
    revised_article = strip_evidence_purpose_tags(revised_article)

    # ── Post-revision quality audit ──────────────────────────────────────
    audit_result = audit_revision(article, revised_article)
    if audit_result["warnings"]:
        for w in audit_result["warnings"]:
            logger.warning("Revision audit: %s", w)

    return RevisionResult(
        revised_article=revised_article,
        point_by_point_reply=point_by_point,
        response_data=response_data if response_data else None,
        action_map=action_map,
        audit=audit_result,
        applied_changes=len(applied_ops),
        failed_changes=len(failed_ops),
        change_justifications=[
            str(resp.get("changes_summary", "")).strip()
            for resp in edit_responses
            if str(resp.get("changes_summary", "")).strip()
        ],
        response_qc=response_qc,
    )


async def apply_followup_revision(
    provider: "AIProvider",
    article: str,
    review: PeerReviewReport,
    *,
    journal: str = "",
    article_type: str = "review",
    journal_style: object | None = None,
    action_map: RevisionActionMap | None = None,
    consistency_audit: dict | None = None,
    re_review: dict | None = None,
    editorial_review: dict | None = None,
    aggressive: bool = False,
    user_guidance: str = "",
    last_known_good_article: str = "",
) -> RevisionResult:
    """Apply a minimal second-pass revision driven by audit/re-review/editor findings."""
    section_index = build_section_index(article) if article else []
    repair_tasks, repair_telemetry = _collect_repair_tasks(
        review=review,
        manuscript_text=article,
        consistency_audit=consistency_audit,
        re_review=re_review,
        editorial_review=editorial_review,
    )
    if not repair_tasks:
        return RevisionResult(
            revised_article=article,
            point_by_point_reply="",
            response_data=None,
            action_map=action_map,
            audit=audit_revision(article, article),
            applied_changes=0,
            failed_changes=0,
            change_justifications=[],
            repair_telemetry=repair_telemetry,
        )

    # ── Phase 1: Apply validated safe edit ops directly ───────────────────
    working_article = article
    pre_applied: list[dict] = []
    pre_failed: list[dict] = []
    precomputed_ops: list[dict] = []
    precomputed_justifications: list[str] = []
    for task in repair_tasks:
        for edit in task.safe_edit_ops:
            op: dict = {"type": edit.edit_type, "find": edit.find}
            if edit.edit_type == "insert_after":
                op["anchor"] = edit.find
                op["text"] = edit.replace_with
            else:
                op["replace_with"] = edit.replace_with
            precomputed_ops.append(op)
            precomputed_justifications.append(f"{task.source}: {task.expected_outcome}")

    applied_safe_targets: set[str] = set()
    if precomputed_ops:
        logger.info(
            "Applying %d validated safe QA edits before AI followup",
            len(precomputed_ops),
        )
        working_article, pre_applied, pre_failed = apply_manuscript_changes(
            article, precomputed_ops,
        )
        applied_safe_targets = {
            str(op.get("find", "") or op.get("anchor", "")).strip()
            for op in pre_applied
            if str(op.get("find", "") or op.get("anchor", "")).strip()
        }

    # ── Phase 2: Send remaining findings to AI ────────────────────────────
    remaining_tasks = []
    for task in repair_tasks:
        if task.safe_edit_ops and any(op.find in applied_safe_targets for op in task.safe_edit_ops):
            continue
        remaining_tasks.append(task)

    ai_applied: list[dict] = []
    ai_failed: list[dict] = []
    ai_justifications: list[str] = []

    if remaining_tasks:
        structural_task_count = sum(
            1 for task in remaining_tasks if task.issue_type in {"structural", "placeholder", "corruption"}
        )
        structural_repair_mode = structural_task_count > 0
        if structural_repair_mode:
            repair_telemetry.structural_repair_invocations += 1
            logger.info(
                "Structural repair mode activated for %d repair task(s)",
                structural_task_count,
            )
        base_system = _FOLLOWUP_EDITS_AGGRESSIVE_SYSTEM if aggressive or structural_repair_mode else _FOLLOWUP_EDITS_SYSTEM
        followup_system = _build_revision_system_prompt(base_system, article_type, journal_style)
        passages = [task.quoted_passage for task in remaining_tasks if task.quoted_passage]
        if not passages and action_map:
            passages = [
                action.quoted_passage
                for action in (action_map.actions or [])
                if action.quoted_passage
            ]
        updated_context = build_relevant_passage_context(
            manuscript_text=working_article,
            passages=passages,
            section_index=build_section_index(working_article) if working_article else section_index,
            neighbor_paragraphs=1,
            heading="CURRENT MANUSCRIPT EXCERPTS RELEVANT TO REPAIR TASKS",
        ) if passages else build_full_manuscript_context(
            manuscript_text=working_article,
            section_index=build_section_index(working_article) if working_article else section_index,
        )

        user_prompt = _FOLLOWUP_EDITS_USER_TMPL.format(
            journal=journal or "Not specified",
            manuscript_context=updated_context,
            findings_text=_serialize_repair_tasks(remaining_tasks),
        )
        if user_guidance:
            user_prompt += f"\n\nAUTHOR GUIDANCE:\n{user_guidance}"
        if structural_repair_mode:
            user_prompt += (
                "\n\nSTRUCTURAL REPAIR MODE:\n"
                "- For structural, placeholder, or corruption tasks, rewrite the full affected paragraph block.\n"
                "- Do not emit fragment-only edits for those tasks.\n"
                "- Prefer one complete paragraph-level replace over multiple tiny edits when the same passage is damaged.\n"
            )

        raw_edits = await provider.complete(
            system=followup_system,
            user=user_prompt,
            json_mode=True,
            temperature=0.15 if aggressive else 0.1,
            max_tokens=16384,
        )

        parsed: list[dict] = []
        try:
            cleaned = re.sub(r'^```(?:json)?\s*', '', raw_edits.strip())
            cleaned = re.sub(r'\s*```$', '', cleaned.strip())
            loaded = json.loads(cleaned)
            if isinstance(loaded, list):
                parsed = loaded
        except json.JSONDecodeError:
            logger.warning("Failed to parse follow-up revision JSON")

        ai_changes: list[dict] = []
        for item in parsed:
            if not isinstance(item, dict):
                continue
            summary = str(item.get("changes_summary", "")).strip()
            ids = [str(x).strip() for x in (item.get("justification_ids") or []) if str(x).strip()]
            if summary:
                ai_justifications.append(f"{', '.join(ids) if ids else 'justified'}: {summary}")
            ops = item.get("manuscript_changes", [])
            if isinstance(ops, list):
                ai_changes.extend(ops)

        if ai_changes:
            safe_ai_changes, unsafe_ai_changes = validate_change_operations(
                working_article,
                ai_changes,
            )
            for op in unsafe_ai_changes:
                ai_failed.append({**op, "reason": str(op.get("reason", "unsafe AI edit"))})
            if safe_ai_changes:
                working_article, ai_applied, apply_failed = apply_manuscript_changes(
                    working_article, safe_ai_changes,
                )
                ai_failed.extend(apply_failed)

    # ── Combine results ───────────────────────────────────────────────────
    all_justifications = precomputed_justifications + ai_justifications
    total_applied = len(pre_applied) + len(ai_applied)
    total_failed = len(pre_failed) + len(ai_failed)

    if pre_applied:
        logger.info(
            "Precomputed edits: %d applied, %d failed; AI edits: %d applied, %d failed",
            len(pre_applied), len(pre_failed), len(ai_applied), len(ai_failed),
        )

    # ── Structural integrity check ───────────────────────────────────────
    integrity_warnings = _validate_revised_manuscript(article, working_article)
    if integrity_warnings:
        for w in integrity_warnings:
            logger.error("FOLLOWUP REVISION INTEGRITY FAILURE: %s", w)
        logger.error("Reverting followup to pre-edit state")
        working_article = last_known_good_article or article
        total_applied = 0
        total_failed = len(pre_applied) + len(ai_applied) + len(pre_failed) + len(ai_failed)

    audit_result = audit_revision(article, working_article)
    return RevisionResult(
        revised_article=working_article,
        point_by_point_reply="",
        response_data=None,
        action_map=action_map,
        audit=audit_result,
        applied_changes=total_applied,
        failed_changes=total_failed,
        change_justifications=all_justifications,
        repair_telemetry=repair_telemetry,
    )


async def generate_final_response_letter(
    provider: "AIProvider",
    review: PeerReviewReport,
    revised_article: str,
    *,
    journal: str = "",
    manuscript_title: str = "",
    action_map: RevisionActionMap | None = None,
    change_justifications: list[str] | None = None,
) -> tuple[str, dict, ResponseQCResult]:
    """Generate the final point-by-point response once the manuscript is complete."""
    edits_summary_lines: list[str] = []
    if action_map and action_map.actions:
        for a in action_map.actions:
            edits_summary_lines.append(
                f"[{a.reviewer_comment_id}] {a.action_type}: {a.revision_instruction}"
            )
    if change_justifications:
        edits_summary_lines.extend(change_justifications)
    edits_summary = "\n".join(edits_summary_lines) or "(No edits summary provided)"

    review_concern_lines: list[str] = []
    for i, c in enumerate(review.major_concerns, 1):
        review_concern_lines.append(f"[major_{i}] MAJOR: {c.concern}")
        if c.location:
            review_concern_lines.append(f"  Location: {c.location}")
    for i, c in enumerate(review.minor_concerns, 1):
        review_concern_lines.append(f"[minor_{i}] MINOR: {c.concern}")
        if c.location:
            review_concern_lines.append(f"  Location: {c.location}")
    for i, r in enumerate(review.required_revisions, 1):
        review_concern_lines.append(f"[required_{i}] REQUIRED: {r}")

    relevant_passages = []
    if action_map:
        relevant_passages.extend(
            a.quoted_passage
            for a in (action_map.actions or [])
            if a.quoted_passage
        )
    relevant_passages.extend(
        getattr(c, "quoted_passage", "")
        for c in [*review.major_concerns, *review.minor_concerns]
        if getattr(c, "quoted_passage", "")
    )

    response_prompt = _RESPONSE_USER_TMPL.format(
        journal=journal or "Not specified",
        manuscript_title=manuscript_title or "Untitled",
        review_concerns="\n".join(review_concern_lines) or "(none)",
        edits_summary=edits_summary,
        revised_manuscript=build_relevant_passage_context(
            revised_article,
            passages=relevant_passages,
            section_index=build_section_index(revised_article) if revised_article else [],
            neighbor_paragraphs=1,
            heading="REVISED MANUSCRIPT EXCERPTS RELEVANT TO REVIEWER CONCERNS",
        ) if revised_article else "(No manuscript)",
    )

    raw_response = await provider.complete(
        system=_RESPONSE_SYSTEM,
        user=response_prompt,
        json_mode=True,
        temperature=0.15,
        max_tokens=8192,
    )

    response_data: dict = {}
    try:
        raw_cleaned = re.sub(r'^```(?:json)?\s*', '', raw_response.strip())
        raw_cleaned = re.sub(r'\s*```$', '', raw_cleaned.strip())
        response_data = json.loads(raw_cleaned)
        if not isinstance(response_data, dict):
            response_data = {}
    except json.JSONDecodeError:
        logger.warning("Failed to parse response letter JSON, falling back to raw text")

    point_by_point = _build_table_format_md(
        response_data=response_data,
        journal=journal,
        manuscript_title=manuscript_title,
        raw_fallback=raw_response,
    )
    response_qc = _verify_final_response(
        review=review,
        revised_article=revised_article,
        response_data=response_data if response_data else {},
    )
    return point_by_point, (response_data if response_data else {}), response_qc


async def _resolve_conflicts(
    provider: "AIProvider",
    manuscript: str,
    failed_ops: list[dict],
) -> list[dict]:
    """Send failed operations back to the LLM for conflict resolution."""
    excerpts: list[str] = []
    for op in failed_ops[:10]:
        find_text = op.get("find", "") or op.get("anchor", "")
        if find_text and len(find_text) > 20:
            snippet = find_text[:40]
            idx = manuscript.lower().find(snippet[:20].lower())
            if idx != -1:
                start = max(0, idx - 200)
                end = min(len(manuscript), idx + len(find_text) + 200)
                excerpts.append(f"--- EXCERPT (chars {start}-{end}) ---\n{manuscript[start:end]}")

    manuscript_excerpt = "\n\n".join(excerpts) if excerpts else "\n\n".join(
        build_manuscript_chunk_coverage_context(
            manuscript,
            section_index=build_section_index(manuscript) if manuscript else [],
            max_chars=8000,
            overlap_paragraphs=1,
        )
    )

    raw = await provider.complete(
        system=_CONFLICT_SYSTEM,
        user=_CONFLICT_USER_TMPL.format(
            manuscript_excerpt=manuscript_excerpt,
            failed_ops_json=json.dumps(failed_ops[:10], indent=2),
        ),
        json_mode=True,
        temperature=0.1,
    )

    try:
        raw_cleaned = re.sub(r'^```(?:json)?\s*', '', raw.strip())
        raw_cleaned = re.sub(r'\s*```$', '', raw_cleaned.strip())
        parsed = json.loads(raw_cleaned)
        if isinstance(parsed, list):
            return parsed
    except json.JSONDecodeError:
        pass
    return []
