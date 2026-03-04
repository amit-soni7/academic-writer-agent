"""
services/revision_docx_builder.py

Build three .docx artefacts for a revision round:
  A. Point-by-point reply (4-column table, blue responses, red action taken)
  B. Clean revised manuscript (markdown → Word)
  C. Track-changes manuscript (real OOXML w:ins / w:del via lxml + difflib)

All packages used (python-docx, lxml, difflib) ship with the existing venv.
"""

from __future__ import annotations

import difflib
import io
import re
from datetime import datetime, timezone
from typing import Optional


# ── Manuscript parser (structure-preserving) ──────────────────────────────────

def _parse_manuscript(md_text: str) -> list[tuple[str, str]]:
    """
    Parse markdown into a list of (style, plain_text) tuples.
      style     : 'Heading 1'..'Heading 4', 'Normal', or '' (blank line)
      plain_text: text with inline markdown stripped
    """
    def _clean(text: str) -> str:
        text = re.sub(r'\[CITE:[^\]]*\]|\[CK\]|\[INF\]', '', text)
        text = re.sub(r'\*\*([^*]*)\*\*', r'\1', text)
        text = re.sub(r'\*([^*]*)\*', r'\1', text)
        return text.strip()

    heading_re = re.compile(r'^(#{1,6})\s+(.*)')
    result: list[tuple[str, str]] = []

    for line in md_text.splitlines():
        stripped = line.strip()
        if not stripped:
            result.append(('', ''))
            continue
        m = heading_re.match(stripped)
        if m:
            level = min(len(m.group(1)), 4)
            result.append((f'Heading {level}', _clean(m.group(2))))
        else:
            result.append(('Normal', _clean(stripped)))
    return result


def _tokenize_text(text: str) -> list[str]:
    """
    Split text into alternating word and whitespace tokens.
    'Hello, world foo' → ['Hello,', ' ', 'world', ' ', 'foo']
    Preserves original spacing so no artificial spaces are injected.
    """
    return re.findall(r'\S+|\s+', text)


# ── python-docx helpers ────────────────────────────────────────────────────────

def _heading_level(line: str) -> Optional[int]:
    m = re.match(r'^(#{1,6})\s+', line)
    return len(m.group(1)) if m else None


def _markdown_to_doc(doc, md_text: str):
    """Convert a markdown string into paragraphs/headings in the given Document."""
    for line in md_text.splitlines():
        lvl = _heading_level(line)
        if lvl is not None:
            text = line.lstrip('#').strip()
            doc.add_heading(text, level=min(lvl, 4))
        elif line.strip() == '':
            doc.add_paragraph('')
        else:
            p = doc.add_paragraph()
            # Bold spans
            parts = re.split(r'(\*\*[^*]+\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)


# ── A. Point-by-point reply ────────────────────────────────────────────────────

def _set_cell_width(cell, twips: int) -> None:
    """Set a table cell's fixed width in twips (1440 twips = 1 inch)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn('w:tcW')):
        tcPr.remove(existing)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(twips))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


def _fill_change_column(cell, action_taken: str, diff_str: str, copy_paste_text: str = "", citation_suggestions: list[str] | None = None) -> None:
    """
    Fill the 'Manuscript Change Instructions' table cell with colour-coded
    instructions that tell the author exactly what to ADD / DELETE / CHANGE,
    at which section and approximate line, with the specific text.
    """
    import json as _json
    from docx.shared import Pt, RGBColor

    RED   = RGBColor(0xC0, 0x10, 0x10)
    GREEN = RGBColor(0x0D, 0x6A, 0x35)
    AMBER = RGBColor(0x92, 0x40, 0x09)
    DARK  = RGBColor(0x1E, 0x29, 0x3B)
    GRAY  = RGBColor(0x64, 0x74, 0x8B)

    # Parse the diff JSON (may be string-encoded or a dict)
    diff: dict = {}
    try:
        raw = _json.loads(diff_str or '{}')
        diff = raw if isinstance(raw, dict) else _json.loads(raw)
    except Exception:
        pass

    deleted = (diff.get('deleted') or '').strip()
    added   = (diff.get('added')   or '').strip()
    no_diff = not deleted and not added

    # No text diff — show action_taken or "no change"
    if no_diff:
        no_action = not action_taken or action_taken.lower().startswith('no change')
        r = cell.paragraphs[0].add_run(
            'No manuscript change required.' if no_action else action_taken
        )
        r.font.color.rgb = GRAY if no_action else DARK
        r.font.size = Pt(9)
        r.italic = no_action
        return

    # Parse "Location: description" from action_taken
    location = description = ''
    if action_taken:
        if ':' in action_taken:
            loc_part, desc_part = action_taken.split(':', 1)
            location    = loc_part.strip()
            description = desc_part.strip()
        else:
            location = action_taken.strip()

    # Action label + colour
    if deleted and added:
        action_label, lbl_color = 'CHANGE', AMBER
    elif deleted:
        action_label, lbl_color = 'DELETE', RED
    else:
        action_label, lbl_color = 'ADD', GREEN

    # Line 1: "ACTION   Section, paragraph X, Lines Y–Z"
    p1 = cell.paragraphs[0]
    rl = p1.add_run(action_label + '  ')
    rl.bold = True; rl.font.color.rgb = lbl_color; rl.font.size = Pt(9)
    if location:
        rlo = p1.add_run(location)
        rlo.bold = True; rlo.font.color.rgb = DARK; rlo.font.size = Pt(9)

    # Line 2: italic description of what changes
    if description:
        p2 = cell.add_paragraph()
        r2 = p2.add_run(description)
        r2.italic = True; r2.font.color.rgb = GRAY; r2.font.size = Pt(8)

    # Line 3: specific text to delete
    if deleted:
        p3 = cell.add_paragraph()
        r3a = p3.add_run('DELETE: ')
        r3a.bold = True; r3a.font.color.rgb = RED; r3a.font.size = Pt(9)
        r3b = p3.add_run(f'"{deleted}"')
        r3b.italic = True; r3b.font.color.rgb = RED; r3b.font.size = Pt(9)

    # Line 4: specific text to add
    if added:
        p4 = cell.add_paragraph()
        r4a = p4.add_run('ADD: ')
        r4a.bold = True; r4a.font.color.rgb = GREEN; r4a.font.size = Pt(9)
        r4b = p4.add_run(f'"{added}"')
        r4b.italic = True; r4b.font.color.rgb = GREEN; r4b.font.size = Pt(9)

    if copy_paste_text.strip():
        p5 = cell.add_paragraph()
        r5a = p5.add_run('COPY-PASTE: ')
        r5a.bold = True; r5a.font.color.rgb = DARK; r5a.font.size = Pt(9)
        r5b = p5.add_run(copy_paste_text.strip())
        r5b.font.color.rgb = DARK; r5b.font.size = Pt(9)

    citation_suggestions = citation_suggestions or []
    if citation_suggestions:
        p6 = cell.add_paragraph()
        r6a = p6.add_run('CITATIONS TO CONSIDER: ')
        r6a.bold = True; r6a.font.color.rgb = GRAY; r6a.font.size = Pt(8)
        r6b = p6.add_run('; '.join(citation_suggestions[:4]))
        r6b.font.color.rgb = GRAY; r6b.font.size = Pt(8)


def build_point_by_point_docx(
    revision_round: dict,
    manuscript_title: str = "",
) -> bytes:
    """
    Landscape .docx — 4-column point-by-point reply table.

    Columns:
      Sr. No. | Reviewer's Comment | Author's Response (blue)
      | Manuscript Change Instructions (ADD/DELETE/CHANGE with location + text)

    Page: Letter landscape (11″ × 8.5″), 0.75″ margins → 9.5″ usable width.
    """
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = docx.Document()

    # ── Landscape page setup ──────────────────────────────────────────────
    sec = doc.sections[0]
    sec.orientation   = WD_ORIENT.LANDSCAPE
    sec.page_width    = Inches(11)
    sec.page_height   = Inches(8.5)
    sec.left_margin   = Inches(0.75)
    sec.right_margin  = Inches(0.75)
    sec.top_margin    = Inches(0.75)
    sec.bottom_margin = Inches(0.75)

    # ── Title block ───────────────────────────────────────────────────────
    doc.add_heading('Point-by-Point Reply and Changes Made in the Manuscript', level=1)
    if manuscript_title:
        p = doc.add_paragraph()
        run = p.add_run(manuscript_title)
        run.bold = True; run.font.size = Pt(12)

    journal_name = revision_round.get('journal_name', '')
    if journal_name:
        p = doc.add_paragraph()
        p.add_run('Journal: ').bold = True
        p.add_run(journal_name)

    intro = doc.add_paragraph()
    intro.add_run(
        'We sincerely thank all reviewers for their careful reading and constructive comments. '
        'Below we provide a detailed point-by-point response to each concern. '
        'The Change Plan column (purple) summarises the agreed strategy for each comment. '
        'Author responses are shown in blue. The final column — Manuscript Change '
        'Instructions — specifies exactly what to ADD, DELETE, or CHANGE in the '
        'manuscript, with section name, approximate line numbers, and the specific text. '
        'Please apply these changes directly to your manuscript.'
    ).font.size = Pt(10)

    # ── Column widths (twips) — total 9.5 in = 13 680 twips ──────────────
    #   Sr.No  Comment  ChangePlan  Response  Instructions
    COL_TWIPS = [504, 2520, 2880, 3024, 4752]
    HEADERS = [
        'Sr. No.',
        "Reviewer's Comment",
        "Change Plan",
        "Author's Response",
        "Manuscript Change Instructions\n(what · where · exact text)",
    ]
    BLUE      = RGBColor(0x1A, 0x56, 0xDB)
    PURPLE    = RGBColor(0x5B, 0x21, 0xB6)
    HDR_FILL  = '1E3A5F'  # dark navy

    responses = revision_round.get('responses', [])
    by_reviewer: dict[int, list[dict]] = {}
    for r in responses:
        by_reviewer.setdefault(r.get('reviewer_number', 1), []).append(r)

    for rev_num in sorted(by_reviewer.keys()):
        doc.add_heading(f'Reviewer #{rev_num}', level=2)
        comments = sorted(by_reviewer[rev_num], key=lambda x: x.get('comment_number', 0))

        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        table.autofit = False

        # Header row — dark navy background, white text
        hdr_cells = table.rows[0].cells
        for cell, text, twips in zip(hdr_cells, HEADERS, COL_TWIPS):
            _set_cell_width(cell, twips)
            run = cell.paragraphs[0].add_run(text)
            run.bold = True; run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), HDR_FILL)
            tcPr.append(shd)

        # Data rows
        for c in comments:
            row_cells = table.add_row().cells
            for cell, twips in zip(row_cells, COL_TWIPS):
                _set_cell_width(cell, twips)

            # Sr. No.
            r0 = row_cells[0].paragraphs[0].add_run(str(c.get('comment_number', '')))
            r0.bold = True; r0.font.size = Pt(9)

            # Reviewer's Comment
            r1 = row_cells[1].paragraphs[0].add_run(c.get('original_comment', ''))
            r1.font.size = Pt(9)

            # Change Plan (purple)
            change_plan_text = c.get('change_plan', '').strip()
            r3 = row_cells[2].paragraphs[0].add_run(change_plan_text or '—')
            r3.font.color.rgb = PURPLE; r3.font.size = Pt(9)

            # Author's Response (blue)
            r2 = row_cells[3].paragraphs[0].add_run(c.get('author_response', ''))
            r2.font.color.rgb = BLUE; r2.font.size = Pt(9)

            # Manuscript Change Instructions
            _fill_change_column(
                row_cells[4],
                c.get('action_taken', ''),
                c.get('manuscript_diff', '{}'),
                c.get('copy_paste_text', ''),
                c.get('citation_suggestions', []),
            )

        doc.add_paragraph('')  # spacer between reviewers

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── B. Clean revised manuscript ────────────────────────────────────────────────

def build_clean_revised_docx(revised_article: str, manuscript_title: str = "") -> bytes:
    """Convert the revised manuscript markdown to a clean .docx."""
    import docx

    doc = docx.Document()
    _markdown_to_doc(doc, revised_article)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── C. Track-changes .docx (real OOXML) ───────────────────────────────────────

def _make_ins(text: str, rev_id: int, author: str, date: str):
    from docx.oxml.ns import qn
    from lxml import etree
    ins = etree.Element(qn('w:ins'))
    ins.set(qn('w:id'), str(rev_id))
    ins.set(qn('w:author'), author)
    ins.set(qn('w:date'), date)
    r = etree.SubElement(ins, qn('w:r'))
    t = etree.SubElement(r, qn('w:t'))
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    return ins


def _make_del(text: str, rev_id: int, author: str, date: str):
    from docx.oxml.ns import qn
    from lxml import etree
    d = etree.Element(qn('w:del'))
    d.set(qn('w:id'), str(rev_id))
    d.set(qn('w:author'), author)
    d.set(qn('w:date'), date)
    r = etree.SubElement(d, qn('w:r'))
    t = etree.SubElement(r, qn('w:delText'))
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    return d


def _make_run(text: str):
    from docx.oxml.ns import qn
    from lxml import etree
    r = etree.Element(qn('w:r'))
    t = etree.SubElement(r, qn('w:t'))
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    return r


def _diff_paragraph(p_elem, orig_text: str, rev_text: str,
                    rev_id_start: int, author: str, date: str) -> int:
    """
    Token-level diff between orig_text and rev_text.

    Tokens are alternating word/whitespace units so original spacing is
    preserved exactly — no trailing spaces or double-spaces injected.

    Appends w:ins / w:del / w:r elements to p_elem.
    Returns the next available rev_id.
    """
    orig_tokens = _tokenize_text(orig_text) if orig_text else []
    rev_tokens  = _tokenize_text(rev_text)  if rev_text  else []

    if not orig_tokens and not rev_tokens:
        return rev_id_start

    matcher = difflib.SequenceMatcher(None, orig_tokens, rev_tokens, autojunk=False)
    rev_id = rev_id_start

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            p_elem.append(_make_run(''.join(orig_tokens[i1:i2])))

        elif tag == 'delete':
            chunk = ''.join(orig_tokens[i1:i2])
            if chunk.strip():
                # Only mark non-whitespace deletions as w:del
                p_elem.append(_make_del(chunk, rev_id, author, date))
                rev_id += 1
            else:
                p_elem.append(_make_run(chunk))

        elif tag == 'insert':
            chunk = ''.join(rev_tokens[j1:j2])
            if chunk.strip():
                p_elem.append(_make_ins(chunk, rev_id, author, date))
                rev_id += 1
            else:
                p_elem.append(_make_run(chunk))

        elif tag == 'replace':
            del_text = ''.join(orig_tokens[i1:i2])
            ins_text = ''.join(rev_tokens[j1:j2])
            p_elem.append(_make_del(del_text, rev_id, author, date))
            rev_id += 1
            p_elem.append(_make_ins(ins_text, rev_id, author, date))
            rev_id += 1

    return rev_id


def build_track_changes_docx(
    original_manuscript: str,
    revised_manuscript: str,
    author: str = "Author",
) -> bytes:
    """
    Build a .docx with real OOXML track changes (w:ins / w:del).

    Improvements over the naive approach:
    - Preserves heading styles (Heading 1–4) from markdown structure
    - Token-level diff (word + whitespace tokens) for accurate spacing
    - Two-level diff: paragraph-level then token-level within changed paras
    - Sub-diffing inside replace blocks to maximise inline annotations
    - ISO 8601 date without microseconds (OOXML-compliant)
    """
    import docx

    # OOXML-compliant date: no microseconds
    date = datetime.now(timezone.utc).strftime('%Y-%m-%dT%H:%M:%SZ')

    orig_paras = _parse_manuscript(original_manuscript)
    rev_paras  = _parse_manuscript(revised_manuscript)

    doc = docx.Document()
    # Remove the default empty paragraph python-docx adds
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)

    rev_id = 1

    # ── helpers ──────────────────────────────────────────────────────────────

    def _add_para_elem(style: str):
        """
        Add a styled empty paragraph and return its lxml w:p element.
        Passing '' text avoids auto-adding an empty run.
        """
        if style.startswith('Heading'):
            level = int(style.split()[-1])
            return doc.add_heading('', level=level)._element
        return doc.add_paragraph()._element

    def _render_block(pairs: list[tuple[str, str]], wrap_fn):
        """Render a list of (style, text) paragraphs, wrapping text in wrap_fn."""
        nonlocal rev_id
        for style, text in pairs:
            p_elem = _add_para_elem(style)
            if text:
                p_elem.append(wrap_fn(text, rev_id, author, date))
                rev_id += 1

    def _process_replace(orig_block: list[tuple[str, str]],
                         rev_block:  list[tuple[str, str]]):
        """
        Diff a replace block as granularly as possible:
        1. Sub-diff the paragraph list so identical paragraphs inside the
           block are rendered as plain (equal) text, not del+ins.
        2. For 1:1 para replacements do token-level inline diff.
        3. For N:M mismatches fall back to del-all / ins-all.
        """
        nonlocal rev_id

        sub = difflib.SequenceMatcher(None, orig_block, rev_block, autojunk=False)
        for s_tag, si1, si2, sj1, sj2 in sub.get_opcodes():
            o_sub = orig_block[si1:si2]
            r_sub = rev_block[sj1:sj2]

            if s_tag == 'equal':
                for style, text in o_sub:
                    p_elem = _add_para_elem(style)
                    if text:
                        p_elem.append(_make_run(text))

            elif s_tag == 'delete':
                _render_block(o_sub, _make_del)

            elif s_tag == 'insert':
                _render_block(r_sub, _make_ins)

            elif s_tag == 'replace':
                if len(o_sub) == len(r_sub):
                    # 1:1 → token-level inline diff per paragraph
                    for (o_style, o_text), (r_style, r_text) in zip(o_sub, r_sub):
                        p_elem = _add_para_elem(r_style)
                        rev_id = _diff_paragraph(
                            p_elem, o_text, r_text, rev_id, author, date
                        )
                else:
                    # N:M mismatch → delete originals, insert revised
                    _render_block(o_sub, _make_del)
                    _render_block(r_sub, _make_ins)

    # ── paragraph-level diff ─────────────────────────────────────────────────

    para_matcher = difflib.SequenceMatcher(None, orig_paras, rev_paras, autojunk=False)

    for tag, i1, i2, j1, j2 in para_matcher.get_opcodes():

        if tag == 'equal':
            for style, text in orig_paras[i1:i2]:
                if not style:          # blank line
                    doc.add_paragraph('')
                    continue
                p_elem = _add_para_elem(style)
                if text:
                    p_elem.append(_make_run(text))

        elif tag == 'delete':
            for style, text in orig_paras[i1:i2]:
                if not style:
                    continue           # skip blank-line deletions
                p_elem = _add_para_elem(style)
                if text:
                    p_elem.append(_make_del(text, rev_id, author, date))
                    rev_id += 1

        elif tag == 'insert':
            for style, text in rev_paras[j1:j2]:
                if not style:
                    continue           # skip blank-line insertions
                p_elem = _add_para_elem(style)
                if text:
                    p_elem.append(_make_ins(text, rev_id, author, date))
                    rev_id += 1

        elif tag == 'replace':
            # Strip blank lines before sub-diffing
            orig_block = [(s, t) for s, t in orig_paras[i1:i2] if s]
            rev_block  = [(s, t) for s, t in rev_paras[j1:j2]  if s]
            _process_replace(orig_block, rev_block)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
