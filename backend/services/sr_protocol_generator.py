"""
sr_protocol_generator.py

Generates PRISMA-P 2015 compliant SR protocol documents, maps to PROSPERO fields,
generates search strings for 8 databases, and optionally registers on OSF.
"""

from __future__ import annotations

import asyncio
import json
import logging
import re
import uuid
import io
from datetime import datetime
from typing import Optional

import httpx

from services.ai_provider import AIProvider
from services.completion_guard import CompletionConfig, OutputFormat, TokenBudget

logger = logging.getLogger(__name__)


_PRISMA_P_ITEMS = {
    "1": "Title — identifies the report as a systematic review/meta-analysis",
    "2": "Registration — PROSPERO number with expected completion date",
    "3": "Protocol and registration — DOI or link to registered protocol",
    "4a": "Funding — funding sources for the review",
    "4b": "Institutional affiliations of all authors",
    "5": "Contributions of each author",
    "6": "Background — rationale for the review in context of existing knowledge",
    "7": "Objectives — explicit statement using PICO elements",
    "8": "Eligibility criteria — characteristics of eligible studies (PICO + study design + time/language limits)",
    "9": "Information sources — databases, registers, grey literature",
    "10": "Search strategy — full search for at least one database, including any limits",
    "11": "Study records — data management, selection process (numbers screened/included), automation",
    "12a": "Data collection — data extraction process (forms, pilot testing, dual extraction)",
    "12b": "Data items — list of variables to be collected",
    "13a": "Outcomes and prioritization — primary/secondary outcomes with measurement, aggregation",
    "13b": "Methods for handling alternative or multiple groupings of interventions/comparators",
    "14a": "Risk of bias assessment — tool to be used, per domain",
    "14b": "Individual study level vs summary level risk assessment",
    "15a": "Data synthesis — criteria for meta-analysis vs narrative synthesis",
    "15b": "Effect measure(s) to be used",
    "15c": "Heterogeneity — assessment (I², Q-statistic) and investigation (subgroup, meta-regression)",
    "15d": "Sensitivity analyses planned",
    "15e": "Assessment of reporting biases (funnel plots, Egger's test)",
    "15f": "Assessment of certainty (GRADE framework)",
    "16": "Meta-bias — potential biases that may affect the review process",
    "17": "Funding — sources, support, and role of funders",
}

_PRISMA_P_CHECKLIST_TEMPLATE_ROWS = [
    {"kind": "section", "label": "ADMINISTRATIVE INFORMATION"},
    {"kind": "group", "label": "Title:"},
    {"kind": "item", "topic": "Identification", "item_no": "1a",
     "checklist_item": "Identify the report as a protocol of a systematic review"},
    {"kind": "item", "topic": "Update", "item_no": "1b",
     "checklist_item": "If the protocol is for an update of a previous systematic review, identify as such"},
    {"kind": "item", "topic": "Registration", "item_no": "2",
     "checklist_item": "If registered, provide the name of the registry (such as PROSPERO) and registration number"},
    {"kind": "group", "label": "Authors:"},
    {"kind": "item", "topic": "Contact", "item_no": "3a",
     "checklist_item": "Provide name, institutional affiliation, e-mail address of all protocol authors; provide physical mailing address of corresponding author"},
    {"kind": "item", "topic": "Contributions", "item_no": "3b",
     "checklist_item": "Describe contributions of protocol authors and identify the guarantor of the review"},
    {"kind": "item", "topic": "Amendments", "item_no": "4",
     "checklist_item": "If the protocol represents an amendment of a previously completed or published protocol, identify as such and list changes; otherwise, state plan for documenting important protocol amendments"},
    {"kind": "group", "label": "Support:"},
    {"kind": "item", "topic": "Sources", "item_no": "5a",
     "checklist_item": "Indicate sources of financial or other support for the review"},
    {"kind": "item", "topic": "Sponsor", "item_no": "5b",
     "checklist_item": "Provide name for the review funder and/or sponsor"},
    {"kind": "item", "topic": "Role of sponsor or funder", "item_no": "5c",
     "checklist_item": "Describe roles of funder(s), sponsor(s), and/or institution(s), if any, in developing the protocol"},
    {"kind": "section", "label": "INTRODUCTION"},
    {"kind": "item", "topic": "Rationale", "item_no": "6",
     "checklist_item": "Describe the rationale for the review in the context of what is already known"},
    {"kind": "item", "topic": "Objectives", "item_no": "7",
     "checklist_item": "Provide an explicit statement of the question(s) the review will address with reference to participants, interventions, comparators, and outcomes (PICO)"},
    {"kind": "section", "label": "METHODS"},
    {"kind": "item", "topic": "Eligibility criteria", "item_no": "8",
     "checklist_item": "Specify the study characteristics (such as PICO, study design, setting, time frame) and report characteristics (such as years considered, language, publication status) to be used as criteria for eligibility for the review"},
    {"kind": "item", "topic": "Information sources", "item_no": "9",
     "checklist_item": "Describe all intended information sources (such as electronic databases, contact with study authors, trial registers or other grey literature sources) with planned dates of coverage"},
    {"kind": "item", "topic": "Search strategy", "item_no": "10",
     "checklist_item": "Present draft of search strategy to be used for at least one electronic database, including planned limits, such that it could be repeated"},
    {"kind": "group", "label": "Study records:"},
    {"kind": "item", "topic": "Data management", "item_no": "11a",
     "checklist_item": "Describe the mechanism(s) that will be used to manage records and data throughout the review"},
    {"kind": "item", "topic": "Selection process", "item_no": "11b",
     "checklist_item": "State the process that will be used for selecting studies (such as two independent reviewers) through each phase of the review (that is, screening, eligibility and inclusion in meta-analysis)"},
    {"kind": "item", "topic": "Data collection process", "item_no": "11c",
     "checklist_item": "Describe planned method of extracting data from reports (such as piloting forms, done independently, in duplicate), any processes for obtaining and confirming data from investigators"},
    {"kind": "item", "topic": "Data items", "item_no": "12",
     "checklist_item": "List and define all variables for which data will be sought (such as PICO items, funding sources), any pre-planned data assumptions and simplifications"},
    {"kind": "item", "topic": "Outcomes and prioritization", "item_no": "13",
     "checklist_item": "List and define all outcomes for which data will be sought, including prioritization of main and additional outcomes, with rationale"},
    {"kind": "item", "topic": "Risk of bias in individual studies", "item_no": "14",
     "checklist_item": "Describe anticipated methods for assessing risk of bias of individual studies, including whether this will be done at the outcome or study level, or both; state how this information will be used in data synthesis"},
    {"kind": "item", "topic": "Data synthesis", "item_no": "15a",
     "checklist_item": "Describe criteria under which study data will be quantitatively synthesised"},
    {"kind": "item", "topic": "Data synthesis", "item_no": "15b",
     "checklist_item": "If data are appropriate for quantitative synthesis, describe planned summary measures, methods of handling data and methods of combining data from studies, including any planned exploration of consistency (such as I2, Kendall’s τ)"},
    {"kind": "item", "topic": "Data synthesis", "item_no": "15c",
     "checklist_item": "Describe any proposed additional analyses (such as sensitivity or subgroup analyses, meta-regression)"},
    {"kind": "item", "topic": "Data synthesis", "item_no": "15d",
     "checklist_item": "If quantitative synthesis is not appropriate, describe the type of summary planned"},
    {"kind": "item", "topic": "Meta-bias(es)", "item_no": "16",
     "checklist_item": "Specify any planned assessment of meta-bias(es) (such as publication bias across studies, selective reporting within studies)"},
    {"kind": "item", "topic": "Confidence in cumulative evidence", "item_no": "17",
     "checklist_item": "Describe how the strength of the body of evidence will be assessed (such as GRADE)"},
]

_PRISMA_P_PAGE_ALIASES = {
    "1a": ["administrative information", "title"],
    "1b": ["administrative information", "title"],
    "2": ["administrative information", "registration"],
    "3a": ["administrative information", "authors"],
    "3b": ["administrative information", "authors"],
    "4": ["administrative information", "amendments"],
    "5a": ["administrative information", "support", "funding"],
    "5b": ["administrative information", "support", "sponsor"],
    "5c": ["administrative information", "support", "sponsor"],
    "6": ["background and rationale", "introduction", "background"],
    "7": ["objectives"],
    "8": ["eligibility criteria"],
    "9": ["information sources"],
    "10": ["search strategy"],
    "11a": ["study records", "data management", "records management"],
    "11b": ["selection process", "screening", "study records"],
    "11c": ["data collection process", "study records"],
    "12": ["data items"],
    "13": ["outcomes and prioritization", "outcomes"],
    "14": ["risk of bias assessment", "risk of bias"],
    "15a": ["data synthesis", "synthesis methods", "synthesis plan"],
    "15b": ["data synthesis", "effect measures", "synthesis methods"],
    "15c": ["data synthesis", "subgroup and sensitivity analyses", "subgroup sensitivity"],
    "15d": ["data synthesis", "narrative synthesis", "synthesis methods"],
    "16": ["meta bias assessment", "reporting bias assessment", "meta bias", "reporting bias"],
    "17": ["confidence in evidence", "certainty of evidence", "grade", "confidence"],
}


def _join_readable_list(items: list[str]) -> str:
    cleaned = [str(item).strip() for item in items if str(item).strip()]
    if not cleaned:
        return ""
    if len(cleaned) == 1:
        return cleaned[0]
    if len(cleaned) == 2:
        return f"{cleaned[0]} and {cleaned[1]}"
    return f"{', '.join(cleaned[:-1])}, and {cleaned[-1]}"


def _ensure_sentence(text: str) -> str:
    trimmed = (text or "").strip()
    if not trimmed:
        return ""
    return trimmed if trimmed[-1] in ".!?" else f"{trimmed}."


def _coerce_string_list(value: object) -> list[str]:
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()]
    if isinstance(value, str) and value.strip():
        return [value.strip()]
    return []


def _strip_control_chars(text: str) -> str:
    return "".join(ch for ch in text if ch in "\n\r\t" or ord(ch) >= 32)


def _json_safe(value):
    if isinstance(value, str):
        return _strip_control_chars(value)
    if isinstance(value, list):
        return [_json_safe(item) for item in value]
    if isinstance(value, dict):
        return {key: _json_safe(item) for key, item in value.items()}
    return value


def _load_llm_json(raw: str) -> dict:
    text = (raw or "").strip()
    if not text:
        return {}

    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*", "", text)
        text = re.sub(r"\s*```$", "", text).strip()

    candidates = [text]
    start = text.find("{")
    end = text.rfind("}")
    if start != -1 and end > start:
        candidates.append(text[start:end + 1])

    seen: set[str] = set()
    last_error: Exception | None = None
    for candidate in candidates:
        if candidate in seen:
            continue
        seen.add(candidate)
        try:
            return json.loads(candidate)
        except json.JSONDecodeError as exc:
            last_error = exc
            try:
                return json.loads(candidate, strict=False)
            except json.JSONDecodeError as inner_exc:
                last_error = inner_exc

    if last_error:
        raise last_error
    return {}


def _parse_tagged_text_phase_response(raw: str) -> tuple[str, str]:
    text = (raw or "").strip()
    if not text:
        return "", ""

    if text.startswith("{"):
        try:
            payload = _load_llm_json(text)
            if isinstance(payload, dict):
                reply = str(payload.get("chat_reply") or "").strip()
                body = str(payload.get("text") or "").strip()
                if reply or body:
                    return reply, body
        except Exception:
            pass

    if text.startswith("```"):
        text = re.sub(r"^```(?:markdown|md|text)?\s*", "", text)
        text = re.sub(r"\s*```$", "", text).strip()

    reply = ""
    body = ""

    tagged_match = re.search(
        r"CHAT_REPLY:\s*(.*?)\s*TEXT:\s*<<<\s*(.*?)\s*>>>\s*$",
        text,
        flags=re.DOTALL,
    )
    if tagged_match:
        reply = tagged_match.group(1).strip()
        body = tagged_match.group(2).strip()
        return reply, body

    text_idx = text.find("TEXT:")
    if text_idx != -1:
        before = text[:text_idx].strip()
        after = text[text_idx + len("TEXT:"):].strip()
        if before.startswith("CHAT_REPLY:"):
            reply = before[len("CHAT_REPLY:"):].strip()
        body = after
        if body.startswith("<<<"):
            body = body[3:].lstrip()
        if body.endswith(">>>"):
            body = body[:-3].rstrip()
        return reply, body.strip()

    if text.startswith("CHAT_REPLY:"):
        reply = text[len("CHAT_REPLY:"):].strip()
        return reply, ""

    return "", text


def _build_data_collection_narrative(content: dict | None) -> str:
    if not isinstance(content, dict):
        return ""

    parts: list[str] = []
    extraction_method = str(content.get("extraction_method") or content.get("data_collection_notes") or "").strip()
    extraction_team = str(content.get("extraction_team") or "").strip()
    pilot_testing = str(content.get("pilot_testing") or "").strip()
    disagreement_resolution = str(content.get("disagreement_resolution") or "").strip()
    software = str(content.get("software") or "").strip()
    author_contact = str(content.get("author_contact") or "").strip()

    if extraction_method:
        parts.append(_ensure_sentence(extraction_method))
    if extraction_team:
        parts.append(_ensure_sentence(f"Data extraction will be undertaken by {extraction_team}"))
    if pilot_testing:
        parts.append(_ensure_sentence(f"The extraction approach will be piloted {pilot_testing}"))
    if disagreement_resolution:
        parts.append(_ensure_sentence(f"Disagreements will be resolved through {disagreement_resolution}"))
    if software:
        parts.append(_ensure_sentence(f"Extraction records will be managed using {software}"))
    if author_contact:
        parts.append(_ensure_sentence(f"Where needed, study authors will be contacted {author_contact}"))

    return " ".join(part for part in parts if part)


def _build_data_items_markdown(content: dict | None, section_prefix: str = "") -> str:
    if not isinstance(content, dict):
        return ""

    sections = [
        (
            "1",
            "Study Characteristics",
            "The following study-level descriptors will be recorded for each included study: ",
            _coerce_string_list(content.get("study_characteristics")),
        ),
        (
            "2",
            "Participant Characteristics",
            "Participant-level data to be extracted will include ",
            _coerce_string_list(content.get("participant_characteristics")),
        ),
        (
            "3",
            "Intervention and Comparator Details",
            "Information pertaining to the exposure, intervention, or comparator will include ",
            _coerce_string_list(content.get("intervention_characteristics")),
        ),
        (
            "4",
            "Outcome Data",
            "Outcome data to be extracted will include ",
            _coerce_string_list(content.get("outcome_items")),
        ),
        (
            "5",
            "Methodological Variables",
            "Methodological data to be recorded will include ",
            _coerce_string_list(content.get("methodological_items")),
        ),
    ]

    blocks: list[str] = []
    for suffix, title, intro, items in sections:
        if not items:
            continue
        heading = f"#### {section_prefix}.{suffix} {title}" if section_prefix else f"#### {title}"
        blocks.extend([
            heading,
            "",
            _ensure_sentence(f"{intro}{_join_readable_list(items)}"),
            "",
        ])

    return "\n".join(blocks).strip()

# ── Cochrane/JBI study-design-specific extraction schema templates ─────────────
# Sources: Cochrane Handbook Ch.11, JBI Manual, SRDR+, PRISMA 2020, RoB 2.0, QUADAS-2
# Field types: text | number | list | boolean | dichotomous_outcome | continuous_outcome
# dichotomous_outcome → {events_intervention, n_intervention, events_control, n_control}
# continuous_outcome  → {mean_intervention, sd_intervention, n_intervention, mean_control, sd_control, n_control}

def _f(field: str, ftype: str, required: bool, description: str, section: str = "General") -> dict:
    return {"field": field, "type": ftype, "required": required,
            "description": description, "section": section}


_SCHEMA_TEMPLATES: dict[str, list[dict]] = {

    # ── RCT (Cochrane Handbook Ch.11, RoB 2.0) ──────────────────────────────
    "rct": [
        # Methods
        _f("study_design",           "text",    True,  "Confirm parallel/crossover/cluster/factorial RCT", "Methods"),
        _f("trial_registration_id",  "text",    False, "ClinicalTrials.gov or other registry ID (e.g., NCT01234567)", "Methods"),
        _f("funding_source",         "text",    True,  "Industry, government, charitable, or none declared", "Methods"),
        _f("conflicts_of_interest",  "text",    True,  "Author-reported conflicts; note 'none declared' if stated", "Methods"),
        _f("randomization_method",   "text",    True,  "Computer, block, stratified, minimisation — extract verbatim (Cochrane RoB 2 D1)", "Methods"),
        _f("allocation_concealment", "text",    True,  "Central phone/web, sealed opaque envelopes, pharmacy — extract verbatim (Cochrane RoB 2 D1)", "Methods"),
        _f("blinding_participants",  "boolean", True,  "Were participants blinded to allocation? (Cochrane RoB 2 D2)", "Methods"),
        _f("blinding_assessors",     "boolean", True,  "Were outcome assessors blinded? (Cochrane RoB 2 D4)", "Methods"),
        _f("itt_analysis",           "text",    True,  "Intention-to-treat (full ITT / modified ITT / per-protocol) — extract verbatim", "Methods"),
        _f("country",                "text",    False, "Country or countries where study was conducted", "Methods"),
        _f("setting",                "text",    False, "Hospital, community, primary care, online, etc.", "Methods"),
        # Participants
        _f("sample_size_randomized", "number",  True,  "Total N randomized across all arms", "Participants"),
        _f("sample_size_analyzed",   "number",  True,  "Total N included in primary analysis (often < randomized)", "Participants"),
        _f("age_mean_sd",            "text",    True,  "Mean age ± SD (or median [IQR]) for each arm", "Participants"),
        _f("percent_female",         "number",  False, "% female participants across all arms", "Participants"),
        _f("inclusion_criteria_summary", "text", False, "Key inclusion criteria as stated by authors", "Participants"),
        _f("baseline_comparability", "text",    False, "Were groups comparable at baseline? Note any imbalances", "Participants"),
        _f("n_lost_followup_intervention", "number", False, "N lost to follow-up in intervention arm", "Participants"),
        _f("n_lost_followup_control",      "number", False, "N lost to follow-up in control arm", "Participants"),
        _f("reasons_attrition",      "text",    False, "Reasons for dropout/withdrawal — extract verbatim if given", "Participants"),
        # Interventions
        _f("intervention_description","text",   True,  "Name, dose, frequency, delivery mode, duration — sufficient to replicate", "Interventions"),
        _f("comparator_description",  "text",   True,  "Placebo, usual care, active comparator — describe fully", "Interventions"),
        _f("intervention_duration",   "text",   True,  "Total duration of active intervention (e.g., '12 weeks')", "Interventions"),
        _f("intervention_fidelity",   "text",   False, "How adherence/fidelity was monitored or reported", "Interventions"),
        # Outcomes
        _f("primary_outcomes",       "list",    True,  "List primary outcomes with measurement tool and timepoint (e.g., 'HbA1c at 12 weeks via lab test')", "Outcomes"),
        _f("secondary_outcomes",     "list",    False, "List secondary outcomes with measurement tools", "Outcomes"),
        _f("adverse_events",         "list",    False, "All adverse events reported; note if none or 'not reported'", "Outcomes"),
        _f("outcome_timepoints",     "text",    False, "All measurement timepoints (e.g., baseline, 4w, 8w, 12w)", "Outcomes"),
        # Results
        _f("primary_outcome_data",   "dichotomous_outcome", True,
           "For binary outcomes: events and N per arm. For continuous: use continuous_outcome field instead", "Results"),
        _f("effect_measure_type",    "text",    True,  "OR, RR, RD, MD, SMD, HR — as reported by authors", "Results"),
        _f("effect_estimate",        "number",  True,  "Point estimate of effect (e.g., OR=1.23, MD=-0.5)", "Results"),
        _f("ci_95",                  "text",    True,  "95% confidence interval (e.g., '0.95 to 1.59')", "Results"),
        _f("p_value",                "text",    False, "P-value for primary outcome (exact if available)", "Results"),
        _f("author_conclusion",      "text",    True,  "Authors' main conclusion — extract verbatim from abstract/conclusion", "Results"),
    ],

    # ── Observational (Cohort / Case-Control / Cross-Sectional) ─────────────
    # Sources: Cochrane Handbook Ch.25, ROBINS-I, Newcastle-Ottawa Scale
    "observational": [
        _f("study_design",           "text",    True,  "Prospective cohort / retrospective cohort / case-control / cross-sectional", "Methods"),
        _f("study_duration",         "text",    False, "Total follow-up period for cohort studies", "Methods"),
        _f("funding_source",         "text",    True,  "Industry, government, charitable, or none declared", "Methods"),
        _f("conflicts_of_interest",  "text",    True,  "Author-reported conflicts", "Methods"),
        _f("country",                "text",    False, "Country or countries", "Methods"),
        _f("setting",                "text",    False, "Hospital, community, registry, database", "Methods"),
        # Participants
        _f("sample_size_total",      "number",  True,  "Total N enrolled/analysed", "Participants"),
        _f("age_mean_sd",            "text",    True,  "Mean/median age with SD or IQR", "Participants"),
        _f("percent_female",         "number",  False, "% female", "Participants"),
        _f("inclusion_criteria_summary", "text", False, "Key inclusion criteria", "Participants"),
        _f("response_rate",          "text",    False, "Response or participation rate (for surveys/cross-sectional)", "Participants"),
        # Exposure & confounding (ROBINS-I aligned)
        _f("intervention_description","text",   True,  "Exposure or intervention as defined by authors", "Interventions"),
        _f("comparator_description",  "text",   False, "Unexposed group or comparison group definition", "Interventions"),
        _f("confounders_measured",    "list",   True,  "List all confounders measured (e.g., age, BMI, smoking)", "Interventions"),
        _f("confounders_adjusted",    "list",   True,  "List confounders adjusted for in analysis", "Interventions"),
        _f("method_of_adjustment",    "text",   True,  "Regression, propensity score, stratification, matching", "Interventions"),
        # Outcomes
        _f("primary_outcomes",       "list",    True,  "Primary outcomes with measurement definition", "Outcomes"),
        _f("secondary_outcomes",     "list",    False, "Secondary outcomes", "Outcomes"),
        _f("outcome_assessment",     "text",    True,  "Objective (lab/imaging/registry) or subjective (self-report)?", "Outcomes"),
        _f("blinding_assessors",     "boolean", False, "Were outcome assessors blinded to exposure status?", "Outcomes"),
        # Results
        _f("effect_measure_type",    "text",    True,  "OR, RR, HR, RD, IRR, prevalence ratio — as reported", "Results"),
        _f("effect_estimate",        "number",  True,  "Adjusted point estimate", "Results"),
        _f("ci_95",                  "text",    True,  "95% CI for primary estimate", "Results"),
        _f("p_value",                "text",    False, "P-value", "Results"),
        _f("unadjusted_estimate",    "text",    False, "Unadjusted (crude) estimate if reported alongside adjusted", "Results"),
        _f("subgroup_analyses",      "text",    False, "Any pre-specified subgroup analyses", "Results"),
        _f("author_conclusion",      "text",    True,  "Authors' main conclusion — extract verbatim", "Results"),
    ],

    # ── Diagnostic Test Accuracy (QUADAS-2) ───────────────────────────────────
    "diagnostic": [
        _f("study_design",           "text",    True,  "Prospective/retrospective diagnostic cohort; single-gate or two-gate", "Methods"),
        _f("funding_source",         "text",    True,  "Funding source", "Methods"),
        _f("country",                "text",    False, "Country", "Methods"),
        _f("setting",                "text",    True,  "Primary care, emergency, specialist, community screening", "Methods"),
        # Participants
        _f("sample_size_total",      "number",  True,  "Total N in the diagnostic study", "Participants"),
        _f("age_mean_sd",            "text",    True,  "Mean/median age", "Participants"),
        _f("percent_female",         "number",  False, "% female", "Participants"),
        _f("disease_prevalence",     "number",  True,  "Prevalence of target condition in study sample (%)", "Participants"),
        _f("participant_selection",  "text",    True,  "Consecutive, random, or convenience sampling? QUADAS-2 domain", "Participants"),
        # Index test & reference standard
        _f("index_test_name",        "text",    True,  "Name of index test being evaluated (e.g., 'PCR', 'troponin I')", "Interventions"),
        _f("index_test_threshold",   "text",    True,  "Cut-off/threshold used and how determined (pre-specified?)", "Interventions"),
        _f("reference_standard",     "text",    True,  "Gold standard used for diagnosis — name and how applied", "Interventions"),
        _f("blinding_index_test",    "boolean", True,  "Were index test interpreters blinded to reference standard?", "Interventions"),
        _f("blinding_reference",     "boolean", True,  "Were reference standard assessors blinded to index test?", "Interventions"),
        # Diagnostic accuracy results
        _f("true_positives",         "number",  True,  "True positives (TP) from 2×2 table", "Results"),
        _f("false_positives",        "number",  True,  "False positives (FP) from 2×2 table", "Results"),
        _f("false_negatives",        "number",  True,  "False negatives (FN) from 2×2 table", "Results"),
        _f("true_negatives",         "number",  True,  "True negatives (TN) from 2×2 table", "Results"),
        _f("sensitivity",            "number",  True,  "Sensitivity = TP/(TP+FN) × 100 (%)", "Results"),
        _f("specificity",            "number",  True,  "Specificity = TN/(TN+FP) × 100 (%)", "Results"),
        _f("auc_roc",                "number",  False, "Area Under ROC Curve (0–1)", "Results"),
        _f("ppv",                    "number",  False, "Positive Predictive Value (%)", "Results"),
        _f("npv",                    "number",  False, "Negative Predictive Value (%)", "Results"),
        _f("likelihood_ratio_pos",   "text",    False, "Positive likelihood ratio (+LR)", "Results"),
        _f("likelihood_ratio_neg",   "text",    False, "Negative likelihood ratio (−LR)", "Results"),
        _f("author_conclusion",      "text",    True,  "Authors' main conclusion", "Results"),
    ],

    # ── Qualitative (JBI Meta-Aggregative Approach, SPIDER) ──────────────────
    "qualitative": [
        _f("study_design",           "text",    True,  "Phenomenology, grounded theory, ethnography, narrative, case study, thematic analysis", "Methods"),
        _f("data_collection_method", "text",    True,  "Interviews, focus groups, observation, document review, diary — describe fully", "Methods"),
        _f("analysis_method",        "text",    True,  "Thematic analysis, content analysis, IPA, constant comparative, grounded theory", "Methods"),
        _f("theoretical_framework",  "text",    False, "Underpinning theory or epistemological stance stated by authors", "Methods"),
        _f("funding_source",         "text",    False, "Funding source", "Methods"),
        _f("country",                "text",    False, "Country of study", "Methods"),
        # Participants (Sample in SPIDER)
        _f("sample_size_total",      "number",  True,  "Number of participants", "Participants"),
        _f("participant_description","text",    True,  "Key characteristics: age, gender, role, experience, clinical status", "Participants"),
        _f("recruitment_method",     "text",    True,  "Purposive, snowball, theoretical, convenience", "Participants"),
        _f("setting",                "text",    True,  "Where participants were recruited/studied", "Participants"),
        _f("saturation_reached",     "boolean", False, "Did authors claim data saturation? (important for rigor)", "Participants"),
        # Phenomenon of Interest / Findings
        _f("key_findings",           "list",    True,  "List all JBI 'findings' (author interpretations with supporting data)", "Outcomes"),
        _f("themes",                 "list",    True,  "List of themes/categories identified by authors", "Outcomes"),
        _f("verbatim_quotes",        "list",    True,  "2–3 illustrative verbatim participant quotes per theme", "Outcomes"),
        _f("negative_cases",         "text",    False, "Did authors discuss disconfirming or negative cases?", "Outcomes"),
        # Methodological rigor (JBI QARI tool)
        _f("congruity_methodology",  "boolean", True,  "Is methodology congruent with research question? (JBI criterion 1)", "Methods"),
        _f("reflexivity_addressed",  "boolean", True,  "Did authors address researcher reflexivity? (JBI criterion 6)", "Methods"),
        _f("member_checking",        "boolean", False, "Was member-checking or participant validation conducted?", "Methods"),
        _f("author_conclusion",      "text",    True,  "Authors' main conclusion — extract verbatim", "Results"),
    ],

    # ── Prevalence / Incidence (CoCoPop, JBI Prevalence tool) ────────────────
    "prevalence": [
        _f("study_design",           "text",    True,  "Point prevalence / period prevalence / lifetime prevalence / incidence study", "Methods"),
        _f("sampling_method",        "text",    True,  "Random, systematic, stratified, cluster, convenience, census — Cochrane/JBI require representative sample", "Methods"),
        _f("funding_source",         "text",    False, "Funding source", "Methods"),
        _f("country",                "text",    True,  "Country or region", "Methods"),
        _f("setting",                "text",    True,  "Community, hospital, school, occupational, national registry", "Methods"),
        _f("study_year",             "text",    False, "Year or period of data collection", "Methods"),
        # Participants
        _f("sample_size_total",      "number",  True,  "Total N sampled", "Participants"),
        _f("response_rate",          "number",  False, "Participation/response rate (%)", "Participants"),
        _f("age_range",              "text",    True,  "Age range or mean of study population", "Participants"),
        _f("percent_female",         "number",  False, "% female", "Participants"),
        _f("inclusion_criteria_summary", "text", False, "Key inclusion/exclusion criteria", "Participants"),
        # Case definition & measurement
        _f("case_definition",        "text",    True,  "Exact diagnostic criteria or case definition used — extract verbatim", "Outcomes"),
        _f("outcome_measurement",    "text",    True,  "Objective (lab/registry) or subjective (self-report/screening tool)?", "Outcomes"),
        _f("measurement_tool",       "text",    False, "Specific scale, questionnaire, or test used", "Outcomes"),
        # Results
        _f("numerator",              "number",  True,  "Number with the condition (cases)", "Results"),
        _f("denominator",            "number",  True,  "Total number in denominator population", "Results"),
        _f("prevalence_estimate",    "number",  True,  "Prevalence/incidence proportion or rate (raw value)", "Results"),
        _f("prevalence_unit",        "text",    True,  "% / per 1000 / per 100,000 / person-years", "Results"),
        _f("ci_95",                  "text",    True,  "95% confidence interval for prevalence estimate", "Results"),
        _f("stratified_estimates",   "list",    False, "Estimates stratified by age, sex, region if reported", "Results"),
        _f("author_conclusion",      "text",    True,  "Authors' main conclusion", "Results"),
    ],

    # ── Case Report / Case Series (JBI) ──────────────────────────────────────
    "case_report": [
        _f("study_design",           "text",    True,  "Single case report / case series (specify N)", "Methods"),
        _f("funding_source",         "text",    False, "Funding source or none declared", "Methods"),
        _f("country",                "text",    False, "Country where case occurred", "Methods"),
        _f("setting",                "text",    False, "Hospital, outpatient, community", "Methods"),
        # Case characteristics
        _f("patient_age",            "text",    True,  "Patient age (or age range for series)", "Participants"),
        _f("patient_sex",            "text",    True,  "Patient sex/gender", "Participants"),
        _f("relevant_history",       "text",    False, "Relevant past medical history and comorbidities", "Participants"),
        # Clinical presentation
        _f("presenting_complaint",   "text",    True,  "Chief complaint or reason for presentation — extract verbatim", "Outcomes"),
        _f("symptom_duration",       "text",    False, "Duration of symptoms before presentation", "Outcomes"),
        _f("diagnostic_criteria",    "text",    True,  "Criteria used to establish diagnosis; tests performed and results", "Outcomes"),
        # Intervention
        _f("intervention_description","text",   True,  "Treatment provided — drug/dose/route or procedure; describe fully", "Interventions"),
        _f("intervention_duration",  "text",    False, "Duration of treatment", "Interventions"),
        # Outcomes
        _f("clinical_response",      "text",    True,  "Patient response to treatment — extract verbatim", "Results"),
        _f("adverse_events",         "list",    True,  "Any adverse events or complications; note 'none reported' if stated", "Results"),
        _f("followup_duration",      "text",    True,  "Total follow-up period after treatment", "Results"),
        _f("final_outcome",          "text",    True,  "Resolution / improvement / no change / deterioration / death", "Results"),
        _f("case_uniqueness",        "text",    True,  "What makes this case clinically significant or unique?", "Results"),
        _f("alternative_diagnoses",  "text",    False, "Alternative diagnoses considered and ruled out", "Results"),
        _f("author_conclusion",      "text",    True,  "Authors' main learning points and conclusion", "Results"),
    ],
}


def _select_template(question_type: str, study_design: str) -> tuple[str, list[dict]]:
    """
    Select the best Cochrane/JBI extraction schema template based on
    question type and study design. Returns (template_name, fields).
    """
    qt = (question_type or "").lower()
    sd = (study_design or "").lower()

    if qt == "diagnosis":
        return "diagnostic", _SCHEMA_TEMPLATES["diagnostic"]
    if qt in ("qualitative-experience",) or "qualitative" in sd:
        return "qualitative", _SCHEMA_TEMPLATES["qualitative"]
    if qt == "prevalence":
        return "prevalence", _SCHEMA_TEMPLATES["prevalence"]
    if "case series" in sd or "case report" in sd:
        return "case_report", _SCHEMA_TEMPLATES["case_report"]
    if "rct" in sd or "randomis" in sd or "randomiz" in sd:
        return "rct", _SCHEMA_TEMPLATES["rct"]
    if any(x in sd for x in ("cohort", "case-control", "cross-sectional", "quasi")):
        return "observational", _SCHEMA_TEMPLATES["observational"]
    # Default: RCT template is the most complete and applicable to most effectiveness reviews
    return "rct", _SCHEMA_TEMPLATES["rct"]


def _merge_schema(template_fields: list[dict], ai_fields: list[dict]) -> list[dict]:
    """
    Merge AI-suggested domain-specific fields into template.
    Template fields take priority; AI adds fields not already present.
    """
    existing = {f["field"] for f in template_fields}
    extra = [f for f in ai_fields if f.get("field") and f["field"] not in existing]
    # Ensure extra fields have description and section
    for f in extra:
        f.setdefault("description", "")
        f.setdefault("section", "General")
    return template_fields + extra


_PROSPERO_FIELDS = [
    "review_title", "review_type", "review_question", "population",
    "intervention", "comparator", "outcome", "timing_of_outcome_measurement",
    "study_design", "eligibility_criteria", "information_sources", "main_outcome_measures",
    "strategy_for_data_extraction", "data_management", "synthesis_of_results",
    "subgroup_analysis", "sensitivity_analysis", "language", "country",
    "background", "health_condition", "anticipated_start_date", "anticipated_completion_date",
    "guidance_document_reference", "funding_source", "conflicts_of_interest",
    "named_contact", "named_contact_email", "organizational_affiliation",
    "collaborative_group", "contributors", "review_status", "registration_source",
    "protocol_link", "stage_of_review_when_registered",
    "anticipated_international_prospective_register_of_systematic_reviews",
    "article_type", "registration_number", "mesh_terms", "date_of_last_amendment",
]


async def parse_pico_from_text(text: str, ai_provider: AIProvider, review_type: str = "systematic_review", framework: str = "") -> dict:
    """
    Parse a free-form research question, abstract, or uploaded protocol text
    using a Cochrane/JBI-trained methodologist approach.

    Automatically selects the best framework (PICO, SPIDER, PEO, PICOTS, etc.),
    classifies the question type, and returns a rich structured output.

    Returns a dict with keys:
      question_type, framework, review_title, review_objective, review_question,
      alternative_phrasings, methodological_cautions,
      pico, inclusion_criteria, exclusion_criteria, extraction_schema
    """
    system = """You are a Cochrane/JBI-trained systematic review methodologist with expertise in evidence synthesis methodology.

The user will provide a research topic, question, or protocol excerpt. Do the following in order:

STEP 1 — Classify the review question type as exactly one of:
  effectiveness | prevention | diagnosis | prognosis | etiology-risk | qualitative-experience | mixed-methods | prevalence | policy-service-delivery

STEP 2 — Select the best framework:
  - PICO or PICOS for effectiveness/intervention/prevention questions
  - PICOTS (adding Timing and Setting) when temporal or contextual scope matters
  - PEO (Population, Exposure, Outcome) for etiology/risk/qualitative-experience questions
  - SPIDER (Sample, Phenomenon of Interest, Design, Evaluation, Research type) for qualitative reviews
  - CoCoPop (Condition, Context, Population) for prevalence/incidence questions
  - ECLIPSE for policy/service-delivery questions
  Do NOT force PICO when the question type calls for a different framework.

STEP 3 — Fill the chosen framework's elements in the "pico" object using the framework's NATIVE element keys:
  PICO:    population, intervention, comparator, outcome
  PICOS:   population, intervention, comparator, outcome, study_design
  PCC:     population, concept, context          ← JBI scoping review framework
  SPIDER:  sample, phenomenon_of_interest, design, evaluation, research_type
  PEO:     population, exposure, outcome
  ECLIPSE: expectation, client_group, location, impact, professionals, service
  SPICE:   setting, perspective, interest, comparison, evaluation
  PICOTS:  population, intervention, comparator, outcome, timing, setting
  CoCoPop: condition, context, population

  PICO guidance (Cochrane Handbook):
  - Population: Be specific. Include age group, clinical diagnosis, setting if relevant.
  - Intervention: Name the intervention explicitly; include dose/delivery if known.
  - Comparator: OPTIONAL. Omit if the question is prevalence, etiology, or qualitative experience.
  - Outcome: Primary; Secondary; Adverse outcomes. Prespecify for prioritization — rarely a strict gate.

  PCC guidance (JBI Manual for Scoping Reviews):
  - Population: Who is the review about?
  - Concept: What is being mapped (phenomenon, activity, exposure)?
  - Context: Where or in what circumstances (setting, culture, geographic area)?

  Always include these additional fields in "pico" regardless of framework:
    study_design, health_area, language_restriction, date_from, date_to, review_type, target_registries

STEP 4 — Produce a JSON response with this exact structure (no markdown fences, no extra keys):
{
  "question_type": "effectiveness",
  "framework": "PICO",
  "review_title": "A protocol-ready review title",
  "review_objective": "To assess the effectiveness of [intervention] on [outcomes] in [population]",
  "review_question": "In [population], does [intervention], compared with [comparator], affect [outcomes]?",
  "alternative_phrasings": [
    "Broad phrasing of the question",
    "Moderate phrasing",
    "Narrow/focused phrasing"
  ],
  "methodological_cautions": "Concise paragraph on: whether comparator can remain broad; whether outcomes should be strict eligibility criteria or flexible prioritization; whether another framework would perform better; any known heterogeneity or indexing issues to warn the user about.",
  "pico": {
    "<<native element keys for the selected framework>>": "filled values",
    "study_design": "Comma-separated from: RCT, Quasi-experimental, Cohort, Case-control, Cross-sectional, Case series, Mixed methods, Qualitative, Any",
    "health_area": "Clinical specialty or subject area",
    "language_restriction": "No restriction",
    "date_from": "2000-01-01",
    "date_to": "",
    "review_type": "systematic_review",
    "target_registries": []
  },
  "inclusion_criteria": [
    "Studies involving [specific population/sample]",
    "Studies evaluating [concept/intervention/exposure]",
    "Peer-reviewed primary research",
    "Additional domain-specific criteria..."
  ],
  "exclusion_criteria": [
    "Non-human studies (animal, in vitro)",
    "Conference abstracts without full text",
    "Additional domain-specific exclusions..."
  ],
  "extraction_schema": [
    {"field": "study_design", "type": "text", "required": true},
    {"field": "sample_size", "type": "number", "required": true},
    {"field": "population_description", "type": "text", "required": true},
    {"field": "intervention_description", "type": "text", "required": true},
    {"field": "comparator_description", "type": "text", "required": false},
    {"field": "primary_outcomes", "type": "list", "required": true},
    {"field": "follow_up_duration", "type": "text", "required": false},
    {"field": "country", "type": "text", "required": false},
    {"field": "funding_source", "type": "text", "required": false}
  ]
}

IMPORTANT: The "pico" object MUST use the native element keys for the selected framework.
  Examples:
    PCC  → "pico": { "population": "...", "concept": "...", "context": "...", "study_design": "...", ... }
    SPIDER → "pico": { "sample": "...", "phenomenon_of_interest": "...", "design": "...", "evaluation": "...", "research_type": "...", "study_design": "...", ... }
    PEO  → "pico": { "population": "...", "exposure": "...", "outcome": "...", "study_design": "...", ... }
  Do NOT map non-PICO elements into PICO field names.

Additional rules:
- For extraction_schema: add domain-specific fields beyond the defaults when clearly relevant (e.g. "risk_ratio", "nnt", "adverse_events" for clinical trials; "odds_ratio", "incidence_rate" for epidemiology; "theme", "participant_quote" for qualitative reviews).
- inclusion_criteria and exclusion_criteria must be specific to the topic, not generic boilerplate.
- Generate exactly 3 alternative_phrasings: broad → moderate → narrow.
- methodological_cautions must be substantive (2–4 sentences), referencing Cochrane or JBI guidance where applicable.
- Return ONLY valid JSON. Do NOT add explanations outside the JSON."""

    framework_hint = f"\nPreferred framework (override — use this framework if suitable): {framework}" if framework else ""
    user = f"""Topic / research question:

---
{text[:4000]}
---

Review type (override): {review_type}{framework_hint}

Apply the Cochrane/JBI methodologist workflow and return the structured JSON."""

    try:
        resp = await ai_provider.guarded_complete(
            system=system, user=user,
            config=CompletionConfig(
                output_format=OutputFormat.JSON,
                budget=TokenBudget(target_json_keys=40, format=OutputFormat.JSON),
                max_continuations=1,
            ),
            json_mode=True, temperature=0.2,
        )
        result = _load_llm_json(resp.text)
        # Guarantee review_type override and required keys
        if "pico" in result:
            result["pico"]["review_type"] = review_type
            result["pico"].setdefault("target_registries", [])
        result.setdefault("question_type", "effectiveness")
        result.setdefault("framework", "PICO")
        result.setdefault("review_title", "")
        result.setdefault("review_objective", "")
        result.setdefault("review_question", "")
        result.setdefault("alternative_phrasings", [])
        result.setdefault("methodological_cautions", "")

        # Apply Cochrane/JBI study-design-specific schema template
        question_type = result.get("question_type", "effectiveness")
        study_design = result.get("pico", {}).get("study_design", "")
        template_name, template_fields = _select_template(question_type, study_design)
        ai_extra_fields = result.get("extraction_schema", [])
        result["extraction_schema"] = _merge_schema(template_fields, ai_extra_fields)
        result["schema_template"] = template_name

        return result
    except Exception as e:
        logger.error("parse_pico_from_text failed: %s", e)
        return {
            "question_type": "effectiveness",
            "framework": "PICO",
            "review_title": "",
            "review_objective": "",
            "review_question": "",
            "alternative_phrasings": [],
            "methodological_cautions": "",
            "schema_template": "rct",
            "pico": {
                "population": "",
                "intervention": "",
                "comparator": "",
                "outcome": "",
                "study_design": "",
                "health_area": "",
                "language_restriction": "No restriction",
                "date_from": "2000-01-01",
                "date_to": "",
                "review_type": review_type,
                "target_registries": [],
            },
            "inclusion_criteria": ["Human participants", "Original research", "Peer-reviewed publications"],
            "exclusion_criteria": ["Animal studies", "Non-English without translation", "Conference abstracts without full text"],
            "extraction_schema": _SCHEMA_TEMPLATES["rct"],
            "error": str(e),
        }


async def generate_protocol_document(
    pico: dict,
    ai_provider: AIProvider,
    prisma_p_data: Optional[dict] = None,
) -> str:
    """
    Generate a full PRISMA-P 2015 compliant protocol document in Markdown.
    When prisma_p_data is provided (all 17 items filled), generates a richer document.
    """
    system = """You are an expert systematic review methodologist. Generate a complete, PRISMA-P 2015 compliant protocol document in Markdown.

The document MUST include ALL 17 PRISMA-P items and these sections in order:
1. Administrative Information (Title, Registration, Authors, Funding)
2. Background and Rationale
3. Objectives (explicit PICO/SPIDER/PEO statement)
4. Methods:
   4.1 Eligibility Criteria
   4.2 Information Sources (databases, grey literature)
   4.3 Search Strategy
   4.4 Study Records (data management, selection process, data collection process)
   4.5 Data Items
   4.6 Outcomes and Prioritization
   4.7 Risk of Bias Assessment
   4.8 Data Synthesis
   4.9 Meta-Bias Assessment
   4.10 Confidence in Evidence (GRADE)
5. Strengths and Limitations
6. Dissemination Plans

Use formal academic writing. Be specific about tools, software, and statistical methods.
Within section 4.5 Data Items, use brief subsection headings for the major domains when details are available, such as Study Characteristics, Participant Characteristics, Intervention and Comparator Details, Outcome Data, and Methodological Variables.
Return ONLY the Markdown text — no JSON, no preamble."""

    population = pico.get("population", "")
    intervention = pico.get("intervention", "")
    comparator = pico.get("comparator", "")
    outcome = pico.get("outcome", "")
    study_design = pico.get("study_design", "randomised controlled trials")
    review_type = pico.get("review_type", "systematic_review")
    health_area = pico.get("health_area", "")
    date_from = pico.get("date_from", "inception")
    date_to = pico.get("date_to", "present")
    language = pico.get("language_restriction", "English")

    context_lines = [
        f"**Review Type:** {review_type}",
        "**PICO:**",
        f"- Population: {population}",
        f"- Intervention: {intervention}",
        f"- Comparator: {comparator or 'Not specified / not applicable'}",
        f"- Outcome: {outcome}",
        f"- Study Design: {study_design}",
        f"- Health Area: {health_area}",
        f"- Date Range: {date_from} to {date_to}",
        f"- Language: {language}",
    ]

    if prisma_p_data:
        admin = prisma_p_data.get("administrative", {})
        intro = prisma_p_data.get("introduction", {})
        elig = prisma_p_data.get("methods_eligibility", {})
        dc = prisma_p_data.get("methods_data_collection", {})
        synth = prisma_p_data.get("methods_synthesis", {})
        builder = prisma_p_data.get("protocol_builder", {})
        phase_snapshots = builder.get("phases", {}) if isinstance(builder, dict) else {}
        data_collection_phase = phase_snapshots.get("data_collection", {}) if isinstance(phase_snapshots, dict) else {}
        data_items_phase = phase_snapshots.get("data_items", {}) if isinstance(phase_snapshots, dict) else {}
        data_collection_content = data_collection_phase.get("content", {}) if isinstance(data_collection_phase, dict) else {}
        data_items_content = data_items_phase.get("content", {}) if isinstance(data_items_phase, dict) else {}

        if admin.get("review_title"):
            context_lines.insert(0, f"**Title:** {admin['review_title']}")
        if admin.get("registration_name"):
            registration_line = str(admin["registration_name"])
            if admin.get("registration_number"):
                registration_line += f" ({admin['registration_number']})"
            context_lines.append(f"**Registration / Submission Target:** {registration_line}")
        if admin.get("authors") and isinstance(admin["authors"], list):
            author_str = "; ".join(
                a.get("name", "") + (f" ({a.get('affiliation','')})" if a.get("affiliation") else "")
                for a in admin["authors"] if isinstance(a, dict) and a.get("name")
            )
            if author_str:
                context_lines.append(f"**Authors:** {author_str}")
        if admin.get("contributions"):
            context_lines.append(f"**Author Contributions:** {admin['contributions']}")
        if admin.get("funding_sources"):
            context_lines.append(f"**Funding:** {admin['funding_sources']}")
        if admin.get("competing_interests"):
            context_lines.append(f"**Competing Interests:** {admin['competing_interests']}")
        if admin.get("amendment_plan"):
            context_lines.append(f"**Amendment Policy:** {admin['amendment_plan']}")
        if admin.get("sponsor_role"):
            context_lines.append(f"**Sponsor Role:** {admin['sponsor_role']}")
        if intro.get("rationale"):
            context_lines.append(f"\n**Rationale:**\n{intro['rationale']}")
        if elig.get("inclusion_criteria") and isinstance(elig["inclusion_criteria"], list):
            context_lines.append("**Inclusion Criteria:** " + "; ".join(elig["inclusion_criteria"]))
        if elig.get("exclusion_criteria") and isinstance(elig["exclusion_criteria"], list):
            context_lines.append("**Exclusion Criteria:** " + "; ".join(elig["exclusion_criteria"]))
        if elig.get("databases") and isinstance(elig["databases"], list):
            context_lines.append("**Databases:** " + ", ".join(elig["databases"]))
        if dc.get("selection_process"):
            context_lines.append(f"**Selection Process:** {dc['selection_process']}")
        data_collection_narrative = _build_data_collection_narrative(data_collection_content or dc)
        if data_collection_narrative:
            context_lines.append(f"**Data Collection Process:** {data_collection_narrative}")
        data_items_markdown = _build_data_items_markdown(data_items_content, section_prefix="4.5")
        if data_items_markdown:
            context_lines.append(f"**Data Items Structure:**\n{data_items_markdown}")
        if dc.get("outcome_prioritization"):
            context_lines.append(f"**Outcome Prioritization:** {dc['outcome_prioritization']}")
        if synth.get("rob_tool"):
            context_lines.append(f"**Risk of Bias Tool:** {synth['rob_tool']}")
        if synth.get("synthesis_type"):
            context_lines.append(f"**Synthesis Type:** {synth['synthesis_type']}")
        if synth.get("effect_measure"):
            context_lines.append(f"**Effect Measure:** {synth['effect_measure']}")
        if synth.get("publication_bias_plan"):
            context_lines.append(f"**Publication Bias Plan:** {synth['publication_bias_plan']}")
        if synth.get("grade_plan"):
            context_lines.append(f"**GRADE Plan:** {synth['grade_plan']}")

        phase_snapshot_lines: dict[str, dict] = {}
        if isinstance(phase_snapshots, dict):
            for phase_id, snapshot in phase_snapshots.items():
                if not isinstance(snapshot, dict):
                    continue
                content = snapshot.get("content", {})
                if not isinstance(content, dict):
                    continue
                compact = {
                    key: value for key, value in content.items()
                    if key != "__type" and value not in ("", [], {}, None)
                }
                if compact:
                    phase_snapshot_lines[phase_id] = compact
        if phase_snapshot_lines:
            context_lines.append("**Protocol Builder Section Data:**\n" + json.dumps(phase_snapshot_lines, indent=2)[:4000])

    user = "Generate a complete PRISMA-P 2015 protocol document for:\n\n" + "\n".join(context_lines)

    return await ai_provider.complete(
        system=system,
        user=user,
        temperature=0.3,
        max_tokens=8192,
    )


async def map_to_prospero_fields(
    pico: dict,
    protocol_text: str,
    ai_provider: AIProvider,
) -> dict:
    """
    Map PICO data and protocol text to all 40 PROSPERO fields.
    Returns a structured dict ready for PROSPERO form submission.
    """
    system = """You are an expert in PROSPERO systematic review registration.
Extract and map all required PROSPERO fields from the provided PICO data and protocol text.
Return ONLY valid JSON with these exact field names (no markdown fences):

{
  "review_title": "...",
  "review_type": "Systematic review" or "Meta-analysis",
  "review_question": "...",
  "population": "...",
  "intervention": "...",
  "comparator": "...",
  "outcome": "...",
  "timing_of_outcome_measurement": "...",
  "study_design": "...",
  "eligibility_criteria": "...",
  "information_sources": "...",
  "main_outcome_measures": "...",
  "strategy_for_data_extraction": "...",
  "data_management": "...",
  "synthesis_of_results": "...",
  "subgroup_analysis": "...",
  "sensitivity_analysis": "...",
  "language": "...",
  "country": "...",
  "background": "...",
  "health_condition": "...",
  "anticipated_start_date": "...",
  "anticipated_completion_date": "...",
  "funding_source": "None declared",
  "conflicts_of_interest": "None declared",
  "review_status": "Ongoing",
  "stage_of_review_when_registered": "Review planning stage",
  "article_type": "Systematic review",
  "copy_paste_guide": "Step-by-step instructions for submitting to PROSPERO at https://www.crd.york.ac.uk/prospero/"
}"""

    user = f"""PICO Data:
{json.dumps(pico, indent=2)}

Protocol Document (excerpt):
{protocol_text[:3000]}

Map ALL PROSPERO fields. For fields not determinable from the PICO/protocol, provide reasonable placeholder text."""

    try:
        raw = await ai_provider.complete(system=system, user=user, json_mode=True, temperature=0.2)
        result = _load_llm_json(raw)
        # Ensure copy_paste_guide is always present
        if "copy_paste_guide" not in result:
            result["copy_paste_guide"] = (
                "1. Go to https://www.crd.york.ac.uk/prospero/\n"
                "2. Click 'Register a new review'\n"
                "3. Copy each field value from this document into the corresponding PROSPERO field\n"
                "4. The review will be assessed within 5 working days\n"
                "5. You will receive a PROSPERO ID (format CRDxxxxxxxxxxxx) upon registration"
            )
        return result
    except Exception as e:
        logger.error("PROSPERO field mapping failed: %s", e)
        return {"error": str(e), "pico": pico}


async def map_to_campbell_fields(
    pico: dict,
    protocol_text: str,
    ai_provider: AIProvider,
) -> dict:
    """
    Map protocol to Campbell Collaboration template structure.
    Emphasizes broader study design, grey literature, and 8 Campbell coordinating groups.
    """
    system = """You are an expert in Campbell Collaboration systematic review standards.
Map the PICO and protocol to the Campbell Collaboration template structure.
Return ONLY valid JSON:
{
  "title": "...",
  "coordinating_group": "...",  // One of: Crime and Justice, Disability, Education, International Development, Knowledge Translation and Implementation, Methods, Social Welfare, User Engagement
  "background": "...",
  "objectives": "...",
  "search_strategy": "...",
  "selection_criteria": "...",
  "data_collection": "...",
  "data_synthesis": "...",
  "study_designs_included": [...],  // RCTs, quasi-experimental, interrupted time series, etc.
  "grey_literature_sources": [...],
  "contact_person": "...",
  "review_type": "..."
}"""

    user = f"""PICO:
{json.dumps(pico, indent=2)}

Protocol Excerpt:
{protocol_text[:2000]}

Map to Campbell Collaboration template."""

    try:
        raw = await ai_provider.complete(system=system, user=user, json_mode=True, temperature=0.2)
        return _load_llm_json(raw)
    except Exception as e:
        logger.error("Campbell field mapping failed: %s", e)
        return {"error": str(e)}


async def generate_prisma_p_checklist(protocol_text: str) -> dict:
    """
    Check the generated protocol text against all 26 PRISMA-P items.
    Returns completion status for each item.
    """
    checklist = {}
    for item_id, description in _PRISMA_P_ITEMS.items():
        # Heuristic: check if key words from the description appear in the protocol
        keywords = description.lower().split()[:4]  # First 4 words as keywords
        found = any(kw in protocol_text.lower() for kw in keywords if len(kw) > 4)
        checklist[item_id] = {
            "description": description,
            "status": "complete" if found else "missing",
            "location": "See protocol document",
        }
    return checklist


async def generate_database_search_strings(pico: dict, ai_provider: AIProvider) -> dict:
    """
    Generate ready-to-use search strings for 9 databases.
    Returns dict with database names as keys.
    """
    system = """You are an expert medical librarian specializing in systematic review search strategies.
Generate complete, ready-to-use search strings for ALL 9 databases listed below.

Include:
- Correct field codes (PubMed: [tiab][mesh], Scopus: TITLE-ABS-KEY, Ovid: .mp. .sh.)
- MeSH terms (PubMed), Emtree (Embase), CINAHL Headings
- Truncation (* for PubMed/Embase, $ for CINAHL)
- Boolean operators (AND/OR/NOT)
- Date and language limits
- Publication type filters where appropriate

For databases requiring institutional access (EMBASE, PsycINFO, CINAHL), include a note.

Return ONLY valid JSON:
{
  "pubmed": "full search string...",
  "scopus": "full search string...",
  "embase_ovid": "full search string... [NOTE: Requires institutional Ovid subscription]",
  "psycinfo_ebsco": "full search string... [NOTE: Requires institutional EBSCO subscription]",
  "cinahl_ebsco": "full search string... [NOTE: Requires institutional EBSCO subscription]",
  "cochrane_central": "full search string...",
  "clinicaltrials": "simple terms for clinicaltrials.gov advanced search...",
  "who_ictrp": "simple terms for WHO ICTRP search...",
  "eric": "search string for ERIC education database...",
  "notes": "General notes about the search strategy and recommended date ranges"
}"""

    user = f"""PICO for search string generation:
Population: {pico.get('population', '')}
Intervention: {pico.get('intervention', '')}
Comparator: {pico.get('comparator', '')}
Outcome: {pico.get('outcome', '')}
Study Design: {pico.get('study_design', '')}
Date From: {pico.get('date_from', 'inception')}
Language: {pico.get('language_restriction', 'English')}
Health Area: {pico.get('health_area', '')}

Generate comprehensive search strings for all 9 databases."""

    try:
        resp = await ai_provider.guarded_complete(
            system=system, user=user,
            config=CompletionConfig(
                output_format=OutputFormat.JSON,
                budget=TokenBudget(target_json_keys=50, format=OutputFormat.JSON),
                max_continuations=1,
            ),
            json_mode=True, temperature=0.2,
        )
        return _load_llm_json(resp.text)
    except Exception as e:
        logger.error("Search string generation failed: %s", e)
        return {"error": str(e)}


async def register_on_osf(
    protocol_text: str,
    pico: dict,
    osf_token: str,
) -> dict:
    """
    Programmatically register the protocol on OSF.
    Returns {osf_id, doi, url, status} or {error: ...}

    NOTE: osf_token must never be logged.
    """
    headers = {
        "Authorization": f"Bearer {osf_token}",
        "Content-Type": "application/vnd.api+json",
    }

    async with httpx.AsyncClient(timeout=httpx.Timeout(30.0)) as client:
        # Step 1: Create an OSF project node
        project_payload = {
            "data": {
                "type": "nodes",
                "attributes": {
                    "title": f"SR Protocol: {pico.get('population', '')[:100]} — {pico.get('intervention', '')[:80]}",
                    "category": "project",
                    "description": f"Systematic review protocol. PICO: {pico.get('population', '')} | {pico.get('intervention', '')} | {pico.get('outcome', '')}",
                    "public": False,
                }
            }
        }
        try:
            r = await client.post(
                "https://api.osf.io/v2/nodes/",
                json=project_payload,
                headers=headers,
            )
            r.raise_for_status()
            node_id = r.json()["data"]["id"]
        except Exception as e:
            logger.error("OSF node creation failed: %s", e)  # Do NOT log token
            return {"error": f"Failed to create OSF project: {r.status_code if 'r' in dir() else str(e)}"}

        # Step 2: Find the Generalized Systematic Review schema
        try:
            r = await client.get(
                "https://api.osf.io/v2/schemas/registrations/",
                params={"filter[name]": "Preregistration"},
                headers=headers,
            )
            r.raise_for_status()
            schemas = r.json().get("data", [])
            schema_id = schemas[0]["id"] if schemas else None
        except Exception:
            schema_id = None

        if not schema_id:
            return {
                "node_id": node_id,
                "status": "draft_only",
                "url": f"https://osf.io/{node_id}/",
                "message": "Project created on OSF. Schema not found — please complete registration manually at https://osf.io/registries/",
            }

        # Step 3: Create draft registration
        draft_payload = {
            "data": {
                "type": "draft_registrations",
                "attributes": {
                    "registration_supplement": {
                        "q1": protocol_text[:2000],
                    }
                },
                "relationships": {
                    "registration_schema": {
                        "data": {"type": "registration-schemas", "id": schema_id}
                    }
                }
            }
        }
        try:
            r = await client.post(
                f"https://api.osf.io/v2/nodes/{node_id}/draft_registrations/",
                json=draft_payload,
                headers=headers,
            )
            draft_id = r.json().get("data", {}).get("id", "")
        except Exception as e:
            logger.warning("OSF draft creation: %s", e)
            draft_id = ""

        return {
            "node_id": node_id,
            "draft_id": draft_id,
            "status": "draft",
            "url": f"https://osf.io/{node_id}/",
            "message": "Draft registration created on OSF. Please review and submit at the URL provided.",
        }


def _normalize_heading(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", " ", (value or "").lower()).strip()


def _estimate_line_units(text: str, kind: str) -> float:
    words = re.findall(r"\S+", text or "")
    count = len(words)
    if kind == "heading1":
        return 2.5 + max(0, count - 6) * 0.08
    if kind == "heading2":
        return 2.1 + max(0, count - 8) * 0.05
    if kind == "heading3":
        return 1.7 + max(0, count - 8) * 0.05
    if kind == "bullet":
        return max(1.0, count / 9.0)
    if kind == "code":
        return max(1.0, count / 7.0)
    if kind == "blank":
        return 0.4
    return max(1.0, count / 11.0)


def _build_protocol_heading_pages(protocol_text: str, lines_per_page: float = 42.0) -> dict[str, int]:
    heading_pages: dict[str, int] = {}
    cursor = 0.0
    in_code_block = False

    for raw_line in protocol_text.splitlines():
        line = raw_line.strip()
        if line.startswith("```"):
            in_code_block = not in_code_block
            cursor += 0.4
            continue

        if not line:
            cursor += _estimate_line_units("", "blank")
            continue

        kind = "paragraph"
        content = line
        if in_code_block:
            kind = "code"
        elif line.startswith("# "):
            kind = "heading1"
            content = line[2:].strip()
        elif line.startswith("## "):
            kind = "heading2"
            content = line[3:].strip()
        elif line.startswith("### "):
            kind = "heading3"
            content = line[4:].strip()
        elif line.startswith("#### "):
            kind = "heading4"
            content = line[5:].strip()
        elif line.startswith("- ") or line.startswith("* "):
            kind = "bullet"
            content = line[2:].strip()

        page_no = max(1, int(cursor // lines_per_page) + 1)
        if kind.startswith("heading"):
            heading_pages.setdefault(_normalize_heading(content), page_no)

        cursor += _estimate_line_units(content, kind)

    return heading_pages


def _resolve_page_number(
    aliases: list[str],
    heading_pages: dict[str, int],
    fallback_page: int,
) -> int:
    normalized_aliases = [_normalize_heading(alias) for alias in aliases]

    for alias in normalized_aliases:
        if alias in heading_pages:
            return heading_pages[alias]

    for alias in normalized_aliases:
        for heading, page_no in heading_pages.items():
            if alias and (alias in heading or heading in alias):
                return page_no

    return fallback_page


def _build_prisma_p_page_refs(protocol_text: str) -> dict[str, str]:
    heading_pages = _build_protocol_heading_pages(protocol_text)
    admin_page = _resolve_page_number(["administrative information", "title"], heading_pages, 1)
    intro_page = _resolve_page_number(["background and rationale", "introduction"], heading_pages, admin_page)
    objectives_page = _resolve_page_number(["objectives"], heading_pages, intro_page)
    methods_page = _resolve_page_number(["methods"], heading_pages, objectives_page)

    refs: dict[str, str] = {}
    for item_no, aliases in _PRISMA_P_PAGE_ALIASES.items():
        fallback_page = methods_page
        if item_no in {"1a", "1b", "2", "3a", "3b", "4", "5a", "5b", "5c"}:
            fallback_page = admin_page
        elif item_no == "6":
            fallback_page = intro_page
        elif item_no == "7":
            fallback_page = objectives_page
        refs[item_no] = str(_resolve_page_number(aliases, heading_pages, fallback_page))
    return refs


def _protocol_title(pico: dict, prisma_p_data: Optional[dict] = None) -> str:
    admin = (prisma_p_data or {}).get("administrative", {})
    if isinstance(admin, dict):
        title = str(admin.get("review_title") or "").strip()
        if title:
            return title

    population = str(pico.get("population") or "").strip()
    intervention = str(pico.get("intervention") or "").strip()
    outcome = str(pico.get("outcome") or "").strip()
    if population or intervention or outcome:
        parts = [part for part in (population, intervention, outcome) if part]
        return "Protocol: " + " / ".join(parts[:3])
    return "Systematic Review Protocol"


def _append_field_run(paragraph, instruction: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    display_text = OxmlElement("w:t")
    display_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_separate)
    run._r.append(display_text)
    run._r.append(fld_end)


def _add_page_number_footer(doc) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.text = "Page "
    _append_field_run(paragraph, "PAGE")


async def generate_protocol_docx(
    protocol_text: str,
    pico: dict,
    prisma_p_data: Optional[dict] = None,
) -> bytes:
    """
    Generate a downloadable .docx of the protocol following PRISMA-P structure.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx required: pip install python-docx")

    doc = Document()
    _add_page_number_footer(doc)

    # Title
    title_para = doc.add_heading(_protocol_title(pico, prisma_p_data), level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # PICO summary box
    doc.add_heading("Review Overview (PICO)", level=1)
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    cells = [
        ("Population", pico.get("population", "")),
        ("Intervention", pico.get("intervention", "")),
        ("Comparator", pico.get("comparator", "")),
        ("Outcome", pico.get("outcome", "")),
        ("Study Design", pico.get("study_design", "")),
    ]
    for i, (label, value) in enumerate(cells):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
    doc.add_paragraph("")

    # Protocol document body
    doc.add_heading("Full Protocol", level=1)
    for line in protocol_text.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph("")
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("#### "):
            doc.add_heading(line[5:], level=4)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


async def generate_prisma_p_checklist_docx(
    protocol_text: str,
    pico: dict,
    prisma_p_data: Optional[dict] = None,
) -> bytes:
    """
    Generate a separate PRISMA-P checklist .docx with reported protocol page numbers.
    """
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx required: pip install python-docx")

    doc = Document()
    _add_page_number_footer(doc)

    title_para = doc.add_heading("PRISMA-P 2015 Checklist", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    protocol_title_para = doc.add_paragraph()
    protocol_title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    protocol_title_run = protocol_title_para.add_run(_protocol_title(pico, prisma_p_data))
    protocol_title_run.italic = True

    note_para = doc.add_paragraph()
    note_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_run = note_para.add_run("Reported page numbers refer to the exported protocol manuscript.")
    note_run.font.size = Pt(9)

    page_refs = _build_prisma_p_page_refs(protocol_text)

    checklist_table = doc.add_table(rows=1, cols=4)
    checklist_table.style = "Table Grid"
    header_cells = checklist_table.rows[0].cells
    header_cells[0].text = "Section and topic"
    header_cells[1].text = "Item No"
    header_cells[2].text = "Checklist item"
    header_cells[3].text = "Reported on page No."

    for row_def in _PRISMA_P_CHECKLIST_TEMPLATE_ROWS:
        row_cells = checklist_table.add_row().cells
        if row_def["kind"] in {"section", "group"}:
            merged = row_cells[0]
            for idx in range(1, 4):
                merged = merged.merge(row_cells[idx])
            merged.text = row_def["label"]
            if merged.paragraphs and merged.paragraphs[0].runs:
                merged.paragraphs[0].runs[0].bold = True
            continue

        row_cells[0].text = row_def["topic"]
        row_cells[1].text = row_def["item_no"]
        row_cells[2].text = row_def["checklist_item"]
        row_cells[3].text = page_refs.get(row_def["item_no"], "")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── New AI-first Protocol Builder functions ───────────────────────────────────

async def generate_background_draft(
    query: str,
    ai_provider: AIProvider,
    section: str = "background",
    review_type: str = "systematic_review",
) -> dict:
    """
    Generate a background or rationale section for a systematic review protocol.
    Uses AI knowledge to write a well-informed draft based on the research topic.

    section: "background" | "rationale"
    Returns: { draft, summary }
    """
    if section == "rationale":
        system = """You are a systematic review methodologist writing the 'Rationale & Why This Review Is Needed' section of a protocol.

Write 2–3 focused paragraphs that:
1. Summarise what is already known about this topic (existing evidence, prior reviews)
2. Identify the specific gap in the evidence this review will fill
3. Explain why this review is important NOW and who will benefit from its findings

Use formal academic writing (third person, past tense for existing evidence, future tense for review plans).
Reference relevant study designs, populations, or methodological issues where applicable.
Do NOT fabricate specific paper citations — use general references like "existing reviews have shown..." or "the literature suggests...".

Return a plain Markdown string (no JSON, no fences)."""

        user = f"""Research topic: {query}
Review type: {review_type}

Write the 'Rationale and Why This Review Is Needed' section."""

    else:  # background
        system = """You are a systematic review methodologist writing a comprehensive 'Background' section for a PRISMA-P 2015 compliant systematic review protocol.

The background must be structured under EXACTLY these four subheadings, in this order:

## The Problem, Condition, or Issue
Write 3–4 paragraphs covering:
- Definition and scope of the condition, problem, or phenomenon
- Global or regional prevalence, incidence, and burden (epidemiological data)
- Affected populations, settings, and demographic patterns
- Consequences: clinical, social, economic, and quality-of-life impacts
- Risk factors or determinants relevant to the review topic

## The Intervention (or Exposure / Phenomenon of Interest)
Write 3–4 paragraphs covering:
- Description of the intervention, exposure, or phenomenon under review
- How it is currently delivered, implemented, or studied
- Variations, modalities, or subtypes relevant to the review scope
- Context in which the intervention is typically applied (healthcare setting, policy, community)
- Existing guidelines or recommendations, if any

## How the Intervention Could Work
Write 2–3 paragraphs covering:
- The theoretical mechanism of action or logical pathway through which the intervention is expected to produce change
- Biological, psychological, social, or structural mechanisms (as appropriate to the topic)
- Evidence from experimental or observational studies suggesting plausible effects
- Potential moderating or mediating factors

## Why It Is Important to Do This Review
Write 2–3 paragraphs covering:
- Overview of prior systematic reviews on this topic (if any), and their limitations (scope, quality, recency)
- The specific evidence gap(s) this review will address
- Decision-makers or stakeholders who will benefit from the findings
- The potential policy, clinical, or practice implications of this review

WRITING STANDARDS:
- Target total length: 1000–1500 words across all four sections
- Formal academic prose, third person, present or past tense as appropriate
- Do NOT fabricate specific statistics, author names, or DOIs — use hedged language: "evidence suggests...", "studies indicate...", "it is estimated that...", "a body of literature has shown..."
- Each subheading must be written as a Markdown ## heading exactly as shown above
- No bullet lists in the body — prose paragraphs only

Return a plain Markdown string with the four ## subheadings and body paragraphs. No JSON, no code fences."""

        user = f"""Research topic: {query}
Review type: {review_type}

Write the full structured 'Background' section with all four subheadings."""

    try:
        resp = await ai_provider.guarded_complete(
            system=system, user=user,
            config=CompletionConfig(
                output_format=OutputFormat.PROSE,
                budget=TokenBudget(target_words=1200),
                max_continuations=2,
            ),
            temperature=0.4,
        )
        draft = resp.text
        # Generate a summary for the chat thread
        if section == "background":
            summary = (
                f"I've written a detailed background section (~1000+ words) for your review on '{query[:80]}'. "
                "It covers four structured subsections: (1) the problem/condition and its burden, "
                "(2) the intervention or phenomenon of interest, "
                "(3) the mechanism of action, and "
                "(4) why this review is needed. "
                "Edit the text directly or tell me what to expand, change, or emphasise."
            )
        else:
            summary = f"I've drafted a {section} section based on the topic '{query[:80]}'. It covers the key context, existing evidence landscape, and the gap this review will address. Review it above and let me know what to add, emphasise, or change."
        return {"draft": draft.strip(), "summary": summary}
    except Exception as e:
        logger.error("generate_background_draft failed: %s", e)
        return {
            "draft": f"[Background draft could not be generated: {e}]",
            "summary": "There was an error generating the draft. Please try again.",
        }


async def generate_review_question_from_elements(
    framework: str,
    elements: dict,
    ai_provider: AIProvider,
    feedback: str = "",
    review_type: str = "systematic_review",
) -> dict:
    """
    Generate a structured research question from framework elements (PICO/SPIDER/PCC/etc.).
    This is the REVERSE of parsePicoFromText — elements first, then question.

    Returns: { review_question, alternative_phrasings, methodological_cautions }
    """
    element_str = "\n".join(f"- {k.capitalize()}: {v}" for k, v in elements.items() if v)

    feedback_clause = f"\n\nUser feedback for this revision: {feedback}" if feedback.strip() else ""

    system = """You are a Cochrane/JBI-trained systematic review methodologist.
Given structured framework elements, formulate a precise, answerable research question.

Rules:
- The research question must flow naturally from the elements — do NOT invent elements not listed
- Use the appropriate question format for the framework (PICO: "In [P], does [I] compared to [C] affect [O]?"; PCC: "What [concept] exists in [context] for [population]?"; SPIDER: "What are the experiences/perceptions of [S] regarding [PI] as studied by [D]?"; etc.)
- Generate exactly 3 alternative phrasings: broad → moderate → narrow
- methodological_cautions: 2–3 sentences on potential heterogeneity, indexing issues, or framework limitations
- Return ONLY valid JSON (no markdown fences):
{
  "review_question": "...",
  "alternative_phrasings": ["broad", "moderate", "narrow"],
  "methodological_cautions": "..."
}"""

    user = f"""Framework: {framework}
Review type: {review_type}

Elements:
{element_str}{feedback_clause}

Generate the research question."""

    try:
        resp = await ai_provider.guarded_complete(
            system=system, user=user,
            config=CompletionConfig(
                output_format=OutputFormat.JSON,
                budget=TokenBudget(target_json_keys=20, format=OutputFormat.JSON),
                max_continuations=1,
            ),
            json_mode=True, temperature=0.2,
        )
        result = _load_llm_json(resp.text)
        result.setdefault("review_question", "")
        result.setdefault("alternative_phrasings", ["", "", ""])
        result.setdefault("methodological_cautions", "")
        return result
    except Exception as e:
        logger.error("generate_review_question_from_elements failed: %s", e)
        return {
            "review_question": "",
            "alternative_phrasings": ["", "", ""],
            "methodological_cautions": "",
            "error": str(e),
        }


_PHASE_SYSTEMS: dict[str, str] = {
    "objectives": """You are a systematic review methodologist writing the 'Objectives' section.
Draft 1–2 formal objective statements. Use standard SR language:
"This systematic review will [collate/synthesise/characterise] evidence on [X] in order to [Y]."
Also include the explicit research question restated as a formal objective.
Keep it to 100–150 words. Return plain text (no JSON).""",

    "outcomes": """You are a systematic review methodologist defining outcomes for a review.
Return JSON: { "primary": ["..."], "secondary": ["..."], "chat_reply": "..." }
- primary: 2–3 primary outcomes (measurable, time-specified where possible)
- secondary: 3–5 secondary outcomes
- chat_reply: 2–3 sentence explanation of your choices and how they relate to the PICO
Base outcomes on the provided PICO/framework context.""",

    "eligibility": """You are a Cochrane/JBI-trained methodologist writing eligibility criteria.
Return JSON: { "inclusion": ["..."], "exclusion": ["..."], "chat_reply": "..." }
- inclusion: 4–6 specific inclusion criteria (population, design, setting, language, date range)
- exclusion: 4–6 exclusion criteria (mirror of inclusion + non-primary data, grey literature if applicable)
- chat_reply: brief explanation referencing MECIR standards
Be specific to the topic — no generic boilerplate.""",

    "search_sources": """You are a systematic review librarian recommending databases.
Return JSON: { "databases": ["..."], "grey_literature": "...", "chat_reply": "..." }
- databases: 5–8 database names appropriate for this topic and review type
- grey_literature: 1 sentence describing grey literature strategy (trial registries, dissertations, conference abstracts, etc.)
- chat_reply: 2–3 sentences explaining why these databases were chosen (coverage of health area, indexing)
Common databases: PubMed/MEDLINE, Embase, Cochrane CENTRAL, Web of Science, Scopus, PsycINFO, CINAHL, WHO ICTRP, ProQuest Dissertations, OpenGrey""",

    "search_string": """You are a medical librarian creating a systematic review search string.
Return JSON: { "database": "PubMed/MEDLINE", "search_string": "...", "chat_reply": "..." }
- database: the primary database (choose PubMed/MEDLINE by default unless context indicates otherwise)
- search_string: a complete, ready-to-use boolean search string with MeSH terms, free-text synonyms, truncation (*), field codes ([tiab][mesh]), AND/OR operators, and date/language limits
- chat_reply: 2–3 sentences explaining the strategy structure and any notable choices
The search string must be complete enough to run directly in the specified database.""",

    "screening": """You are a systematic review methodologist writing the screening workflow.
Return JSON: { "selection_process": "...", "data_management_tool": "...", "chat_reply": "..." }
- selection_process: 2–3 sentences describing independent dual screening (title/abstract then full text), how conflicts are resolved, and who screens
- data_management_tool: name of the software (Covidence, Rayyan, or Endnote — recommend Covidence)
- chat_reply: brief explanation with reference to PRISMA 2020 screening guidance""",

    "extraction": """You are a systematic review methodologist designing a data extraction form.
Return JSON: { "schema": [{"field": "...", "type": "text|number|list|boolean", "required": true|false, "description": "...", "section": "Methods|Participants|Interventions|Outcomes|Results"}], "chat_reply": "..." }
Generate 10–15 fields appropriate for this review type and PICO.
Organise into sections: Methods, Participants, Interventions, Outcomes, Results.
chat_reply: 2–3 sentences explaining the schema design and any domain-specific fields.""",

    "bias_synthesis": """You are a systematic review methodologist recommending risk of bias and synthesis approaches.
Return JSON: {
  "rob_tool": "...",
  "rob_level": "study|outcome|both",
  "rob_rationale": "...",
  "synthesis_type": "quantitative|narrative|both",
  "effect_measure": "...",
  "i2_threshold": "50%",
  "synthesis_rationale": "...",
  "chat_reply": "..."
}
Choose the RoB tool based on study design:
- RCTs → RoB 2.0 (Cochrane Handbook Ch. 8)
- Non-randomised → ROBINS-I (Sterne et al. BMJ 2016)
- Diagnostic → QUADAS-2
- Qualitative → JBI QARI or CASP
- Mixed → MMAT or multiple tools
chat_reply: 2–3 sentences justifying the RoB tool and synthesis approach.""",

    "dissemination": """You are a systematic review methodologist drafting the dissemination plan.
Return JSON: {
  "review_title": "...",
  "registration_name": "PROSPERO",
  "target_journals": ["...", "..."],
  "planned_outputs": "...",
  "chat_reply": "..."
}
- review_title: a concise, PROSPERO-ready title following format: "[Intervention/topic] for [condition] in [population]: a systematic review [and meta-analysis]" (≤25 words)
- registration_name: recommend PROSPERO for health topics, Campbell for social sciences
- target_journals: 2–3 appropriate journals (e.g. Systematic Reviews, BMJ, Campbell Systematic Reviews, Cochrane Database of Systematic Reviews)
- planned_outputs: 1 sentence on planned dissemination (peer-reviewed article + plain language summary)
- chat_reply: 2–3 sentences on registration timing and dissemination strategy""",

    "background": """You are a systematic review methodologist revising the 'Background' section of a protocol.
The current background draft and conversation history are provided in the context.
Revise the background based on the user's latest feedback, preserving what is good.

The background MUST retain these four ## subheadings (add any that are missing):
## The Problem, Condition, or Issue
## The Intervention (or Exposure / Phenomenon of Interest)
## How the Intervention Could Work
## Why It Is Important to Do This Review

Return the revised section in the response format requested by the caller.
- text: revised background (1000–1500 words) with all four ## subheadings intact, prose paragraphs only
- chat_reply: 1–2 sentences summarising what was revised
Maintain formal scholarly academic tone.""",

    "rationale": """You are a systematic review methodologist writing the 'Rationale & Gap' section of a protocol.
The current rationale draft and conversation history are provided in the context.
Revise the rationale based on the user's latest feedback, preserving what is good.
Return the revised section in the response format requested by the caller.
- text: revised 2–3 paragraph rationale explaining: (1) existing systematic reviews on this topic and their limitations or outdatedness, (2) the specific evidence gap this review fills, (3) why this review is needed now and who will benefit
- chat_reply: 1–2 sentences summarising what was revised
Reference Campbell Collaboration, Cochrane, or JBI guidance where relevant.""",

    "protocol_chat": """You are a systematic review methodologist reviewing a complete protocol draft.
The full protocol document and conversation history are provided in the context.
Answer the user's question or act on their request — be specific, referencing protocol sections by name.
Return JSON: { "chat_reply": "..." }
- chat_reply: your full response (3–6 sentences). If the user requests a change, describe exactly what to update and in which section. If answering a question, be precise and cite the relevant section of the protocol.""",

    "review_setup": """You are a systematic review methodologist helping the researcher select the right review type and question framework.
Return JSON: {
  "review_family_rationale": "...",
  "recommended_framework": "PICO",
  "framework_rationale": "...",
  "downstream_notes": "...",
  "chat_reply": "..."
}
- review_family_rationale: 1–2 sentences explaining why this review family is appropriate for the stated topic
- recommended_framework: suggest PICO (intervention), PECO (exposure), PEO (qualitative/exposure), PCC (scoping), SPIDER (qualitative), ECLIPSE (policy), SPICE (service evaluation)
- framework_rationale: 1–2 sentences explaining the framework choice
- downstream_notes: brief note on how the selection affects effect measures, RoB tools, and certainty assessment
- chat_reply: concise welcome message acknowledging the review family and framework choice. Also explain what each applicable review type (epidemiological/observational, qualitative, scoping, systematic) would look like for this specific topic — help the researcher choose the most appropriate one.""",

    "research_question": """You are a Cochrane/JBI-trained systematic review methodologist helping the researcher define their research question framework and elements.

Your role in this phase is conversational and advisory:
1. Help the researcher choose the right framework (PICO, PCC, SPIDER, PEO, ECLIPSE, SPICE) by explaining the purpose of each and which fits their topic.
2. Explain what each element of the chosen framework means in the context of their specific topic.
3. Suggest concrete values for empty or weak elements.
4. Identify potential problems with how elements are phrased (too broad, too narrow, confounded, missing key terms).
5. Explain how the framework connects to the rest of the review (search strategy, eligibility criteria, synthesis).

Framework decision guidance:
- PICO/PICOS: when there is a clear intervention and you want to compare outcomes
- PCC (JBI scoping): when you want to map evidence on a concept in a population across contexts — do NOT use PICO if there is no control group or comparator
- SPIDER: qualitative or mixed-methods reviews about experiences, perceptions, meanings
- PEO: observational/epidemiological association questions (exposure → outcome, no intervention)
- ECLIPSE: policy, management, service-delivery questions
- SPICE: service evaluation

Return JSON: {
  "suggested_elements": {},
  "chat_reply": "..."
}
- suggested_elements: a dict of framework element keys → suggested text (only include elements you are suggesting changes for; empty dict if no changes)
- chat_reply: conversational, warm, expert response (3–6 sentences). Explain your reasoning clearly. If suggesting element changes, describe WHY. Reference JBI or Cochrane guidance by name when applicable.""",

    "search_strategy": """You are a medical librarian designing a comprehensive systematic review search strategy.
Return JSON: {
  "primary_database": "PubMed/MEDLINE",
  "primary_search_string": "...",
  "date_limits": "...",
  "language_restrictions": "...",
  "publication_type_filters": "...",
  "platform_notes": "...",
  "translation_notes": "...",
  "handsearching": "...",
  "press_review": "...",
  "deduplication_plan": "...",
  "chat_reply": "..."
}
- primary_search_string: complete boolean string with MeSH/controlled vocab + free text synonyms, truncation (*), field tags ([tiab][mesh]), AND/OR, ready to run
- date_limits: specify from/to years and rationale
- language_restrictions: English only or multilingual with translation plan
- publication_type_filters: e.g. exclude case reports, editorials, conference abstracts
- platform_notes: note any platform-specific syntax differences (Ovid vs PubMed vs Embase)
- translation_notes: how string will be adapted across databases
- handsearching: journals to handsearch + citation chasing plan
- press_review: whether PRESS peer review of search string is planned
- deduplication_plan: tool (Endnote/Covidence) and deduplication workflow""",

    "records_management": """You are a systematic review methodologist writing the records management and deduplication section.
Return JSON: {
  "deduplication_tool": "...",
  "deduplication_method": "...",
  "record_keeping": "...",
  "automation_tools": "...",
  "chat_reply": "..."
}
- deduplication_tool: recommend Endnote, Covidence, or Rayyan
- deduplication_method: automated + manual deduplication steps
- record_keeping: how records of search results will be stored and documented for PRISMA flow
- automation_tools: any AI-assisted screening tools planned (optional)""",

    "data_collection": """You are a systematic review methodologist writing the data collection process section.
Return JSON: {
  "extraction_method": "...",
  "extraction_team": "...",
  "pilot_testing": "...",
  "disagreement_resolution": "...",
  "software": "...",
  "author_contact": "...",
  "chat_reply": "..."
}
- extraction_method: a polished paragraph suitable for the manuscript that explains the overall data extraction process in prose, including whether extraction is independent/duplicate or single with verification
- extraction_team: who extracts (e.g. two independent reviewers, student + senior)
- pilot_testing: calibration on a sample before full extraction
- disagreement_resolution: consensus, third reviewer arbitration
- software: Covidence, REDCap, Excel, etc.
- author_contact: will corresponding authors be contacted for missing data""",

    "data_items": """You are a systematic review methodologist writing the data items section.
Return JSON: {
  "study_characteristics": ["..."],
  "participant_characteristics": ["..."],
  "intervention_characteristics": ["..."],
  "outcome_items": ["..."],
  "methodological_items": ["..."],
  "chat_reply": "..."
}
- study_characteristics: design, country, setting, funding, year published
- participant_characteristics: age, sex, diagnosis criteria, severity, comorbidities
- intervention_characteristics: type, dose, duration, delivery mode, provider, comparator details
- outcome_items: primary outcomes (measure, instrument, time points), secondary outcomes
- methodological_items: sample size, randomisation, blinding, follow-up rate, ITT analysis
- These lists will be rendered into manuscript subheadings in this order: Study Characteristics, Participant Characteristics, Intervention and Comparator Details, Outcome Data, Methodological Variables
Each list: 4–6 items appropriate for this review type and PICO.""",

    "rob_assessment": """You are a systematic review methodologist recommending risk of bias assessment tools.
Select based on the review_family field in the PICO context.
Return JSON: {
  "primary_tool": "...",
  "primary_tool_rationale": "...",
  "secondary_tool": "...",
  "assessment_level": "study",
  "domains": ["..."],
  "calibration_plan": "...",
  "chat_reply": "..."
}
Tool selection rules:
- intervention RCTs → RoB 2.0 (Cochrane Handbook Ch. 8)
- intervention non-randomised → ROBINS-I (Sterne et al. BMJ 2016)
- diagnostic accuracy → QUADAS-2
- prognostic → PROBAST or QUIPS
- qualitative → CASP qualitative checklist or JBI QARI
- mixed methods → MMAT
- prevalence/descriptive → JBI Critical Appraisal Tools
domains: list the specific assessment domains for the chosen tool
calibration_plan: pilot assessment on 5–10 papers with consensus before full assessment""",

    "synthesis_plan": """You are a systematic review methodologist writing the synthesis methods section.
Return JSON: {
  "synthesis_type": "quantitative",
  "quantitative_method": "...",
  "narrative_method": "...",
  "heterogeneity_assessment": "...",
  "i2_threshold": "50%",
  "pooling_decision_rule": "...",
  "software": "...",
  "chat_reply": "..."
}
- synthesis_type: 'quantitative' | 'narrative' | 'both' — based on expected study homogeneity
- quantitative_method: random-effects meta-analysis (DerSimonian-Laird or REML), fixed-effects
- narrative_method: structured narrative synthesis using harvest plots, vote counting, logic models
- heterogeneity_assessment: Cochran's Q test + I² statistic + prediction interval
- pooling_decision_rule: criteria for deciding whether to pool (clinical homogeneity + I²<threshold)
- software: R (meta package), RevMan, Stata, or OpenMeta-Analyst""",

    "effect_measures": """You are a systematic review statistician recommending effect measures for this review.
Return JSON: {
  "primary_effect_measure": "...",
  "secondary_effect_measures": ["..."],
  "rationale": "...",
  "confidence_intervals": "95% CI",
  "significance_threshold": "p < 0.05",
  "heterogeneity_measure": "I²",
  "chat_reply": "..."
}
Choose based on review_family and outcome types:
- Continuous outcomes → Mean Difference (MD) if same scale; Standardised Mean Difference (SMD/Hedges' g) if different scales; Cohen's d for psychological outcomes
- Binary/dichotomous → Risk Ratio (RR) preferred; Odds Ratio (OR); Risk Difference (RD); Number Needed to Treat (NNT)
- Time-to-event/survival → Hazard Ratio (HR) with log-scale pooling
- Ordinal → Common Odds Ratio using proportional OR model
- Diagnostic → Sensitivity, Specificity, Positive/Negative Likelihood Ratio, Diagnostic Odds Ratio (DOR)
- Prevalence → Pooled proportion with Freeman-Tukey double arcsine transformation for stabilisation""",

    "subgroup_sensitivity": """You are a systematic review methodologist writing subgroup and sensitivity analyses.
Return JSON: {
  "subgroup_analyses": ["..."],
  "subgroup_rationale": "...",
  "sensitivity_analyses": ["..."],
  "missing_data_plan": "...",
  "heterogeneity_threshold": "I² > 50%",
  "chat_reply": "..."
}
- subgroup_analyses: 3–5 pre-specified moderator variables (e.g. age group, severity, intervention type, setting, study design, follow-up duration, risk of bias category)
- subgroup_rationale: why these specific subgroups are clinically or methodologically important
- sensitivity_analyses: 3–4 planned checks — e.g. exclude high-RoB studies, leave-one-out analysis, restrict to peer-reviewed RCTs, restrict to studies with low attrition, fixed vs random effects comparison
- missing_data_plan: best-case/worst-case scenarios for binary outcomes; multiple imputation for continuous; contact authors
IMPORTANT: All subgroup and sensitivity analyses must be pre-specified to avoid post-hoc data-dredging.""",

    "reporting_certainty": """You are a systematic review methodologist writing the reporting bias and certainty of evidence assessment plan.
Return JSON: {
  "reporting_bias_methods": ["..."],
  "funnel_plot_threshold": "≥10 studies per meta-analysis",
  "statistical_tests": ["..."],
  "certainty_tool": "GRADE",
  "certainty_software": "GRADEpro GDT",
  "certainty_domains": ["..."],
  "downgrade_criteria": ["..."],
  "upgrade_criteria": ["..."],
  "sof_table": true,
  "chat_reply": "..."
}
Reporting bias methods: funnel plot asymmetry (visual), Egger's linear regression test, Begg's rank correlation test, trim-and-fill method, p-curve analysis
Note: funnel plots are only interpretable when ≥10 studies per meta-analysis
GRADE certainty domains: risk of bias, inconsistency (heterogeneity), indirectness, imprecision, publication bias
Downgrade for: serious/very serious limitations in any domain
Upgrade for: large effect (RR>2 or <0.5), dose-response, plausible confounding opposing the effect
For qualitative reviews: use CERQual (Confidence in the Evidence from Reviews of Qualitative research) instead of GRADE
For scoping reviews: certainty assessment is not applicable — omit this section""",

    "admin": """You are a systematic review methodologist writing the administrative section of a PRISMA-P 2015 protocol.
Return JSON: {
  "registry_recommendation": "PROSPERO",
  "registry_rationale": "...",
  "registration_timing": "Prior to beginning database searches",
  "amendments_policy": "...",
  "protocol_deposit": "...",
  "funding_note": "...",
  "chat_reply": "..."
}
Registry selection:
- Health/clinical/biomedical topics → PROSPERO (York)
- Social science, education, criminal justice → Campbell Collaboration Open Library
- Multidisciplinary or open science → OSF Registries (osf.io)
- Not eligible for other registries → INPLASY
registration_timing: protocol must be registered BEFORE starting database searches (PRISMA-P requirement)
amendments_policy: any post-registration protocol deviations must be documented with date, reason, and impact assessment
protocol_deposit: recommend depositing full protocol on OSF or Zenodo (open access) in addition to registry
Note: The admin form also captures author names, affiliations, ORCIDs, CRediT contributions, funding sources, competing interests — these are entered directly in the form fields, not generated by AI.""",
}


async def generate_phase_content(
    phase: str,
    pico_context: dict,
    context_data: dict,
    current_content: dict | None,
    messages: list[dict],
    ai_provider: AIProvider,
    review_type: str = "systematic_review",
    mode: str = "direct",
) -> dict:
    """
    Generate or update content for a specific protocol phase, incorporating
    the full chat history for iterative refinement.

    phase: one of the keys in _PHASE_SYSTEMS
    pico_context: the finalized PICO/framework elements from Phase 4
    context_data: all previously finalized phase data
    messages: list of { role: "ai"|"user", text: "..." }
    Returns: { reply: str, content: dict }
    mode: "direct" (default) — apply changes; "plan" — return options only, no content update
    """
    system_prompt = _PHASE_SYSTEMS.get(phase)
    if not system_prompt:
        return {"reply": f"Unknown phase: {phase}", "content": {}}

    # Plan mode: override system prompt to return options/questions only
    if mode == "plan":
        system_prompt = (
            system_prompt
            + "\n\nIMPORTANT: The user is in PLAN MODE. Do NOT generate new content yet. "
            "Instead return JSON: { \"chat_reply\": \"...\" } where chat_reply presents "
            "2–4 concrete options (labeled A, B, C...) describing different ways to fulfill the user's request, "
            "or asks 1–2 clarifying questions. Explain what each option would change and its approximate scope."
        )

    # Build framework context string dynamically from whatever elements are present
    _SKIP_SUMMARY_KEYS = {"review_type", "framework", "target_registries", "language_restriction",
                          "date_from", "date_to"}
    framework = context_data.get("framework", pico_context.get("framework", "PICO"))
    review_family = context_data.get("review_family", "")
    review_question = context_data.get("review_question", "")
    background = context_data.get("background", "")
    protocol_document = context_data.get("protocol_document", "")

    element_lines = []
    for key, val in pico_context.items():
        if key in _SKIP_SUMMARY_KEYS or not val:
            continue
        label = key.replace("_", " ").title()
        element_lines.append(f"{label}: {val}")

    pico_summary = f"""Framework: {framework}
Review type: {review_type}
Review family: {review_family}
Research question: {review_question}
{chr(10).join(element_lines) if element_lines else "(No framework elements provided)"}"""

    if background:
        pico_summary += f"\n\nBackground context:\n{background[:800]}"

    current_draft = context_data.get("current_draft", "")
    if current_draft:
        pico_summary += f"\n\nCurrent draft to revise:\n{current_draft[:2000]}"

    if current_content:
        try:
            current_content_json = json.dumps(current_content, indent=2, ensure_ascii=False)
        except Exception:
            current_content_json = str(current_content)
        pico_summary += f"\n\nCurrent section content:\n{current_content_json[:2500]}"

    evidence_pack = context_data.get("evidence_pack")
    evidence_context = ""
    if phase in {"background", "rationale"} and isinstance(evidence_pack, dict):
        summaries = evidence_pack.get("summaries", []) or []
        papers = evidence_pack.get("ranked_papers", []) or []
        evidence_context = (
            _format_summary_sources_with_src_markers(summaries)
            if summaries else _format_papers_with_src_markers(papers)
        )

    if protocol_document:
        pico_summary += f"\n\nFull protocol document:\n{protocol_document[:5000]}"

    # Build conversation history for context
    conversation = "\n".join(
        f"{'AI' if m['role'] == 'ai' else 'USER'}: {m['text']}"
        for m in messages[-6:]  # last 6 messages for context window efficiency
    )

    is_first_call = not any(m["role"] == "user" for m in messages)

    if is_first_call:
        user_prompt = f"""PICO Context:
{pico_summary}

This is the first draft for the '{phase}' phase. Generate appropriate content."""
    else:
        user_prompt = f"""PICO Context:
{pico_summary}

Conversation history:
{conversation}

Update the content based on the latest user feedback.
Return the updated structured content for this phase, preserving useful existing fields unless the user explicitly asks to remove or replace them."""

    if evidence_context:
        user_prompt += (
            "\n\nEvidence pack for citation-grounded revision "
            "(use ONLY these [SRC{n}] markers if you cite evidence):\n"
            f"{evidence_context[:14000]}"
        )

    # For text-only phases (objectives), return plain text wrapped in dict
    if phase == "objectives":
        try:
            draft = await ai_provider.complete(
                system=system_prompt,
                user=user_prompt,
                temperature=0.3,
            )
            reply = "Here are the formal objectives for your review. They translate your research question into precise aims aligned with PRISMA-P item 7. Edit them directly or let me know if you want a different framing."
            if not is_first_call:
                reply = "I've updated the objectives based on your feedback."
            return {"reply": reply, "content": {"objectives": draft.strip()}}
        except Exception as e:
            return {"reply": f"Error: {e}", "content": {}}

    if phase in {"background", "rationale"} and mode == "direct":
        tagged_system = (
            system_prompt
            + "\nIf an evidence pack is provided, cite ONLY with the supplied [SRC{n}] markers and do not invent any citation."
            + "\n\nReturn plain text in EXACTLY this format:\n"
              "CHAT_REPLY:\n"
              "<1-2 sentence summary of what you revised>\n\n"
              "TEXT:\n"
              "<<<\n"
              "<the full revised section text>\n"
              ">>>\n"
              "Do NOT return JSON. Do NOT use code fences."
        )
        try:
            resp = await ai_provider.guarded_complete(
                system=tagged_system,
                user=user_prompt,
                config=CompletionConfig(
                    output_format=OutputFormat.PROSE,
                    budget=TokenBudget(target_words=1800),
                    max_continuations=2,
                ),
                temperature=0.3,
            )
            reply, text = _parse_tagged_text_phase_response(resp.text)
            if not text.strip():
                raise ValueError("Model returned no revised text for this section.")
            if not reply:
                reply = "I've updated this section based on your feedback."
            if isinstance(evidence_pack, dict) and _evidence_sources_from_pack(evidence_pack):
                updated_pack, _new_cited_ids, warnings = _refresh_pack_references(evidence_pack, phase, text)
                if warnings:
                    reply += f" {len(warnings)} citation marker(s) still need manual review."
                return {
                    "reply": reply,
                    "content": {
                        "text": updated_pack.get(f"{phase}_draft", text.strip()),
                        "references_md": updated_pack.get("references_md", ""),
                        "references_json": updated_pack.get("references_json", []),
                        "bibtex": updated_pack.get("bibtex", ""),
                        "cited_ids": updated_pack.get("cited_ids", []),
                        "citation_warnings": warnings,
                        "evidence_pack": updated_pack,
                    },
                }
            return {"reply": reply, "content": {"text": text.strip()}}
        except Exception as e:
            logger.error("generate_phase_content(%s) tagged-text path failed: %s", phase, e)
            return {"reply": f"Error generating content: {e}", "content": {}}

    # For structured JSON phases
    try:
        resp = await ai_provider.guarded_complete(
            system=system_prompt, user=user_prompt,
            config=CompletionConfig(
                output_format=OutputFormat.JSON,
                budget=TokenBudget(target_json_keys=30, format=OutputFormat.JSON),
                max_continuations=2,
            ),
            json_mode=True, temperature=0.25,
        )
        result = _load_llm_json(resp.text)
        reply = result.pop("chat_reply", "Content generated. Review it above and let me know if you want any changes.")
        if not reply:
            reply = "I've updated this section based on your feedback."
        return {"reply": reply, "content": result}
    except Exception as e:
        logger.error("generate_phase_content(%s) failed: %s", phase, e)
        return {"reply": f"Error generating content: {e}", "content": {}}


# ── Evidence Pack — citation-validated background/rationale generation ─────────

_BACKGROUND_CITATION_SYSTEM = """You are a systematic review methodologist writing the 'Background' section for a PRISMA-P 2015 compliant protocol.

You have been provided a numbered set of evidence-backed source summaries derived from the literature search. These are the ONLY sources you may cite.

CITATION RULES (strictly enforced):
- Cite papers using ONLY the [SRC{n}] marker format (e.g. [SRC1], [SRC3])
- Every factual claim must be supported by at least one [SRC{n}] citation
- Do NOT invent author names, DOIs, or statistics not present in the provided paper list
- Do NOT cite papers not in the provided list

STRUCTURE — write exactly four ## subheadings in this order:

## The Problem, Condition, or Issue
3–4 paragraphs: definition, global prevalence/burden (cite data), affected populations, consequences, risk factors.

## The Intervention (or Exposure / Phenomenon of Interest)
3–4 paragraphs: description of intervention/exposure, how it is delivered/implemented, variations and subtypes, existing guidelines or recommendations (cite them).

## How the Intervention Could Work
2–3 paragraphs: theoretical mechanism of action, biological/psychological/social pathways, evidence for plausible effects [SRC{n}], moderating and mediating factors.

## Why It Is Important to Do This Review
2–3 paragraphs: prior systematic reviews on this topic and their limitations (cite them with [SRC{n}]), the specific evidence gap, stakeholders who will benefit, policy/clinical implications.

WRITING STANDARDS:
- Total length: 1000–1500 words across all four sections
- Formal academic prose, third person, no bullet lists in body
- Each [SRC{n}] marker must correspond to a paper in the provided list
- Return plain Markdown with the four ## subheadings. No JSON, no code fences."""

_RATIONALE_CITATION_SYSTEM = """You are a systematic review methodologist writing the 'Rationale & Gap' section for a PRISMA-P 2015 compliant protocol.

You have been provided a numbered set of evidence-backed source summaries derived from full-text or abstract analysis. Cite ONLY these sources using [SRC{n}] markers.

Write 3–4 focused paragraphs (400–600 words total):
1. Summarise existing systematic reviews or evidence syntheses on this topic [SRC{n}] and their limitations (recency, scope, methodological gaps)
2. Identify the specific evidence gap this review will address
3. Explain why this review is needed now and who will benefit
4. State what this review will add beyond what already exists

CITATION RULES: cite ONLY using [SRC{n}] markers from the provided paper list. No fabricated citations.
Return plain Markdown paragraphs (no headings, no JSON)."""


def _build_background_queries(query: str, pico_context: dict | None) -> list[str]:
    """Build 3 query variants for richer coverage: core + outcome-focused + prior-reviews."""
    queries = [query]
    if pico_context:
        pop = pico_context.get("population", "")
        intv = pico_context.get("intervention", "") or pico_context.get("exposure", "")
        out = pico_context.get("outcome", "")
        if pop and intv:
            queries.append(f"{pop} {intv}")
        if pop and out:
            queries.append(f"{pop} {out}")
    # Always add a prior-reviews query
    queries.append(f"systematic review {query}")
    return list(dict.fromkeys(q.strip() for q in queries if q.strip()))[:4]


def _normalize_doi(doi: str) -> str:
    """Lowercase, strip URL prefix."""
    doi = doi.lower().strip()
    for prefix in ("https://doi.org/", "http://doi.org/", "doi:"):
        if doi.startswith(prefix):
            doi = doi[len(prefix):]
    return doi


def _deduplicate_papers(papers: list) -> list:
    """Deduplicate by normalized DOI > PMID > title+year+first_author."""
    seen: set[str] = set()
    unique = []
    for p in papers:
        # Try DOI
        doi = getattr(p, "doi", None) or (p.get("doi") if isinstance(p, dict) else None)
        if doi:
            key = _normalize_doi(doi)
            if key in seen:
                continue
            seen.add(key)
            unique.append(p)
            continue
        # Try PMID
        pmid = getattr(p, "pmid", None) or (p.get("pmid") if isinstance(p, dict) else None)
        if pmid:
            key = f"pmid:{pmid}"
            if key in seen:
                continue
            seen.add(key)
            unique.append(p)
            continue
        # Fallback: title slice + year + first author
        title = getattr(p, "title", None) or (p.get("title") if isinstance(p, dict) else None) or ""
        year = getattr(p, "year", None) or (p.get("year") if isinstance(p, dict) else None) or ""
        authors = getattr(p, "authors", None) or (p.get("authors") if isinstance(p, dict) else None) or []
        first_author = (authors[0] if isinstance(authors, list) and authors else str(authors))[:10]
        key = f"{title[:60].lower()}:{year}:{first_author.lower()}"
        if key in seen:
            continue
        seen.add(key)
        unique.append(p)
    return unique


def _paper_to_dict(p) -> dict:
    """Convert a Paper object or dict to a serializable dict with a stable id."""
    if isinstance(p, dict):
        d = dict(p)
    else:
        d = {
            "title": getattr(p, "title", ""),
            "authors": getattr(p, "authors", []),
            "year": getattr(p, "year", None),
            "doi": getattr(p, "doi", None),
            "pmid": getattr(p, "pmid", None),
            "journal": getattr(p, "journal", None),
            "abstract": getattr(p, "abstract", ""),
            "citation_count": getattr(p, "citation_count", 0),
        }
    if "id" not in d or not d["id"]:
        d["id"] = str(uuid.uuid4())
    return _json_safe(d)


def _format_authors_short(authors) -> str:
    """Return 'Last et al.' or 'Last & Last' or 'Last' from author list or string."""
    if isinstance(authors, list):
        names = [str(a).split(",")[0].strip() for a in authors if a]
    else:
        names = [str(authors).split(",")[0].strip()]
    if len(names) == 0:
        return "Unknown"
    if len(names) == 1:
        return names[0]
    if len(names) == 2:
        return f"{names[0]} & {names[1]}"
    return f"{names[0]} et al."


def _format_papers_with_src_markers(papers: list[dict]) -> str:
    """Format paper list with [SRC{n}] markers for AI citation injection."""
    lines = []
    for i, p in enumerate(papers, 1):
        authors = _format_authors_short(p.get("authors", ""))
        year = p.get("year") or "n.d."
        title = p.get("title", "Untitled")
        journal = p.get("journal") or ""
        doi = p.get("doi") or ""
        line = f"[SRC{i}] {authors} ({year}). {title}."
        if journal:
            line += f" {journal}."
        if doi:
            line += f" https://doi.org/{doi}"
        lines.append(line)
    return "\n".join(lines)


def _resolve_citation_markers(text: str, papers: list[dict]) -> tuple[str, list[str], list[str]]:
    """
    Replace [SRC{n}] markers with (Author et al., Year) in-text citations.
    Returns (resolved_text, cited_ids, warnings).
    """
    cited_ids: list[str] = []
    warnings: list[str] = []
    max_n = len(papers)

    def _citation_for_index(n: int) -> str:
        if 1 <= n <= max_n:
            p = papers[n - 1]
            author = _format_authors_short(p.get("authors", ""))
            year = p.get("year") or "n.d."
            pid = p.get("id", str(n))
            if pid not in cited_ids:
                cited_ids.append(pid)
            return f"({author}, {year})"
        warnings.append(f"[SRC{n}] out of range (only {max_n} papers)")
        return f"[SRC{n}]"

    def multi_replacer(m: re.Match) -> str:
        numbers = [int(num) for num in re.findall(r"(?:SRC)?(\d+)", m.group(1), flags=re.IGNORECASE)]
        if not numbers:
            return m.group(0)
        citations = []
        for n in numbers:
            resolved_citation = _citation_for_index(n)
            if resolved_citation.startswith("(") and resolved_citation.endswith(")"):
                citations.append(resolved_citation[1:-1])
            else:
                citations.append(resolved_citation)
        if len(citations) == 1:
            return f"({citations[0]})"
        return f"({'; '.join(citations)})"

    resolved = re.sub(
        r"\[((?:SRC?\d+\s*[,;]\s*)+SRC?\d+)\]",
        multi_replacer,
        text,
        flags=re.IGNORECASE,
    )
    resolved = re.sub(r"\[SRC(\d+)\]", lambda m: _citation_for_index(int(m.group(1))), resolved, flags=re.IGNORECASE)
    return resolved, cited_ids, warnings


def _format_apa_references(papers: list[dict]) -> str:
    """Format APA 7th edition reference list from cited papers."""
    if not papers:
        return ""
    lines = ["## References\n"]
    for p in papers:
        authors_raw = p.get("authors", [])
        if isinstance(authors_raw, list):
            formatted_authors = ", ".join(
                f"{str(a).split(',')[0].strip()}, {'. '.join(n[0] for n in str(a).split(',')[1:]).strip()}" if "," in str(a) else str(a)
                for a in authors_raw
            )
        else:
            formatted_authors = str(authors_raw)
        year = p.get("year") or "n.d."
        title = p.get("title", "Untitled")
        journal = p.get("journal") or ""
        doi = p.get("doi") or ""
        ref = f"{formatted_authors} ({year}). {title}."
        if journal:
            ref += f" *{journal}*."
        if doi:
            ref += f" https://doi.org/{doi}"
        lines.append(ref + "\n")
    return "\n".join(lines)


def _format_csl_json(papers: list[dict]) -> list[dict]:
    """Convert papers to basic CSL-JSON for style switching."""
    items = []
    for p in papers:
        authors_raw = p.get("authors", [])
        author_list = []
        if isinstance(authors_raw, list):
            for a in authors_raw:
                parts = str(a).split(",", 1)
                author_list.append({"family": parts[0].strip(), "given": parts[1].strip() if len(parts) > 1 else ""})
        elif authors_raw:
            author_list.append({"literal": str(authors_raw)})
        year = p.get("year")
        issued = {"date-parts": [[int(year)]]} if year else {}
        item: dict = {
            "id": p.get("id", str(uuid.uuid4())),
            "type": "article-journal",
            "title": p.get("title", ""),
            "author": author_list,
            "issued": issued,
            "container-title": p.get("journal") or "",
            "DOI": p.get("doi") or "",
        }
        items.append(item)
    return items


def _format_bibtex(papers: list[dict]) -> str:
    """Generate BibTeX entries for cited papers."""
    entries = []
    for p in papers:
        pid = p.get("id", str(uuid.uuid4()))[:8]
        authors_raw = p.get("authors", [])
        if isinstance(authors_raw, list):
            author_str = " and ".join(str(a) for a in authors_raw)
        else:
            author_str = str(authors_raw)
        year = p.get("year") or ""
        title = p.get("title", "").replace("{", "").replace("}", "")
        journal = (p.get("journal") or "").replace("{", "").replace("}", "")
        doi = p.get("doi") or ""
        entry = (
            f"@article{{{pid},\n"
            f"  author = {{{author_str}}},\n"
            f"  year = {{{year}}},\n"
            f"  title = {{{{{title}}}}},\n"
            f"  journal = {{{journal}}},\n"
            + (f"  doi = {{{doi}}},\n" if doi else "")
            + "}"
        )
        entries.append(entry)
    return "\n\n".join(entries)


def _evidence_sources_from_pack(pack: dict) -> list[dict]:
    summaries = pack.get("summaries", []) or []
    if summaries:
        return [_summary_to_source_dict(summary) for summary in summaries]
    return list(pack.get("ranked_papers", []) or [])


def _ordered_cited_papers(sources: list[dict], cited_ids: list[str]) -> list[dict]:
    source_by_id = {
        str(source.get("id")): source
        for source in sources
        if source.get("id") is not None
    }
    ordered: list[dict] = []
    seen: set[str] = set()
    for cited_id in cited_ids:
        sid = str(cited_id)
        if sid in seen:
            continue
        paper = source_by_id.get(sid)
        if paper:
            ordered.append(paper)
            seen.add(sid)
    return ordered


def _refresh_pack_references(pack: dict, section: str, draft_raw: str) -> tuple[dict, list[str], list[str]]:
    sources = _evidence_sources_from_pack(pack)
    resolved_text, new_cited_ids, warnings = _resolve_citation_markers(draft_raw, sources)
    has_src_markers = bool(re.search(r"\[(?:SRC\d+)(?:\s*[,;]\s*SRC\d+)*\]", draft_raw, flags=re.IGNORECASE))
    if not new_cited_ids and not has_src_markers:
        preserved_pack = dict(pack)
        preserved_pack[f"{section}_draft"] = draft_raw.strip()
        return preserved_pack, [], ["No [SRCn] markers were returned, so the existing references were preserved."]
    all_cited_ids = (
        new_cited_ids
        if section == "background"
        else list(dict.fromkeys([*(pack.get("cited_ids", []) or []), *new_cited_ids]))
    )
    cited_papers = _ordered_cited_papers(sources, all_cited_ids)

    updated_pack = dict(pack)
    updated_pack.update({
        "cited_ids": all_cited_ids,
        f"{section}_draft": resolved_text.strip(),
        "references_md": _format_apa_references(cited_papers),
        "references_json": _format_csl_json(cited_papers),
        "bibtex": _format_bibtex(cited_papers),
    })
    return updated_pack, new_cited_ids, warnings


def _compact_evidence_summary(summary: dict, paper: dict) -> dict:
    sentence_bank = summary.get("sentence_bank", []) or []
    ranked_bank = sorted(
        sentence_bank,
        key=lambda sent: (
            sent.get("importance") != "high",
            sent.get("use_in") not in {"introduction", "discussion", "results"},
            sent.get("section") not in {"background", "discussion", "results"},
        ),
    )

    return _json_safe({
        "source_id": paper.get("id") or summary.get("paper_key") or str(uuid.uuid4()),
        "paper_key": summary.get("paper_key") or paper.get("doi") or (paper.get("title") or "")[:60].lower().strip(),
        "bibliography": summary.get("bibliography") or {
            "title": paper.get("title", ""),
            "authors": paper.get("authors", []),
            "year": paper.get("year"),
            "journal": paper.get("journal", ""),
            "doi": paper.get("doi", ""),
        },
        "methods": summary.get("methods") or {},
        "results": (summary.get("results") or [])[:2],
        "critical_appraisal": summary.get("critical_appraisal") or {},
        "one_line_takeaway": summary.get("one_line_takeaway", ""),
        "sentence_bank": ranked_bank[:6],
        "full_text_used": bool(summary.get("full_text_used")),
        "text_source": summary.get("text_source", ""),
    })


def _summary_to_source_dict(summary: dict) -> dict:
    bibliography = summary.get("bibliography") or {}
    return {
        "id": summary.get("source_id") or summary.get("paper_key") or str(uuid.uuid4()),
        "paper_key": summary.get("paper_key") or "",
        "authors": bibliography.get("authors") or [],
        "year": bibliography.get("year"),
        "title": bibliography.get("title") or "",
        "journal": bibliography.get("journal") or "",
        "doi": bibliography.get("doi") or "",
        "one_line_takeaway": summary.get("one_line_takeaway") or "",
        "methods": summary.get("methods") or {},
        "results": summary.get("results") or [],
        "critical_appraisal": summary.get("critical_appraisal") or {},
        "sentence_bank": summary.get("sentence_bank") or [],
        "full_text_used": bool(summary.get("full_text_used")),
        "text_source": summary.get("text_source") or "",
    }


def _format_summary_sources_with_src_markers(summaries: list[dict]) -> str:
    lines: list[str] = []
    for i, summary in enumerate(summaries, 1):
        src = _summary_to_source_dict(summary)
        authors = _format_authors_short(src.get("authors", ""))
        year = src.get("year") or "n.d."
        title = src.get("title") or "Untitled"
        journal = src.get("journal") or ""
        doi = src.get("doi") or ""

        header = f"[SRC{i}] {authors} ({year}). {title}."
        if journal:
            header += f" {journal}."
        if doi:
            header += f" https://doi.org/{doi}"
        lines.append(header)

        takeaway = str(src.get("one_line_takeaway") or "").strip()
        if takeaway and takeaway != "NR":
            lines.append(f"  Takeaway: {takeaway}")

        methods = src.get("methods") or {}
        study_design = str(methods.get("study_design") or "").strip()
        sample_n = str(methods.get("sample_n") or "").strip()
        if study_design or sample_n:
            study_bits = [bit for bit in [study_design, f"N={sample_n}" if sample_n else ""] if bit]
            lines.append(f"  Study: {' | '.join(study_bits)}")

        evidence = str((src.get("critical_appraisal") or {}).get("evidence_grade") or "").strip()
        if evidence and evidence != "NR":
            lines.append(f"  Evidence: {evidence}")

        sentence_bank = src.get("sentence_bank") or []
        if sentence_bank:
            for sent in sentence_bank[:6]:
                text = str(sent.get("text") or "").strip()
                if not text:
                    continue
                stats = str(sent.get("stats") or "").strip()
                importance = "★" if sent.get("importance") == "high" else "-"
                use_in = str(sent.get("use_in") or "discussion").upper()
                line = f"  {importance} [{use_in}] {text}"
                if stats and stats != "NR":
                    line += f" ({stats})"
                lines.append(line)
        else:
            first_result = (src.get("results") or [{}])[0] or {}
            finding = str(first_result.get("finding") or "").strip()
            if finding and finding != "NR":
                lines.append(f"  Finding: {finding}")

        lines.append("")

    return "\n".join(lines).strip()


async def build_evidence_pack(
    query: str,
    ai_provider: AIProvider,
    n_articles: int = 20,
    pico_context: dict | None = None,
    fetch_settings=None,
) -> dict:
    """
    Run a scoping literature search (NOT the formal review search) to build
    an Evidence Pack for writing the Background + Rationale sections.

    Returns an EvidencePack dict with ranked_papers and metadata.
    Text generation (background/rationale) is handled separately by
    write_background_from_pack() and write_rationale_from_pack().
    """
    from models import Paper
    from services.literature_engine import LiteratureEngine
    from services.paper_summarizer import summarize_paper
    from services.query_expander import expand_query, heuristic_expand_query

    engine = LiteratureEngine()
    base_queries = _build_background_queries(query, pico_context)
    try:
        expanded = await expand_query(ai_provider, query, article_type="systematic_review")
        queries = list(dict.fromkeys([*(expanded.queries or []), *base_queries]))[:8]
        pubmed_queries = expanded.pubmed_queries or None
    except Exception as exc:
        logger.warning("build_evidence_pack query expansion failed: %s", exc)
        fallback = heuristic_expand_query(query, article_type="systematic_review")
        queries = list(dict.fromkeys([*(fallback.queries or []), *base_queries]))[:8]
        pubmed_queries = fallback.pubmed_queries or None

    logger.info("build_evidence_pack: queries=%s n_articles=%d", queries, n_articles)

    raw_papers: list = []
    final_papers: list[dict] = []
    try:
        async for event in engine.search_all_streaming(
            queries,
            total_limit=n_articles,
            pubmed_queries=pubmed_queries,
        ):
            etype = event.get("type", "")
            if etype in {"papers", "result"}:
                raw_papers.extend(event.get("papers", []))
            elif etype == "complete":
                final_papers = event.get("papers", []) or []
    except Exception as exc:
        logger.warning("build_evidence_pack search error: %s", exc)

    unique = final_papers or _deduplicate_papers(raw_papers)[:n_articles]
    paper_dicts = [_paper_to_dict(p) for p in unique]

    summary_limit = min(len(paper_dicts), 20)
    summaries: list[dict] = []
    summary_warnings: list[str] = []

    async def _summarize_one(paper_dict: dict) -> tuple[dict | None, str | None]:
        try:
            paper = Paper(
                title=str(paper_dict.get("title") or ""),
                authors=list(paper_dict.get("authors") or []),
                abstract=paper_dict.get("abstract"),
                doi=paper_dict.get("doi"),
                pmid=paper_dict.get("pmid"),
                pmcid=paper_dict.get("pmcid"),
                year=paper_dict.get("year"),
                journal=paper_dict.get("journal"),
                citation_count=paper_dict.get("citation_count"),
                oa_pdf_url=paper_dict.get("oa_pdf_url"),
                source=str(paper_dict.get("source") or "openalex"),
            )
            summary = await summarize_paper(
                ai_provider,
                paper,
                query,
                fetch_settings=fetch_settings,
                session_id="",
            )
            return _compact_evidence_summary(summary.model_dump(), paper_dict), None
        except Exception as exc:
            title = str(paper_dict.get("title") or "Untitled")[:80]
            logger.warning("Evidence-pack summarization failed for %s: %s", title, exc)
            return None, f"Could not summarize '{title}': {exc}"

    if summary_limit:
        tasks = [asyncio.create_task(_summarize_one(paper_dict)) for paper_dict in paper_dicts[:summary_limit]]
        for task in asyncio.as_completed(tasks):
            compact_summary, warning = await task
            if compact_summary:
                summaries.append(compact_summary)
            if warning:
                summary_warnings.append(warning)

    return {
        "search_terms": queries,
        "databases_searched": ["PubMed", "OpenAlex", "Semantic Scholar", "Europe PMC", "Crossref"],
        "search_date": datetime.utcnow().date().isoformat(),
        "retrieved_count": len(raw_papers),
        "deduplicated_count": len(paper_dicts),
        "ranked_papers": paper_dicts,
        "summaries": summaries,
        "summary_count": len(summaries),
        "cited_ids": [],
        "background_draft": "",
        "rationale_draft": "",
        "references_md": "",
        "references_json": [],
        "bibtex": "",
        "warnings": summary_warnings,
    }


async def write_background_from_pack(
    pack: dict,
    query: str,
    ai_provider: AIProvider,
    review_type: str = "systematic_review",
) -> dict:
    """
    Write Background section using the Evidence Pack with [SRC{n}] citation
    validation. Returns updated pack dict + list of any citation warnings.
    """
    papers = pack.get("ranked_papers", [])
    summaries = pack.get("summaries", []) or []
    sources = _evidence_sources_from_pack(pack)
    paper_context = _format_summary_sources_with_src_markers(summaries) if summaries else _format_papers_with_src_markers(papers)

    user_msg = (
        f"Research topic: {query}\n"
        f"Review type: {review_type}\n\n"
        f"Evidence sources (cite using [SRC{{n}}] markers only):\n{paper_context}\n\n"
        "Write the full structured Background section (1000–1500 words) with four ## subheadings."
    )

    try:
        resp = await ai_provider.guarded_complete(
            system=_BACKGROUND_CITATION_SYSTEM, user=user_msg,
            config=CompletionConfig(
                output_format=OutputFormat.PROSE,
                budget=TokenBudget(target_words=1500),
                max_continuations=2,
            ),
            temperature=0.3,
        )
        draft_raw = resp.text
    except Exception as exc:
        logger.error("write_background_from_pack AI call failed: %s", exc)
        return {"pack": pack, "warnings": [str(exc)]}

    updated_pack, cited_ids, warnings = _refresh_pack_references(pack, "background", draft_raw)

    # If >20% citations unresolved, re-prompt once
    if warnings and len(warnings) > max(1, len(cited_ids)) * 0.2:
        logger.warning("write_background_from_pack: %d citation warnings, re-prompting", len(warnings))
        try:
            resp = await ai_provider.guarded_complete(
                system=_BACKGROUND_CITATION_SYSTEM,
                user=user_msg + "\n\nIMPORTANT: Use ONLY the [SRC{n}] markers listed above. Do not invent any citation.",
                config=CompletionConfig(
                    output_format=OutputFormat.PROSE,
                    budget=TokenBudget(target_words=1500),
                    max_continuations=2,
                ),
                temperature=0.2,
            )
            draft_raw = resp.text
            updated_pack, cited_ids, warnings = _refresh_pack_references(pack, "background", draft_raw)
        except Exception:
            pass

    pack = updated_pack
    sources_used = len(_ordered_cited_papers(sources, pack.get("cited_ids", []) or []))
    summary = (
        f"I searched the literature and found **{pack['deduplicated_count']} papers** on '{query[:60]}'. "
        f"**{sources_used} sources** are cited in this background. "
        "It covers four subheadings: The Problem, The Intervention, How It Works, and Why This Review Is Needed. "
        "Review the text and tell me what to expand, change, or add."
    )
    return {"pack": pack, "warnings": [*(pack.get("warnings", []) or []), *warnings], "summary": summary}


async def write_rationale_from_pack(
    pack: dict,
    query: str,
    ai_provider: AIProvider,
    review_type: str = "systematic_review",
) -> dict:
    """
    Write the Rationale & Gap section by reusing the existing Evidence Pack.
    Optionally supplements with a prior-reviews search.
    Returns updated pack + summary.
    """
    papers = pack.get("ranked_papers", [])
    summaries = pack.get("summaries", []) or []
    sources = _evidence_sources_from_pack(pack)
    paper_context = _format_summary_sources_with_src_markers(summaries) if summaries else _format_papers_with_src_markers(papers)

    user_msg = (
        f"Research topic: {query}\n"
        f"Review type: {review_type}\n\n"
        f"Background already written:\n{pack.get('background_draft', '')[:400]}\n\n"
        f"Evidence sources (cite using [SRC{{n}}] markers only):\n{paper_context}\n\n"
        "Write the Rationale & Gap section (400–600 words, 3–4 paragraphs)."
    )

    try:
        resp = await ai_provider.guarded_complete(
            system=_RATIONALE_CITATION_SYSTEM, user=user_msg,
            config=CompletionConfig(
                output_format=OutputFormat.PROSE,
                budget=TokenBudget(target_words=600),
                max_continuations=1,
            ),
            temperature=0.3,
        )
        draft_raw = resp.text
    except Exception as exc:
        logger.error("write_rationale_from_pack failed: %s", exc)
        return {"pack": pack, "warnings": [str(exc)], "summary": "Error generating rationale."}

    pack, new_cited_ids, warnings = _refresh_pack_references(pack, "rationale", draft_raw)
    summary = (
        f"Rationale drafted using the same evidence pack ({len(sources)} sources). "
        f"{len(new_cited_ids)} sources cited in this section. "
        "Tell me what to revise or expand."
    )
    return {"pack": pack, "warnings": [*(pack.get("warnings", []) or []), *warnings], "summary": summary}
