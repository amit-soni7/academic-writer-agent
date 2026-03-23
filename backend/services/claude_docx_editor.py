"""
services/claude_docx_editor.py

AI-powered track changes .docx generation.

Flow:
  1. Original manuscript + finalized change plans → AI prompt
  2. AI returns python-docx code that modifies a Document in-place
  3. Code is executed in a restricted sandbox
  4. Modified .docx bytes are returned
"""

from __future__ import annotations

import io
import logging
import re
from datetime import datetime, timezone
from typing import TYPE_CHECKING

import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

if TYPE_CHECKING:
    from services.ai_provider import AIProvider

logger = logging.getLogger(__name__)

# ── System prompt ─────────────────────────────────────────────────────────────

_SYSTEM = """\
You are a python-docx expert. Generate Python code that applies reviewer-requested
changes to a Word document using OOXML track changes (w:ins / w:del).

AVAILABLE VARIABLES (already defined — do NOT import anything):
  doc       — a python-docx Document pre-loaded with the original manuscript
  author    — string, the track-changes author name
  date      — string, ISO 8601 date for change timestamps
  qn        — docx.oxml.ns.qn  (namespace helper)
  OxmlElement — docx.oxml.OxmlElement
  re        — the re module
  docx      — the docx module

HELPER FUNCTIONS — define these at the top of your code, then use them:

def make_ins(text, rev_id):
    ins = OxmlElement('w:ins')
    ins.set(qn('w:id'), str(rev_id))
    ins.set(qn('w:author'), author)
    ins.set(qn('w:date'), date)
    run = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    run.append(t)
    ins.append(run)
    return ins

def make_del(text, rev_id):
    d = OxmlElement('w:del')
    d.set(qn('w:id'), str(rev_id))
    d.set(qn('w:author'), author)
    d.set(qn('w:date'), date)
    run = OxmlElement('w:r')
    t = OxmlElement('w:delText')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    run.append(t)
    d.append(run)
    return d

HOW TO APPLY CHANGES:

1. Iterate `doc.paragraphs` to find the paragraph(s) matching the location described
   in each change plan. Use substring matching (paragraph.text) to locate text.

2. To REPLACE text in a paragraph:
   - Clear existing runs: for each run in paragraph.runs, access run._element
   - Build new XML with make_del(old_text, id) + make_ins(new_text, id+1)
   - Append the elements to the paragraph._element

3. To DELETE a paragraph entirely:
   - p._element.getparent().remove(p._element)
   - Then add a make_del paragraph

4. To INSERT a new paragraph after an existing one:
   - Create a new paragraph element and add make_ins with the new text
   - Insert it after the reference paragraph in the body

5. For word-level changes within a paragraph:
   - Split the paragraph text to isolate the changed portion
   - Rebuild the paragraph with: unchanged runs + make_del(old) + make_ins(new) + unchanged runs

RULES:
- Output ONLY executable Python code. No markdown fences, no comments outside code, no explanation text.
- Do NOT import anything — everything you need is already available.
- Do NOT create a new Document — modify `doc` in-place.
- Do NOT save the document — just modify it.
- Apply ONLY the changes described in the plans. Do not change anything else.
- Use a sequential rev_id counter starting from 1.
- Each change plan may describe multiple small edits — apply all of them.
- If a change plan says to add text, INSERT it. If it says to remove text, DELETE it.
  If it says to change/replace text, DELETE the old and INSERT the new.
- CRITICAL: Ensure all strings are properly closed. Use triple-quotes for multi-line strings.
  Never put literal newlines inside single-quoted or double-quoted strings.
"""

_USER_TMPL = """\
ORIGINAL MANUSCRIPT:
---
{manuscript}
---

REVIEWER CHANGE PLANS ({plan_count} comments):
{formatted_plans}

Generate Python code that applies ALL these changes to `doc` with proper OOXML track changes.
Each change must be a separate w:ins/w:del so the user can accept/reject individually in Word.
"""


# ── Helpers ───────────────────────────────────────────────────────────────────

def _format_plans(plans: list[dict]) -> str:
    """Format finalized change plans into a readable block for the AI prompt."""
    parts: list[str] = []
    for p in plans:
        rev = p.get("reviewer_number", "?")
        cnum = p.get("comment_number", "?")
        block = (
            f"--- Change {rev}.{cnum} ---\n"
            f"Reviewer {rev}, Comment {cnum}:\n"
            f"  Original comment: {p.get('original_comment', '(none)')}\n"
            f"  Change plan: {p.get('current_plan', p.get('finalized_plan', ''))}\n"
            f"  Action taken: {p.get('action_taken', '')}\n"
            f"  Manuscript changes: {p.get('manuscript_changes', '{}')}\n"
        )
        parts.append(block)
    return "\n".join(parts) if parts else "(no changes)"


def _text_to_docx(text: str) -> docx.Document:
    """Convert markdown-ish text to a python-docx Document.

    Handles # / ## / ### headings and regular paragraphs.
    """
    doc = docx.Document()

    for line in text.split('\n'):
        stripped = line.rstrip()

        if not stripped:
            doc.add_paragraph('')
            continue

        # Detect markdown headings
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2)
            doc.add_heading(heading_text, level=min(level, 9))
        else:
            doc.add_paragraph(stripped)

    return doc


def _extract_code(response: str) -> str:
    """Extract Python code from AI response, stripping markdown fences if present."""
    # Remove ```python ... ``` fences
    code = re.sub(r'^```(?:python)?\s*\n', '', response.strip())
    code = re.sub(r'\n```\s*$', '', code)
    return code.strip()


# ── Main function ─────────────────────────────────────────────────────────────

async def generate_track_changes_docx(
    provider: "AIProvider",
    original_manuscript: str,
    finalized_plans: list[dict],
    author_name: str = "Author",
    original_docx_bytes: bytes | None = None,
) -> bytes:
    """Generate a .docx with OOXML track changes by asking AI to write python-docx code.

    Args:
        provider: The user's configured AI provider.
        original_manuscript: Full original manuscript text (for the AI prompt context).
        finalized_plans: Per-comment finalized change plans with action_taken, manuscript_changes.
        author_name: Name to use as the track changes author in Word.
        original_docx_bytes: Raw .docx bytes from import. If provided, the AI modifies
            the real document (preserving all formatting, styles, references).
            If None, falls back to converting markdown text to a plain .docx.

    Returns:
        .docx file bytes with track changes markup.

    Raises:
        RuntimeError: If code generation or execution fails.
    """
    if not finalized_plans:
        raise RuntimeError("No finalized change plans provided. Finalize at least one comment first.")

    # Build the prompt
    formatted_plans = _format_plans(finalized_plans)
    user_prompt = _USER_TMPL.format(
        manuscript=original_manuscript,
        plan_count=len(finalized_plans),
        formatted_plans=formatted_plans,
    )

    # Ask AI for python-docx code (with retry on syntax/execution errors)
    MAX_ATTEMPTS = 2
    code = ""
    last_error = ""

    for attempt in range(1, MAX_ATTEMPTS + 1):
        if attempt == 1:
            prompt = user_prompt
        else:
            # Retry: send the error back to the AI so it can fix its code
            prompt = (
                f"Your previous code had an error:\n\n"
                f"```\n{last_error}\n```\n\n"
                f"Here is the failing code (first 3000 chars):\n"
                f"```python\n{code[:3000]}\n```\n\n"
                f"Fix the error and regenerate the COMPLETE code. "
                f"Remember: all strings must be properly closed, use triple-quotes "
                f"for any multi-line text.\n\n"
                f"{user_prompt}"
            )

        logger.info("Requesting track changes code from AI (attempt %d/%d, %d plans)",
                     attempt, MAX_ATTEMPTS, len(finalized_plans))
        raw_response = await provider.complete(
            system=_SYSTEM,
            user=prompt,
            temperature=0.1,
            max_tokens=16384,
        )

        code = _extract_code(raw_response)
        if not code:
            last_error = "AI returned empty code."
            logger.warning("Attempt %d: empty code from AI", attempt)
            continue

        logger.debug("AI-generated track changes code (attempt %d):\n%s", attempt, code[:2000])

        # Syntax check before execution
        try:
            compile(code, "<track_changes>", "exec")
        except SyntaxError as exc:
            last_error = f"SyntaxError: {exc}"
            logger.warning("Attempt %d: syntax error in AI code: %s", attempt, exc)
            continue

        # Syntax OK — break out to execute
        break
    else:
        # All attempts exhausted
        raise RuntimeError(f"AI-generated code failed after {MAX_ATTEMPTS} attempts: {last_error}")

    # Load document: use saved .docx if available (preserves formatting), else convert text
    if original_docx_bytes:
        doc = docx.Document(io.BytesIO(original_docx_bytes))
        logger.info("Using original .docx file (%d bytes) — formatting preserved", len(original_docx_bytes))
    else:
        doc = _text_to_docx(original_manuscript)
        logger.info("No .docx file saved — converting markdown to plain .docx")
    now_str = datetime.now(timezone.utc).strftime('%Y-%m-%dT%H:%M:%SZ')

    # Execute in restricted sandbox
    safe_globals = {
        '__builtins__': {
            'len': len, 'range': range, 'enumerate': enumerate,
            'str': str, 'int': int, 'float': float, 'bool': bool,
            'list': list, 'dict': dict, 'tuple': tuple, 'set': set,
            'print': lambda *a, **kw: None,  # silent
            'isinstance': isinstance, 'hasattr': hasattr, 'getattr': getattr,
            'setattr': setattr,
            'min': min, 'max': max, 'sorted': sorted, 'reversed': reversed,
            'zip': zip, 'map': map, 'filter': filter, 'any': any, 'all': all,
            'ValueError': ValueError, 'TypeError': TypeError, 'KeyError': KeyError,
            'IndexError': IndexError, 'AttributeError': AttributeError,
            'Exception': Exception, 'StopIteration': StopIteration,
            'True': True, 'False': False, 'None': None,
        },
        'doc': doc,
        'author': author_name,
        'date': now_str,
        'docx': docx,
        'qn': qn,
        'OxmlElement': OxmlElement,
        're': re,
        'datetime': datetime,
        'timezone': timezone,
    }

    try:
        exec(code, safe_globals)  # noqa: S102 — sandboxed execution
    except Exception as exc:
        logger.error("AI-generated code execution failed:\n%s\n\nError: %s", code[:3000], exc)
        raise RuntimeError(f"AI-generated code failed to execute: {exc}") from exc

    # Save the modified document to bytes
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
