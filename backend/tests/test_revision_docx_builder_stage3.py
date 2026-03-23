import io
import json
import zipfile

import docx
from docx.oxml.ns import qn
from lxml import etree

from services.revision_docx_builder import (
    accept_all_changes,
    apply_direct_track_changes,
    build_point_by_point_docx,
    build_track_changes_docx,
    prepare_revision_manuscript_docx,
)


def test_point_by_point_docx_accepts_copy_paste_and_citations():
    revision_round = {
        "journal_name": "Test Journal",
        "responses": [
            {
                "reviewer_number": 1,
                "comment_number": 1,
                "original_comment": "Please improve clarity in the introduction.",
                "author_response": "We revised the introduction for clarity.",
                "action_taken": "Introduction, paragraph 2, Lines 30-36: clarified rationale.",
                "manuscript_diff": '{"deleted":"Old sentence.","added":"New clearer sentence."}',
                "copy_paste_text": "This study addresses the gap by...",
                "citation_suggestions": ["10.1000/testdoi", "Smith et al., 2021"],
            }
        ],
    }
    data = build_point_by_point_docx(revision_round, manuscript_title="Title")
    assert isinstance(data, (bytes, bytearray))
    assert len(data) > 200


# ── Helper to build a .docx with tracked changes ─────────────────────────────

def _make_docx_with_tracked_changes():
    """Create a small .docx with w:ins and w:del elements for testing."""
    doc = docx.Document()

    # Paragraph 1: unchanged
    doc.add_paragraph("This paragraph is unchanged.")

    # Paragraph 2: has a deletion and an insertion
    p = doc.add_paragraph()
    body = doc.element.body

    # Add a normal run
    r1 = p._element.makeelement(qn('w:r'), {})
    t1 = r1.makeelement(qn('w:t'), {})
    t1.text = "The study used "
    t1.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r1.append(t1)
    p._element.append(r1)

    # Add w:del
    del_elem = p._element.makeelement(qn('w:del'), {
        qn('w:id'): '1',
        qn('w:author'): 'Amit',
        qn('w:date'): '2026-03-16T00:00:00Z',
    })
    r_del = del_elem.makeelement(qn('w:r'), {})
    dt = r_del.makeelement(qn('w:delText'), {})
    dt.text = "old method"
    dt.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r_del.append(dt)
    del_elem.append(r_del)
    p._element.append(del_elem)

    # Add w:ins
    ins_elem = p._element.makeelement(qn('w:ins'), {
        qn('w:id'): '2',
        qn('w:author'): 'Amit',
        qn('w:date'): '2026-03-16T00:00:00Z',
    })
    r_ins = ins_elem.makeelement(qn('w:r'), {})
    t_ins = r_ins.makeelement(qn('w:t'), {})
    t_ins.text = "new approach"
    t_ins.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r_ins.append(t_ins)
    ins_elem.append(r_ins)
    p._element.append(ins_elem)

    # Add trailing run
    r2 = p._element.makeelement(qn('w:r'), {})
    t2 = r2.makeelement(qn('w:t'), {})
    t2.text = " for analysis."
    t2.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r2.append(t2)
    p._element.append(r2)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _extract_all_text(docx_bytes: bytes) -> str:
    """Extract all paragraph text from a .docx."""
    doc = docx.Document(io.BytesIO(docx_bytes))
    return "\n".join(p.text for p in doc.paragraphs)


def _read_docx_xml(docx_bytes: bytes, member: str) -> str:
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
        return zf.read(member).decode("utf-8", errors="ignore")


def _docx_members(docx_bytes: bytes) -> list[str]:
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
        return zf.namelist()


# ── Tests for accept_all_changes() ───────────────────────────────────────────

def test_accept_all_changes_removes_deletions():
    """w:del elements should be completely removed."""
    src = _make_docx_with_tracked_changes()
    result = accept_all_changes(src)
    text = _extract_all_text(result)
    assert "old method" not in text


def test_accept_all_changes_unwraps_insertions():
    """w:ins wrapper should be gone but inserted text survives."""
    src = _make_docx_with_tracked_changes()
    result = accept_all_changes(src)
    text = _extract_all_text(result)
    assert "new approach" in text

    # Verify no w:ins elements remain in the XML
    doc = docx.Document(io.BytesIO(result))
    ins_elements = doc.element.body.findall(f'.//{qn("w:ins")}')
    assert len(ins_elements) == 0


def test_accept_all_changes_mixed_replace():
    """After accepting, deleted text gone and inserted text present."""
    src = _make_docx_with_tracked_changes()
    result = accept_all_changes(src)
    text = _extract_all_text(result)
    assert "old method" not in text
    assert "new approach" in text
    assert "The study used " in text
    assert " for analysis." in text


def test_accept_all_changes_no_changes():
    """A doc without tracked changes should pass through unchanged."""
    doc = docx.Document()
    doc.add_paragraph("No changes here.")
    doc.add_paragraph("Still nothing.")
    buf = io.BytesIO()
    doc.save(buf)
    original = buf.getvalue()

    result = accept_all_changes(original)
    text = _extract_all_text(result)
    assert "No changes here." in text
    assert "Still nothing." in text


def test_accept_all_changes_no_del_or_ins_elements_remain():
    """After accepting, no w:del or w:ins elements should exist."""
    src = _make_docx_with_tracked_changes()
    result = accept_all_changes(src)
    doc = docx.Document(io.BytesIO(result))
    body = doc.element.body
    assert len(body.findall(f'.//{qn("w:del")}')) == 0
    assert len(body.findall(f'.//{qn("w:ins")}')) == 0


def test_accept_all_changes_valid_docx():
    """Result should be a valid .docx openable by python-docx."""
    src = _make_docx_with_tracked_changes()
    result = accept_all_changes(src)
    doc = docx.Document(io.BytesIO(result))
    assert len(doc.paragraphs) >= 2


def test_apply_direct_track_changes_enables_word_revision_settings():
    doc = docx.Document()
    doc.add_paragraph(
        "Machine learning has transformed NLP. Traditional methods relied on hand-crafted features."
    )
    doc.add_paragraph("We used a convolutional neural network for classification.")
    doc.add_paragraph("These models learn directly from raw text data.")
    buf = io.BytesIO()
    doc.save(buf)

    plans = [{
        "manuscript_changes": json.dumps([
            {
                "type": "replace",
                "find": "convolutional neural network",
                "replace_with": "transformer-based architecture",
            },
            {
                "type": "delete",
                "find": "Traditional methods relied on hand-crafted features.",
            },
            {
                "type": "insert_after",
                "anchor": "raw text data.",
                "text": " Furthermore, pretrained language models improved benchmark scores.",
            },
        ])
    }]

    result = apply_direct_track_changes(buf.getvalue(), plans, author="Reviewer-1")

    settings_xml = _read_docx_xml(result, "word/settings.xml")
    document_xml = _read_docx_xml(result, "word/document.xml")
    members = _docx_members(result)

    assert "trackRevisions" in settings_xml
    assert "revisionView" in settings_xml
    assert 'w:comments="0"' in settings_xml
    assert "lnNumType" in document_xml
    assert 'w:restart="continuous"' in document_xml
    assert "<w:ins" in document_xml
    assert "<w:del" in document_xml
    assert "commentRangeStart" not in document_xml
    assert "commentRangeEnd" not in document_xml
    assert "commentReference" not in document_xml
    assert "word/comments.xml" not in members
    assert "transformer-based architecture" in document_xml
    assert "Traditional methods relied on hand-crafted features." in document_xml


def test_build_track_changes_docx_enables_word_revision_settings():
    result = build_track_changes_docx(
        "The cat sat on the mat.",
        "The dog sat on the warm mat.",
        author="Reviewer-1",
    )

    settings_xml = _read_docx_xml(result, "word/settings.xml")
    document_xml = _read_docx_xml(result, "word/document.xml")
    members = _docx_members(result)

    assert "trackRevisions" in settings_xml
    assert "revisionView" in settings_xml
    assert 'w:comments="0"' in settings_xml
    assert "lnNumType" in document_xml
    assert 'w:restart="continuous"' in document_xml
    assert "<w:ins" in document_xml
    assert "<w:del" in document_xml
    assert "commentRangeStart" not in document_xml
    assert "commentRangeEnd" not in document_xml
    assert "commentReference" not in document_xml
    assert "word/comments.xml" not in members


def test_prepare_revision_manuscript_docx_enables_track_changes_and_line_numbers():
    doc = docx.Document()
    doc.add_paragraph("Introduction")
    doc.add_paragraph("This is the first paragraph of the manuscript.")
    buf = io.BytesIO()
    doc.save(buf)

    result = prepare_revision_manuscript_docx(buf.getvalue())

    settings_xml = _read_docx_xml(result, "word/settings.xml")
    document_xml = _read_docx_xml(result, "word/document.xml")

    assert "trackRevisions" in settings_xml
    assert "revisionView" in settings_xml
    assert "lnNumType" in document_xml
    assert 'w:restart="continuous"' in document_xml
    assert 'w:countBy="1"' in document_xml
