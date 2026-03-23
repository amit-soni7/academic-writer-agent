import io

import pytest
from docx import Document

from services.sr_protocol_generator import (
    generate_prisma_p_checklist_docx,
    generate_protocol_docx,
)


def _sample_protocol_text() -> str:
    long_background = " ".join(["Background evidence sentence."] * 420)
    return f"""# Administrative Information

Registry: PROSPERO

## Background and Rationale

{long_background}

## Objectives

This review will evaluate intervention effects on the primary outcome.

## Methods

### Eligibility Criteria

Adults with the target condition will be included.

### Information Sources

PubMed, OpenAlex, and trial registries will be searched.

### Search Strategy

The PubMed strategy will combine subject headings and free-text terms.

### Study Records

Records will be managed in Rayyan and screened in duplicate.

### Data Items

Study design, sample size, funding source, and outcome data will be extracted.

### Outcomes and Prioritization

Primary and secondary outcomes will be prespecified.

### Risk of Bias Assessment

Risk of bias will be assessed with RoB 2.

### Data Synthesis

Meta-analysis will be performed when studies are sufficiently homogeneous.

### Meta-Bias Assessment

Publication bias will be assessed with funnel plots and Egger's test.

### Confidence in Evidence

GRADE will be used to assess certainty of evidence.
"""


def _sample_protocol_text_with_data_item_subheadings() -> str:
    return """# Administrative Information

Registry: PROSPERO

## Background and Rationale

This protocol addresses an important evidence gap.

## Objectives

This review will evaluate intervention effects on the primary outcome.

## Methods

### Data Items

#### 4.5.1 Study Characteristics

Study design, country, setting, and funding source will be extracted.

#### 4.5.2 Participant Characteristics

Age, sex, baseline risk profile, and eligibility features will be extracted.
"""


@pytest.mark.asyncio
async def test_generate_protocol_docx_uses_review_title_and_page_field():
    payload = await generate_protocol_docx(
        protocol_text=_sample_protocol_text(),
        pico={
            "population": "Adults with hypertension",
            "intervention": "Intervention X",
            "comparator": "Usual care",
            "outcome": "Blood pressure",
            "study_design": "Randomized trials",
        },
        prisma_p_data={"administrative": {"review_title": "Effects of Intervention X on Blood Pressure"}},
    )

    doc = Document(io.BytesIO(payload))
    assert doc.paragraphs[0].text == "Effects of Intervention X on Blood Pressure"
    assert "PAGE" in doc.sections[0].footer._element.xml


@pytest.mark.asyncio
async def test_generate_protocol_docx_preserves_level_four_headings():
    payload = await generate_protocol_docx(
        protocol_text=_sample_protocol_text_with_data_item_subheadings(),
        pico={
            "population": "Adults with hypertension",
            "intervention": "Intervention X",
            "comparator": "Usual care",
            "outcome": "Blood pressure",
            "study_design": "Randomized trials",
        },
        prisma_p_data={"administrative": {"review_title": "Effects of Intervention X on Blood Pressure"}},
    )

    doc = Document(io.BytesIO(payload))
    heading_styles = {
        paragraph.text: paragraph.style.name
        for paragraph in doc.paragraphs
        if paragraph.text in {"4.5.1 Study Characteristics", "4.5.2 Participant Characteristics"}
    }

    assert heading_styles["4.5.1 Study Characteristics"] == "Heading 4"
    assert heading_styles["4.5.2 Participant Characteristics"] == "Heading 4"


@pytest.mark.asyncio
async def test_generate_prisma_p_checklist_docx_includes_page_number_column():
    payload = await generate_prisma_p_checklist_docx(
        protocol_text=_sample_protocol_text(),
        pico={"population": "Adults", "intervention": "Intervention X", "outcome": "Blood pressure"},
        prisma_p_data={"administrative": {"review_title": "Protocol Checklist Export"}},
    )

    doc = Document(io.BytesIO(payload))
    table = doc.tables[0]
    headers = [cell.text for cell in table.rows[0].cells]
    assert headers == [
        "Section and topic",
        "Item No",
        "Checklist item",
        "Reported on page No.",
    ]

    objectives_row = next(row for row in table.rows if len(row.cells) >= 4 and row.cells[1].text == "7")
    rationale_row = next(row for row in table.rows if len(row.cells) >= 4 and row.cells[1].text == "6")
    identification_row = next(row for row in table.rows if len(row.cells) >= 4 and row.cells[1].text == "1a")

    assert identification_row.cells[3].text == "1"
    assert rationale_row.cells[3].text.isdigit()
    assert objectives_row.cells[3].text.isdigit()
    assert int(objectives_row.cells[3].text) >= int(rationale_row.cells[3].text)
