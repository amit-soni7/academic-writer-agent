from pathlib import Path
from types import SimpleNamespace

import docx
import pytest

from services.docx_pdf_converter import convert_docx_to_pdf


def test_convert_docx_to_pdf_moves_generated_pdf_to_target_path(tmp_path: Path, monkeypatch: pytest.MonkeyPatch):
    source_docx = tmp_path / "manuscript.docx"
    doc = docx.Document()
    doc.add_paragraph("Sample manuscript paragraph.")
    doc.save(source_docx)
    output_pdf = tmp_path / "exports" / "manuscript_reference.pdf"

    monkeypatch.setattr("services.docx_pdf_converter._resolve_soffice_binary", lambda: "/usr/bin/soffice")

    def fake_run(cmd, capture_output, text, check, timeout):
        outdir = Path(cmd[cmd.index("--outdir") + 1])
        (outdir / "manuscript.pdf").write_bytes(b"%PDF-1.4 fake")
        return SimpleNamespace(returncode=0, stdout="ok", stderr="")

    monkeypatch.setattr("services.docx_pdf_converter.subprocess.run", fake_run)

    result_path = convert_docx_to_pdf(str(source_docx), str(output_pdf))

    assert Path(result_path) == output_pdf
    assert output_pdf.read_bytes() == b"%PDF-1.4 fake"


def test_convert_docx_to_pdf_falls_back_when_libreoffice_fails(tmp_path: Path, monkeypatch: pytest.MonkeyPatch):
    source_docx = tmp_path / "manuscript.docx"
    doc = docx.Document()
    doc.add_paragraph("Fallback PDF line one.")
    doc.add_paragraph("Fallback PDF line two.")
    doc.save(source_docx)
    output_pdf = tmp_path / "exports" / "manuscript_reference.pdf"

    monkeypatch.setattr("services.docx_pdf_converter._resolve_soffice_binary", lambda: "/usr/bin/soffice")

    def fake_run(cmd, capture_output, text, check, timeout):
        return SimpleNamespace(returncode=1, stdout="", stderr="conversion failed")

    monkeypatch.setattr("services.docx_pdf_converter.subprocess.run", fake_run)

    result_path = convert_docx_to_pdf(str(source_docx), str(output_pdf))

    assert Path(result_path) == output_pdf
    assert output_pdf.read_bytes().startswith(b"%PDF-")


def test_convert_docx_to_pdf_raises_when_source_missing(tmp_path: Path):
    with pytest.raises(RuntimeError, match="Source .docx not found"):
        convert_docx_to_pdf(
            str(tmp_path / "missing.docx"),
            str(tmp_path / "exports" / "manuscript_reference.pdf"),
        )
