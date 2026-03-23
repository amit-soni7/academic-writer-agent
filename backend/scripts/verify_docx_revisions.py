#!/usr/bin/env python3
"""Generate and verify a tracked-changes DOCX via docx-revisions."""

from __future__ import annotations

import argparse
import io
import json
import sys
import zipfile
from pathlib import Path

import docx


BACKEND_DIR = Path(__file__).resolve().parents[1]
if str(BACKEND_DIR) not in sys.path:
    sys.path.insert(0, str(BACKEND_DIR))

from services.revision_docx_builder import apply_direct_track_changes


def build_sample_source_docx() -> bytes:
    doc = docx.Document()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(
        "Machine learning has transformed natural language processing. "
        "Traditional methods relied on hand-crafted features and statistical models."
    )
    doc.add_paragraph(
        "Deep learning approaches can learn representations directly from raw text data."
    )
    doc.add_heading("Methods", level=1)
    doc.add_paragraph(
        "We used a convolutional neural network for text classification."
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_tracked_docx() -> bytes:
    plans = [{
        "manuscript_changes": json.dumps([
            {
                "type": "replace",
                "find": "convolutional neural network",
                "replace_with": "transformer-based architecture",
            },
            {
                "type": "delete",
                "find": "Traditional methods relied on hand-crafted features and statistical models.",
            },
            {
                "type": "insert_after",
                "anchor": "raw text data.",
                "text": " Furthermore, pre-trained language models like BERT have set new benchmarks.",
            },
        ])
    }]
    return apply_direct_track_changes(
        build_sample_source_docx(),
        plans,
        author="Reviewer-1",
    )


def inspect_docx(docx_bytes: bytes) -> dict[str, object]:
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
        names = zf.namelist()
        settings_xml = zf.read("word/settings.xml").decode("utf-8", errors="ignore")
        document_xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")

    return {
        "insertions_present": "<w:ins" in document_xml,
        "deletions_present": "<w:del" in document_xml,
        "trackRevisions_present": "trackRevisions" in settings_xml,
        "revisionView_present": "revisionView" in settings_xml,
        "comments_disabled_in_revisionView": 'w:comments="0"' in settings_xml,
        "comments_part_present": "word/comments.xml" in names,
        "comment_references_present": any(
            token in document_xml
            for token in ("commentRangeStart", "commentRangeEnd", "commentReference")
        ),
        "ins_count": document_xml.count("<w:ins"),
        "del_count": document_xml.count("<w:del"),
    }


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Generate a sample tracked-changes DOCX and verify its XML."
    )
    parser.add_argument(
        "--output",
        default=str(Path("~/Downloads/docx_revisions_verification.docx").expanduser()),
        help="Path where the generated DOCX should be written.",
    )
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    output_path = Path(args.output).expanduser().resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)

    result = build_tracked_docx()
    output_path.write_bytes(result)
    checks = inspect_docx(result)
    checks["output_path"] = str(output_path)

    print(json.dumps(checks, indent=2))

    passed = all([
        checks["insertions_present"],
        checks["deletions_present"],
        checks["trackRevisions_present"],
        checks["revisionView_present"],
        checks["comments_disabled_in_revisionView"],
        not checks["comments_part_present"],
        not checks["comment_references_present"],
    ])
    return 0 if passed else 1


if __name__ == "__main__":
    raise SystemExit(main())
